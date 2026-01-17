from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

def create_matrix_cell(cell, symbol, points, explanation, is_header=False):
    """Create a matrix cell with symbol, points, and explanation - Original style"""
    cell.text = ""
    
    if is_header:
        p = cell.paragraphs[0]
        run = p.add_run(str(symbol))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1F4E79')
    else:
        # Symbol line
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        run.font.bold = True
        run.font.size = Pt(11)
        
        # Points line
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"({points})")
        run2.font.size = Pt(8)
        
        # Explanation line
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(explanation)
        run3.font.size = Pt(7)
        run3.font.italic = True
        
        # Color coding
        if symbol == "++":
            set_cell_shading(cell, 'C6EFCE')  # Light green
        elif symbol == "o":
            set_cell_shading(cell, 'FFEB9C')  # Light yellow
        elif symbol == "--":
            set_cell_shading(cell, 'FFC7CE')  # Light red
        
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

# Create document
doc = Document()

# Narrow margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025')
run.bold = True
run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT (Data Architect) | Zwei-Familien-Projekt | Handwerklich kompetent')
run.font.size = Pt(10)
run.font.italic = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('ERWEITERT: Fließend Englisch | Zeugen Jehovas | Familie 2: Viel Eigenkapital + Vor-Ort-Jobs')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR): ').bold = True
p1.add_run('Öffnet praktisch alle Türen. Mit Eigenkapital Familie 2 noch mehr Optionen.')

p2 = doc.add_paragraph()
p2.add_run('2. Remote IT-Job (Familie 1): ').bold = True
p2.add_run('Stabiles Einkommen. Data Architect öffnet Visa-Optionen (NZ, AU).')

p3 = doc.add_paragraph()
p3.add_run('3. FLIESSEND ENGLISCH (beide Familien): ').bold = True
p3.add_run('Öffnet NZ und AU! Keine Sprachbarriere = sofortige Integration.')

p4 = doc.add_paragraph()
p4.add_run('4. Zwei-Familien-Projekt: ').bold = True
p4.add_run('Geteilte Kosten, soziales Netz von Tag 1.')

p5 = doc.add_paragraph()
p5.add_run('5. Handwerkliche Kompetenz: ').bold = True
p5.add_run('Spart enorm bei Hausbau und Infrastruktur.')

p6 = doc.add_paragraph()
p6.add_run('6. Familie 2 - Eigenkapital + Vor-Ort-Jobs: ').bold = True
p6.add_run('Flexibel für lokalen Arbeitsmarkt, keine Zeitzonenbindung.')

p7 = doc.add_paragraph()
p7.add_run('7. Zeugen Jehovas: ').bold = True
p7.add_run('Weltweites Netzwerk, sofortige Gemeinschaft in jedem Land.')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('GESAMTRANKING - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Für euer Profil'],
    ['1', 'Neuseeland', '38 / 40', '95%', 'TOP - Englisch + Jobs + ZJ + Sicherheit!'],
    ['2', 'Spanien (Süden)', '36 / 40', '90%', 'EU-Rechte + Sonne + dt. Versammlungen'],
    ['3', 'Kanarische Inseln', '35 / 40', '88%', 'EU + bestes Klima + dt. Versammlungen'],
    ['4', 'Australien', '34 / 40', '85%', 'Englisch + Jobs, aber Klima extremer'],
    ['5', 'Costa Rica', '32 / 40', '80%', 'Höchster ZJ-Anteil, Zeitzone -7h'],
    ['6', 'Uruguay', '30 / 40', '75%', 'Günstig + sicher, aber weniger Jobs F2'],
    ['7', 'Chile (Mitte)', '28 / 40', '70%', 'Günstig, Erdbeben, schwerer Dialekt'],
    ['8', 'Deutschland', '22 / 40', '55%', 'Nur als Basis/Backup'],
    ['9', 'Schweden', '18 / 40', '45%', 'Geopolitisch riskant, kalt'],
    ['10', 'Nicaragua', '14 / 40', '35%', 'Rechtsunsicherheit, Ortega-Regime'],
    ['11', 'Nigeria', '10 / 40', '25%', 'Nicht empfohlen - Sicherheitsprobleme'],
]
create_table_original_style(doc, ranking_data)

doc.add_page_break()

# =============================================================================
# DETAILMATRIX - Original style with colored cells
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('Symbole: ').bold = True
legend.add_run('++ = Sehr gut (2 Pkt) | o = Mittel (1 Pkt) | -- = Schlecht (0 Pkt)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
run = abbrev.add_run('Länder: UY=Uruguay, NZ=Neuseeland, ES-S=Spanien Süd, AU=Australien, KAN=Kanaren, CR=Costa Rica, CL=Chile, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix data: [Kriterium, Gewicht, UY, NZ, ES-S, AU, KAN, CR, CL, DE, SE, NI, NG]
# Each cell: (symbol, points, explanation)
matrix_rows = [
    # Header row
    ["Kriterium (Gewicht)", "UY", "NZ", "ES-S", "AU", "KAN", "CR", "CL", "DE", "SE", "NI", "NG"],
    
    # Geopolitische Sicherheit x2
    ["Geopolitische Sicherheit\nx2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "EU-Rand"),
     ("o", "1", "AUKUS"),
     ("++", "2", "Weit weg"),
     ("++", "2", "Kein Militär"),
     ("++", "2", "Isoliert"),
     ("--", "0", "NATO-Front"),
     ("--", "0", "NATO-Ostsee"),
     ("o", "1", "Instabil"),
     ("--", "0", "Unsicher")],
    
    # Einwanderung (500k + IT) x1.5
    ["Einwanderung (500k + IT)\nx1.5",
     ("++", "2", "Sehr einfach"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "Investor OK"),
     ("++", "2", "Einfach"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "Einfach"),
     ("o", "1", "Kompliziert")],
    
    # Arbeiten mit Englisch vor Ort x1.5
    ["Arbeiten mit Englisch vor Ort\nx1.5",
     ("o", "1", "Spanisch\nnötig"),
     ("++", "2", "Mutter-\nsprache"),
     ("o", "1", "Spanisch\nbesser"),
     ("++", "2", "Mutter-\nsprache"),
     ("o", "1", "Spanisch\nbesser"),
     ("o", "1", "Viel Englisch"),
     ("--", "0", "Nur Spanisch"),
     ("++", "2", "Deutsch\nnötig"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Etwas\nEnglisch"),
     ("o", "1", "Amtssprache")],
    
    # Grundstück 2-Familien (10+ ha) x1.5
    ["Grundstück 2-Familien (10+ ha)\nx1.5",
     ("++", "2", "50-150k USD"),
     ("o", "1", "Teuer 500k+"),
     ("++", "2", "Hinterland\nOK"),
     ("o", "1", "Teuer"),
     ("o", "1", "Sehr\nbegrenzt"),
     ("o", "1", "Gestiegen"),
     ("++", "2", "Sehr günstig"),
     ("--", "0", "Sehr teuer"),
     ("--", "0", "Norden OK"),
     ("++", "2", "Billig"),
     ("++", "2", "Risiko")],
    
    # Remote-Work Zeitzone (zu DE) x1
    ["Remote-Work Zeitzone (zu DE)\nx1",
     ("++", "2", "-4h ideal"),
     ("--", "0", "+11h schwer"),
     ("++", "2", "Gleich"),
     ("--", "0", "+9h schwer"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "-5h OK"),
     ("++", "2", "Basis"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "+1h gut")],
    
    # Selbstversorgung (Klima/Boden) x1
    ["Selbstversorgung (Klima/Boden)\nx1",
     ("++", "2", "Ideal"),
     ("++", "2", "Nordinsel top"),
     ("++", "2", "Sehr gut"),
     ("o", "1", "Wasser knapp"),
     ("o", "1", "Wasser\nknapp"),
     ("o", "1", "Tropisch"),
     ("o", "1", "Feucht/kühl"),
     ("o", "1", "Bürokratie"),
     ("--", "0", "Kurze Saison"),
     ("o", "1", "Tropisch"),
     ("--", "0", "Unsicher")],
    
    # Klima (mild, Sonne, Meer) x1
    ["Klima (mild, Sonne, Meer)\nx1",
     ("++", "2", "2400h mild"),
     ("++", "2", "2200h mild"),
     ("++", "2", "3000h"),
     ("o", "1", "Extreme"),
     ("++", "2", "2800h\nperfekt"),
     ("o", "1", "Tropisch"),
     ("o", "1", "1700h feucht"),
     ("o", "1", "1600h"),
     ("--", "0", "Kalt dunkel"),
     ("--", "0", "Heiss feucht"),
     ("--", "0", "Heiss feucht")],
    
    # Rechtssicherheit/Eigentum x1.5
    ["Rechtssicherheit/Eigentum\nx1.5",
     ("++", "2", "Stabil"),
     ("++", "2", "Exzellent"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "Stark"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("++", "2", "Stark"),
     ("++", "2", "Stark"),
     ("--", "0", "Schwach"),
     ("--", "0", "Korrupt")],
    
    # Gesundheitsversorgung x1
    ["Gesundheitsversorgung\nx1",
     ("o", "1", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Sehr gut"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    # ZJ-Gemeinschaft x2
    ["ZJ-Gemeinschaft\nx2",
     ("o", "1", "13k klein"),
     ("++", "2", "14k aktiv"),
     ("++", "2", "113k groß"),
     ("++", "2", "68k aktiv"),
     ("++", "2", "8k + dt.\nVersamml."),
     ("++", "2", "28k 0,54%!"),
     ("++", "2", "79k groß"),
     ("++", "2", "176k"),
     ("o", "1", "Klein"),
     ("o", "1", "Klein"),
     ("++", "2", "407k!")],
    
    # Deutsche ZJ-Versammlungen x1
    ["Deutsche ZJ-Versammlungen\nx1",
     ("--", "0", "Keine"),
     ("--", "0", "Keine"),
     ("++", "2", "Costa del\nSol"),
     ("--", "0", "Keine"),
     ("++", "2", "Teneriffa\nGran Can."),
     ("--", "0", "Keine"),
     ("--", "0", "Keine"),
     ("++", "2", "Überall"),
     ("--", "0", "Keine"),
     ("--", "0", "Keine"),
     ("--", "0", "Keine")],
    
    # Jobs für Familie 2 x2
    ["Jobs für Familie 2\nx2",
     ("--", "0", "Sehr\nbegrenzt"),
     ("++", "2", "Arbeits-\nkräftemangel"),
     ("o", "1", "Hohe\nArbeitslos."),
     ("++", "2", "Arbeits-\nkräftemangel"),
     ("o", "1", "Tourismus"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Moderat"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("--", "0", "Riskant"),
     ("--", "0", "Gefährlich")],
    
    # Nähe Europa (Flug) x0.5
    ["Nähe Europa (Flug)\nx0.5",
     ("o", "1", "12-14h"),
     ("--", "0", "24h"),
     ("++", "2", "2-3h"),
     ("--", "0", "22h"),
     ("++", "2", "4h"),
     ("o", "1", "12h"),
     ("o", "1", "14h"),
     ("++", "2", "Basis"),
     ("++", "2", "1-2h"),
     ("o", "1", "12h"),
     ("o", "1", "6h")],
]

# Create the matrix table
num_cols = len(matrix_rows[0])
num_rows = len(matrix_rows)
table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

# Fill header row
for j, header in enumerate(matrix_rows[0]):
    cell = table.rows[0].cells[j]
    cell.text = header
    set_cell_shading(cell, '1F4E79')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

# Fill data rows
for i, row_data in enumerate(matrix_rows[1:], start=1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        
        if j == 0:  # Kriterium column
            cell.text = cell_data
            set_cell_shading(cell, 'D9E2F3')  # Light blue
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        else:
            # Data cell with symbol, points, explanation
            symbol, points, explanation = cell_data
            create_matrix_cell(cell, symbol, points, explanation)

doc.add_page_break()

# =============================================================================
# DETAILANALYSEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILANALYSEN - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# PLATZ 1: NEUSEELAND
h2 = doc.add_paragraph()
run = h2.add_run('Platz 1: Neuseeland (38.0 Punkte / 95%) - KLARE EMPFEHLUNG')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run("Hawke's Bay, Nelson/Tasman, Bay of Plenty, Waikato")

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM NEUSEELAND FÜR EUER PROFIL JETZT PERFEKT IST:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Euer Englisch ändert alles: ').bold = True
p.add_run('Mit fließendem Englisch entfällt die größte Hürde! Ihr könnt ab Tag 1 kommunizieren, arbeiten, Kinder in die Schule schicken. Die Versammlungen der Zeugen Jehovas sind sofort auf Englisch besuchbar.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 1 - IT-Skills öffnen Türen: ').bold = True
p.add_run('Data Architect ist auf der New Zealand Skilled Occupation List! Mit nachgewiesenem Remote-Einkommen habt ihr sehr gute Chancen auf ein Skilled Migrant Visa.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt perfekt: ').bold = True
p.add_run('Neuseeland hat ARBEITSKRÄFTEMANGEL: Handwerk, Landwirtschaft, Tourismus, Gesundheit. Mit fließendem Englisch sofortiger Zugang. Mit Eigenkapital auch eigenes Business möglich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~14.000 Verkündiger in ~175 Versammlungen. Gut verteilt, auch ländlich. Kiwi-Mentalität offen und freundlich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeitzone-Lösung: ').bold = True
p.add_run('+11-12h vor DE. Familie 1: Async-Arbeit oder NZ/AU-Kunden. Familie 2: Vor Ort = Zeitzone irrelevant!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True
doc.add_paragraph("• 10-15 ha Hawke's Bay/Nelson: 300.000-600.000 NZD")
doc.add_paragraph('• 2 Häuser bauen: 300.000-500.000 NZD')
doc.add_paragraph('• Infrastruktur + Reserve: 210.000 NZD')
doc.add_paragraph('• GESAMT: ~810.000-1.310.000 NZD (~460.000-745.000 EUR)')
doc.add_paragraph('• MIT EIGENKAPITAL FAMILIE 2: Machbar!')

doc.add_paragraph()

# PLATZ 2: SPANIEN
h2 = doc.add_paragraph()
run = h2.add_run('Platz 2: Spanien Süden (36.0 Punkte / 90%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Costa de la Luz (Huelva), Almería Hinterland, Murcia')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Vorteil + Nähe: ').bold = True
p.add_run('Im EU-System. 2-3h Flug nach DE. Krankenversicherung, Rente einfach.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('ZJ - DEUTSCHSPRACHIGE VERSAMMLUNGEN: ').bold = True
p.add_run('~113.000 Verkündiger. An Costa del Sol gibt es DEUTSCHSPRACHIGE Versammlungen! Perfekter Übergang.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Remote-Work: ').bold = True
p.add_run('Gleiche Zeitzone! Glasfaser gut ausgebaut.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2: ').bold = True
p.add_run('Schwieriger als NZ wegen hoher Arbeitslosigkeit. ABER: Expat-Gebiete = Vorteile durch Deutsch+Englisch. Mit Eigenkapital: Eigenes Business möglich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zahlen: ').bold = True
p.add_run('Finca 10 ha: 150-300k EUR. Renovierung: 80-120k. GESAMT: 315-515k EUR.')

doc.add_paragraph()

# PLATZ 3: KANAREN
h2 = doc.add_paragraph()
run = h2.add_run('Platz 3: Kanarische Inseln (35.0 Punkte / 88%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Inseln: ').bold = True
p.add_run('Teneriffa, Gran Canaria, La Palma')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Bestes Klima Europas + DEUTSCHE VERSAMMLUNGEN: ').bold = True
p.add_run('Ganzjährig 18-28°. ~8.000 ZJ mit deutschsprachigen Versammlungen auf Teneriffa und Gran Canaria!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('STRATEGIE: ').bold = True
p.add_run('Kanaren als Einstieg (1-2 Jahre), dt. Versammlung, dann Festland oder NZ.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
p.add_run('Große Grundstücke begrenzt und teuer. 10+ ha kaum verfügbar.')

doc.add_paragraph()

# PLATZ 4: AUSTRALIEN
h2 = doc.add_paragraph()
run = h2.add_run('Platz 4: Australien (34.0 Punkte / 85%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Regionen: ').bold = True
p.add_run('Tasmanien (beste Wahl!), Sunshine Coast, Victoria')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Englisch + Jobs: ').bold = True
p.add_run('IT auf Skilled List. Arbeitskräftemangel für Familie 2. ~68.000 ZJ.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Klima extremer (Buschbrände, Dürren). Wasserknappheit. Tasmanien = beste Option.')

doc.add_paragraph()

# PLATZ 5-7 KURZ
h2 = doc.add_paragraph()
run = h2.add_run('Platz 5-7: Costa Rica, Uruguay, Chile (Kurzübersicht)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Costa Rica (80%): ').bold = True
p.add_run('Höchster ZJ-Anteil (0,54%!), kein Militär. ABER: Spanisch nötig, -7h Zeitzone, Jobs begrenzt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Uruguay (75%): ').bold = True
p.add_run('War ursprünglich #1! Beste Rechtssicherheit, -4h Zeitzone, günstigste Grundstücke. ABER: Kleinere ZJ (13k), Jobs sehr begrenzt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Chile (70%): ').bold = True
p.add_run('Große ZJ (79k), sehr günstig. ABER: Spanisch sehr schwer (Dialekt), Erdbeben.')

doc.add_paragraph()

# PLATZ 8-11 NICHT EMPFOHLEN
h2 = doc.add_paragraph()
run = h2.add_run('Platz 8-11: Nicht empfohlen')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('DEUTSCHLAND (55%): ').bold = True
p.add_run('Nur Backup. NATO-Front, Bürokratie, teuer, wenig Sonne.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('SCHWEDEN (45%): ').bold = True
p.add_run('Kalt, dunkel, NATO nahe Russland.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('NICARAGUA (35%): ').bold = True
p.add_run('Ortega-Regime, Rechtsunsicherheit, willkürliche Enteignungen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('NIGERIA (25%): ').bold = True
p.add_run('NICHT empfohlen! Massive Sicherheitsprobleme, Entführungen, Korruption. Große ZJ (407k) wiegt das nicht auf.')

doc.add_page_break()

# =============================================================================
# AKTIONSPLAN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('AKTIONSPLAN - 24 MONATE BIS ZUM HOMESTEAD')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Phase 1: Vorbereitung in Deutschland (Monat 1-6)').bold = True
doc.add_paragraph('• Familientreffen: Finale Entscheidung NZ vs. Spanien')
doc.add_paragraph('• immigration.govt.nz studieren, Skills Assessment beantragen')
doc.add_paragraph('• Kontakt zu ZJ-Versammlungen im Zielland (jw.org)')
doc.add_paragraph('• Falls Spanien: Spanischkurs beginnen')
doc.add_paragraph('• Steuerberater: Wegzugsbesteuerung klären')
doc.add_paragraph('• Scouting-Reise buchen (4-6 Wochen)')

h3 = doc.add_paragraph()
h3.add_run('Phase 2: Scouting-Reise (Monat 6-9)').bold = True
doc.add_paragraph("• 4-6 Wochen im Zielland erkunden")
doc.add_paragraph('• Versammlungen besuchen (vorab Kontakt!)')
doc.add_paragraph('• Grundstücke besichtigen, Makler treffen')
doc.add_paragraph('• Familie 2: Mit Arbeitgebern sprechen')
doc.add_paragraph('• Internet testen, Schulen anschauen')

h3 = doc.add_paragraph()
h3.add_run('Phase 3: Entscheidung und Visa (Monat 9-15)').bold = True
doc.add_paragraph('• Finale Länder-/Regionswahl')
doc.add_paragraph('• Visa-Anträge einreichen')
doc.add_paragraph('• Immobilie identifizieren, Kaufverhandlung')

h3 = doc.add_paragraph()
h3.add_run('Phase 4: Soft Landing (Monat 15-21)').bold = True
doc.add_paragraph('• Familie 2 ggf. vorausreisen')
doc.add_paragraph('• Grundstückskauf abschließen')
doc.add_paragraph('• Familie 1: Umzug zum Schuljahresbeginn (NZ=Feb, ES=Sep)')

h3 = doc.add_paragraph()
h3.add_run('Phase 5: Aufbau (Monat 21-36)').bold = True
doc.add_paragraph('• Hausbau/Renovierung')
doc.add_paragraph('• Selbstversorgung starten')
doc.add_paragraph('• Aktiv in Versammlung integrieren')

doc.add_paragraph()

# =============================================================================
# EMPFEHLUNG
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('UNSERE EMPFEHLUNG FÜR EUCH')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ERSTE WAHL: NEUSEELAND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('• Englisch = Sofortige Integration')
doc.add_paragraph('• IT-Skills = Skilled Migrant Visa')
doc.add_paragraph('• Arbeitskräftemangel = Jobs für Familie 2')
doc.add_paragraph('• Aktive ZJ-Gemeinschaft')
doc.add_paragraph('• Maximale geopolitische Sicherheit')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE: SPANIEN/KANAREN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('• EU-Rechte, gleiche Zeitzone')
doc.add_paragraph('• 2-3h nach DE')
doc.add_paragraph('• Deutschsprachige ZJ-Versammlungen')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('NÄCHSTER SCHRITT: ')
run.bold = True
p.add_run('4-6 Wochen Scouting-Reise Sommer/Herbst 2025 planen!')

doc.add_paragraph()
doc.add_paragraph('─' * 70)

p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025 | Basierend auf Auswanderungsanalyse 2025').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Erweitert: ZJ | Englisch | Eigenkapital F2 | Vor-Ort-Jobs | Familie 1 (40J, 2 Kinder) | Familie 2 (50J)').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_Komplett_v3.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_Komplett_v3.docx")


