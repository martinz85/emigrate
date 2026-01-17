from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_styled_table(doc, data, header_color='2F5496'):
    """Create a table matching the original style"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            # Style header row
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    return table

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# =============================================================================
# TITLE - Same style as original
# =============================================================================
title = doc.add_heading('AUSWANDERUNGSANALYSE 2025', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 0, 0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

# Tagline
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT (Data Architect) | Zwei-Familien-Projekt | Handwerklich kompetent')
run.font.size = Pt(11)
run.font.italic = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('ERWEITERT: Fließend Englisch | Zeugen Jehovas | Familie 2: Eigenkapital + Vor-Ort-Jobs')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN (Same style as original)
# =============================================================================
h1 = doc.add_heading('EUER PROFIL - STÄRKEN', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph('1. Kapital (500k+ EUR): Öffnet praktisch alle Türen. Mit Eigenkapital Familie 2 sogar noch mehr Optionen.')
doc.add_paragraph('2. Remote IT-Job (Familie 1): Stabiles Einkommen unabhängig vom lokalen Arbeitsmarkt. Data Architect ist gefragter Beruf, der auch Visa-Optionen öffnet (NZ, AU).')
doc.add_paragraph('3. FLIESSEND ENGLISCH (beide Familien): Öffnet englischsprachige Länder! Neuseeland und Australien werden realistisch.')
doc.add_paragraph('4. Zwei-Familien-Projekt: Geteilte Kosten, geteilte Arbeit, soziales Netz von Tag 1.')
doc.add_paragraph('5. Handwerkliche Kompetenz: Spart enorm bei Hausbau und Infrastruktur.')
doc.add_paragraph('6. Familie 2 - Vor-Ort-Jobs: Flexibel für lokalen Arbeitsmarkt, keine Zeitzonenbindung.')
doc.add_paragraph('7. Zeugen Jehovas: Weltweites Netzwerk, sofortige Gemeinschaft in jedem Land.')

doc.add_paragraph()

# =============================================================================
# FAMILIENPROFIL (New section)
# =============================================================================
h1 = doc.add_heading('FAMILIENPROFIL', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

family_data = [
    ['', 'Familie 1', 'Familie 2'],
    ['Alter', 'ca. 40 Jahre', 'ca. 50 Jahre'],
    ['Kinder', '2 (12 und 14 Jahre)', 'Keine'],
    ['Beruf/Einkommen', 'Remote IT (Data Architect)', 'Vor-Ort-Jobs (kein IT)'],
    ['Kapital', 'Anteil an 500k gemeinsam', 'Eigenes Kapital zusätzlich'],
    ['Sprachen', 'Deutsch + Englisch fließend', 'Deutsch + Englisch fließend'],
    ['Zeitzone-Bindung', 'Für Remote-Work relevant', 'Keine (lokale Jobs)'],
    ['Priorität', 'Schulsystem für Kinder', 'Arbeitsmarkt vor Ort'],
]
create_styled_table(doc, family_data)

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING - Same style as original
# =============================================================================
h1 = doc.add_heading('GESAMTRANKING - ANGEPASST AN EUER PROFIL', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Für euer Profil'],
    ['1', 'Neuseeland', '38 / 40', '95%', 'TOP - Englisch + Jobs + ZJ!'],
    ['2', 'Spanien (Süden/Kanaren)', '36 / 40', '90%', 'EU-Rechte + dt. Versammlungen'],
    ['3', 'Australien', '35 / 40', '88%', 'Englisch + Jobs, Klima extremer'],
    ['4', 'Costa Rica', '33 / 40', '83%', 'Höchster ZJ-Anteil, Zeitzone -7h'],
    ['5', 'Uruguay', '31 / 40', '78%', 'Günstig, kleinere ZJ-Gemeinschaft'],
    ['6', 'Chile (Mitte)', '30 / 40', '75%', 'Günstig, Erdbeben, schwerer Dialekt'],
    ['7', 'Deutschland', '22 / 40', '55%', 'Nur als Basis/Backup'],
    ['8', 'Schweden', '18 / 40', '45%', 'Geopolitisch riskant'],
]
create_styled_table(doc, ranking_data, '1F4E79')

doc.add_paragraph()

# =============================================================================
# DETAILMATRIX - Same style as original with symbols
# =============================================================================
h1 = doc.add_heading('DETAILMATRIX MIT BEWERTUNG', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

legend = doc.add_paragraph()
legend.add_run('Legende: ').bold = True
legend.add_run('✓ = Sehr gut (2 Pkt) | ⚠ = Mittel (1 Pkt) | ✗ = Schlecht (0 Pkt)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
abbrev.add_run('NZ=Neuseeland, ES=Spanien/Kanaren, AU=Australien, CR=Costa Rica, UY=Uruguay, CL=Chile').font.size = Pt(9)

doc.add_paragraph()

# Matrix table - same style as original
matrix_data = [
    ['Kriterium', 'Gewicht', 'NZ', 'ES', 'AU', 'CR', 'UY', 'CL'],
    ['Geopolitische Sicherheit', 'x2', '✓ Maximal isoliert', '⚠ Rand Europas', '✓ AUKUS-Schutz', '✓ Kein Militär', '✓ Neutral', '⚠ Stabil'],
    ['Einwanderung mit Kapital', 'x1.5', '✓ IT auf Liste!', '✓ EU-Recht', '✓ IT auf Liste!', '✓ Investor-Visum', '✓ Sehr einfach', '✓ Einfach'],
    ['SPRACHE (Englisch!)', 'x2', '✓ ENGLISCH!', '⚠ Spanisch lernen', '✓ ENGLISCH!', '⚠ Spanisch lernen', '⚠ Spanisch lernen', '⚠ Spanisch schwer'],
    ['Zeugen Jehovas Gemeinschaft', 'x2', '✓ 14k, aktiv', '✓ 113k, dt. Vers.!', '✓ 68k, aktiv', '✓ 28k, 0.54%!', '⚠ 13k, kleiner', '✓ 79k, groß'],
    ['Jobs für Familie 2', 'x2', '✓ Arbeitskräftemangel', '⚠ Hohe Arbeitslos.', '✓ Arbeitskräftemangel', '⚠ Begrenzt', '✗ Sehr begrenzt', '⚠ Moderat'],
    ['Remote-Work Zeitzone (F1)', 'x1', '✗ +11-12h zu DE', '✓ Gleiche TZ', '⚠ +8-10h zu DE', '⚠ -7h zu DE', '✓ -4h zu DE', '⚠ -5h zu DE'],
    ['Grundstück 10+ ha', 'x1.5', '⚠ Teuer 500k+', '✓ Günstig Hinterland', '⚠ Teuer', '✓ 150-300k USD', '✓ 80-150k USD', '✓ Sehr günstig'],
    ['Selbstversorgung (Klima)', 'x1', '✓ Ideal, mild', '✓ Sehr gut', '⚠ Wasserprobleme', '⚠ Tropisch', '✓ Ideal', '⚠ Feucht/kühl'],
    ['Klima (Sonne, Meer)', 'x1', '✓ 2200h, mild', '✓ 2800-3000h', '⚠ Extreme möglich', '⚠ Tropisch', '✓ 2400h, mild', '⚠ 1700h, feucht'],
    ['Rechtssicherheit', 'x1.5', '✓ Exzellent', '✓ EU-Standard', '✓ Exzellent', '✓ Gut', '✓ Sehr stabil', '⚠ OK'],
    ['Nähe Europa (Flug)', 'x0.5', '✗ 24h+', '✓ 2-4h', '✗ 22h+', '⚠ 12h', '⚠ 12-14h', '⚠ 14h'],
    ['Schulen für Kinder', 'x1', '✓ Exzellent', '✓ Gut + Förderung', '✓ Sehr gut', '⚠ Privatschulen', '⚠ OK', '⚠ OK'],
]
create_styled_table(doc, matrix_data, '2F5496')

doc.add_page_break()

# =============================================================================
# DETAILANALYSEN - Same style as original
# =============================================================================
h1 = doc.add_heading('DETAILANALYSEN - ANGEPASST AN EUER PROFIL', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

# -----------------------------------------------------------------------------
# PLATZ 1: NEUSEELAND
# -----------------------------------------------------------------------------
h2 = doc.add_heading('Platz 1: Neuseeland (38.0 Punkte / 95%) - KLARE EMPFEHLUNG', 2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run("Hawke's Bay, Nelson/Tasman, Bay of Plenty, Waikato")

doc.add_paragraph()
h3 = doc.add_heading('WARUM NEUSEELAND FÜR EUER PROFIL JETZT PERFEKT IST:', 3)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Euer Englisch ändert alles: ').bold = True
p.add_run('Mit fließendem Englisch entfällt die größte Hürde! Ihr könnt ab Tag 1 kommunizieren, arbeiten, Kinder in die Schule schicken. Die Versammlungen der Zeugen Jehovas sind sofort auf Englisch besuchbar.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Familie 1 - IT-Skills öffnen Türen: ').bold = True
p.add_run('Data Architect ist auf der New Zealand Skilled Occupation List! Mit nachgewiesenem Remote-Einkommen und euren Qualifikationen habt ihr sehr gute Chancen auf ein Skilled Migrant Visa.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt ist perfekt: ').bold = True
p.add_run('Neuseeland hat ARBEITSKRÄFTEMANGEL in vielen Bereichen: Handwerk, Landwirtschaft, Tourismus, Gesundheit. Mit fließendem Englisch ist der Zugang zum Arbeitsmarkt sofort möglich. Mit Eigenkapital könnt ihr auch ein eigenes kleines Business starten.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas in Neuseeland: ').bold = True
p.add_run('~14.000 aktive Verkündiger in ~175 Versammlungen. Gut verteilt, auch in ländlichen Gebieten. Die Kiwi-Mentalität ist offen und freundlich. Zweigbüro in Auckland gut organisiert.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Das Zeitzonenproblem - und die Lösung: ').bold = True
p.add_run('Neuseeland ist +11-12h vor Deutschland. Für Familie 1 (Remote IT): Async-Arbeit vereinbaren oder neuseeländische/australische Kunden aufbauen. Familie 2 arbeitet VOR ORT - Zeitzone ist irrelevant!')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Konkrete Zahlen für euer Projekt:').bold = True

budget_nz = doc.add_paragraph()
budget_nz.add_run("""
• 10-15 Hektar in Hawke's Bay oder Nelson: ca. 300.000-600.000 NZD
• Hausbau (2 Familienhäuser, je 120m², Qualität): ca. 300.000-500.000 NZD
• Solar-Anlage + Wassersystem: ca. 50.000 NZD
• Fahrzeuge: ca. 60.000 NZD
• Notreserve: 100.000 NZD
• GESAMT: ca. 810.000-1.310.000 NZD (~460.000-745.000 EUR)
• MIT EIGENKAPITAL FAMILIE 2: Machbar mit guter Reserve!
""")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Entfernung zu Europa (24h+ Flug). Zeitzone für Remote-Work nach DE schwierig. Grundstücke teurer als Südamerika.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 2: SPANIEN
# -----------------------------------------------------------------------------
h2 = doc.add_heading('Platz 2: Spanien Süden/Kanaren (36.0 Punkte / 90%)', 2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Costa de la Luz (Huelva), Almería Hinterland, Murcia, Kanaren (Teneriffa, Gran Canaria)')

doc.add_paragraph()
h3 = doc.add_heading('WARUM SPANIEN FÜR EUER PROFIL SEHR GUT IST:', 3)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('EU-Vorteil + Nähe: ').bold = True
p.add_run('Ihr bleibt im EU-System mit allen Rechten. Flug nach Deutschland in 2-3 Stunden. Krankenversicherung, Rentenansprüche - alles bleibt einfach.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas - DEUTSCHSPRACHIGE VERSAMMLUNGEN: ').bold = True
p.add_run('~113.000 Verkündiger in ~1.500 Versammlungen. BESONDERHEIT: Auf den Kanaren und an der Costa del Sol gibt es DEUTSCHSPRACHIGE Versammlungen! Perfekt für einen sanften Übergang.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Sprache: ').bold = True
p.add_run('Spanisch ist für Deutsche relativ leicht zu lernen (12-18 Monate bis B2). Euer Englisch hilft in Touristenregionen als Brücke.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt: ').bold = True
p.add_run('Schwieriger als in NZ/AU wegen hoher Arbeitslosigkeit in Spanien. ABER: In Expat-Gebieten Vorteile durch Deutsch + Englisch (Tourismus, Handwerk). Mit Eigenkapital: Eigenes Business möglich (B&B, Finca).')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True

budget_es = doc.add_paragraph()
budget_es.add_run("""
• Finca 10-15 ha mit Haus (renovierungsbedürftig) im Hinterland: 150.000-250.000 EUR
• Renovierung/Ausbau für zwei Familien: 80.000-120.000 EUR
• Solar + Infrastruktur: 25.000-35.000 EUR
• Reserve: 60.000+ EUR
• GESAMT: ca. 315.000-465.000 EUR
• Bei 500k+ Budget: Gute Reserve vorhanden!
""")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Strategie-Option: ').bold = True
p.add_run('Kanaren als Einstieg (deutschsprachige Versammlung, gleiche Zeitzone, beste Klima) → nach 1-2 Jahren auf Festland oder nach NZ.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 3: AUSTRALIEN
# -----------------------------------------------------------------------------
h2 = doc.add_heading('Platz 3: Australien (35.0 Punkte / 88%)', 2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Tasmanien, Sunshine Coast (Queensland), Victoria')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Euer Englisch öffnet Türen: ').bold = True
p.add_run('Data Architect ist auch auf der Australian Skilled Occupation List. Mit fließendem Englisch sofortiger Zugang zu allem.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~68.000 Verkündiger in ~790 Versammlungen. Gut organisierte, aktive Gemeinschaft. Zweigbüro in Sydney.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Familie 2 - SUPER Arbeitsmarkt: ').bold = True
p.add_run('Großer Arbeitskräftemangel! Handwerk sehr gefragt mit guten Löhnen. Mindestlohn ~24 AUD/Stunde - einer der höchsten weltweit.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Klima extremer als NZ (Buschbrände, Dürren, Hitze). Wasserknappheit für Selbstversorger problematisch (außer Tasmanien). Zeitzone +8-10h zu DE.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('EMPFEHLUNG: ').bold = True
p.add_run('Tasmanien = Beste Region für Homestead. Klima ähnlich NZ, weniger Extreme, günstiger als Festland.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 4-6: KURZÜBERSICHT
# -----------------------------------------------------------------------------
h2 = doc.add_heading('Platz 4-6: Lateinamerika (Kurzübersicht)', 2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Costa Rica (33 Punkte / 83%): ').bold = True
p.add_run('Höchster ZJ-Bevölkerungsanteil (0,54%!). Sehr aktive Gemeinschaft. Kein Militär, stabile Demokratie. ABER: Spanisch nötig, Zeitzone -7h, Arbeitsmarkt begrenzt für Familie 2.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Uruguay (31 Punkte / 78%): ').bold = True
p.add_run('War ursprünglich Platz 1! Beste Rechtssicherheit, perfekte Zeitzone (-4h), günstigste Grundstücke. ABER: Kleinere ZJ-Gemeinschaft (~13.000), begrenzter Arbeitsmarkt für Familie 2.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Chile (30 Punkte / 75%): ').bold = True
p.add_run('Große ZJ-Gemeinschaft (~79.000), sehr günstige Grundstücke. ABER: Chilenisches Spanisch sehr schwer zu verstehen, Erdbebengefahr (höhere Baukosten).')

doc.add_page_break()

# =============================================================================
# JOBS FÜR FAMILIE 2 - VERGLEICH
# =============================================================================
h1 = doc.add_heading('VOR-ORT-JOBS FÜR FAMILIE 2 - LÄNDERVERGLEICH', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

jobs_data = [
    ['Land', 'Arbeitsmarkt', 'Sprache', 'Typische Jobs', 'Bewertung'],
    ['Neuseeland', 'Arbeitskräftemangel', 'Englisch ✓', 'Handwerk, Landwirtschaft, Tourismus', '✓ IDEAL'],
    ['Australien', 'Arbeitskräftemangel', 'Englisch ✓', 'Handwerk, Bergbau, Gesundheit', '✓ IDEAL'],
    ['Spanien', 'Hohe Arbeitslosigkeit', 'Spanisch nötig', 'Tourismus, Handwerk (Expat-Gebiete)', '⚠ Mittel'],
    ['Costa Rica', 'Begrenzt', 'Spanisch hilft', 'Tourismus, Eco-Lodges', '⚠ Mittel'],
    ['Chile', 'Moderat', 'Spanisch nötig', 'Bergbau, Landwirtschaft', '⚠ Mittel'],
    ['Uruguay', 'Sehr begrenzt', 'Spanisch nötig', 'Tourismus, Landwirtschaft', '✗ Schwierig'],
]
create_styled_table(doc, jobs_data)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('FAZIT: ').bold = True
p.add_run('Mit fließendem Englisch sind Neuseeland und Australien für Familie 2 die klaren Gewinner!')

doc.add_paragraph()

# =============================================================================
# AKTIONSPLAN - Same style as original
# =============================================================================
h1 = doc.add_heading('AKTIONSPLAN - 24 MONATE BIS ZUM HOMESTEAD', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

# Phase 1
h3 = doc.add_heading('Phase 1: Vorbereitung in Deutschland (Monat 1-6)', 3)
doc.add_paragraph('• Familientreffen: Finale Entscheidung NZ vs. Spanien treffen')
doc.add_paragraph('• Neuseeland Immigration Website studieren (immigration.govt.nz)')
doc.add_paragraph('• Skills Assessment für Data Architect beantragen')
doc.add_paragraph('• Kontakt zu ZJ-Versammlungen in Zielland aufnehmen (jw.org)')
doc.add_paragraph('• Immigration Advisor konsultieren (für NZ empfohlen)')
doc.add_paragraph('• Steuerberater: Wegzugsbesteuerung klären')
doc.add_paragraph('• Kapital liquide machen: ETFs/Aktien ggf. verkaufen')
doc.add_paragraph('• Scouting-Reise buchen (4-6 Wochen)')

# Phase 2
h3 = doc.add_heading('Phase 2: Scouting-Reise (Monat 6-8)', 3)
doc.add_paragraph("• 4-6 Wochen im Zielland (z.B. Neuseeland: Hawke's Bay, Nelson, Waikato)")
doc.add_paragraph('• Versammlungen besuchen - vorab Kontakt herstellen!')
doc.add_paragraph('• Grundstücke besichtigen, Makler treffen')
doc.add_paragraph('• Schulen für Kinder anschauen')
doc.add_paragraph('• Familie 2: Mit potenziellen Arbeitgebern sprechen')
doc.add_paragraph('• Internet-Qualität testen an potenziellen Standorten')
doc.add_paragraph('• Mit Expats vor Ort sprechen: Was sind die echten Probleme?')

# Phase 3
h3 = doc.add_heading('Phase 3: Entscheidung und Visa (Monat 8-12)', 3)
doc.add_paragraph('• Finale Entscheidung für Land und Region treffen')
doc.add_paragraph('• Visa-Anträge einreichen:')
doc.add_paragraph('  - Familie 1: Skilled Migrant Visa (NZ) oder EU-Freizügigkeit (ES)')
doc.add_paragraph('  - Familie 2: Accredited Employer Work Visa oder Investor Visa')
doc.add_paragraph('• Immobilie identifizieren und Kaufverhandlung starten')
doc.add_paragraph('• Kündigungen in Deutschland vorbereiten')

# Phase 4
h3 = doc.add_heading('Phase 4: Umzug (Monat 12-18)', 3)
doc.add_paragraph('• Familie 2 kann ggf. vorausreisen (mehr Flexibilität)')
doc.add_paragraph('• Grundstückskauf abschließen')
doc.add_paragraph('• Familie 1: Umzug zum Schuljahresbeginn planen')
doc.add_paragraph('  - NZ Schuljahr: Februar - Dezember → Umzug Januar ideal')
doc.add_paragraph('  - Spanien Schuljahr: September - Juni → Umzug August ideal')
doc.add_paragraph('• Remote-Work-Routine etablieren')
doc.add_paragraph('• Bei Versammlung vorstellen, aktiv integrieren')

# Phase 5
h3 = doc.add_heading('Phase 5: Aufbau und Etablierung (Monat 18-36)', 3)
doc.add_paragraph('• Hausbau/Renovierung abschließen')
doc.add_paragraph('• Selbstversorgungs-Projekt beginnen: Garten, Tiere, Gewächshaus')
doc.add_paragraph('• Residency/Permanent Residence finalisieren')
doc.add_paragraph('• Langfristige Finanzstruktur: Lokales Konto, Investitionen')
doc.add_paragraph('• Kinder voll integriert in Schule')
doc.add_paragraph('• Aktiv in Versammlung und Gemeinschaft')

doc.add_page_break()

# =============================================================================
# EMPFEHLUNG
# =============================================================================
h1 = doc.add_heading('UNSERE EMPFEHLUNG FÜR EUCH', 1)
for run in h1.runs:
    run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('ERSTE WAHL: NEUSEELAND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Euer Profil ist wie gemacht für Neuseeland:')
doc.add_paragraph('✓ Fließend Englisch = Sofortige Integration ohne Sprachbarriere')
doc.add_paragraph('✓ IT-Skills (Familie 1) = Skilled Migrant Visa realistisch')
doc.add_paragraph('✓ Arbeitskräftemangel = Familie 2 findet leicht Vor-Ort-Jobs')
doc.add_paragraph('✓ Eigenkapital = Budget für teurere NZ-Grundstücke ausreichend')
doc.add_paragraph('✓ Aktive ZJ-Gemeinschaft = Sofortige Anbindung')
doc.add_paragraph('✓ Maximale geopolitische Sicherheit')
doc.add_paragraph('✓ Exzellentes Bildungssystem für die Kinder')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE: SPANIEN (Kanaren → Festland)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn EU-Nähe und -Rechte priorisiert werden:')
doc.add_paragraph('✓ Gleiche Zeitzone = Perfekt für Remote-Work')
doc.add_paragraph('✓ 2-3h Flug nach Deutschland = Einfache Familienbesuche')
doc.add_paragraph('✓ Deutschsprachige ZJ-Versammlungen = Sanfter Übergang')
doc.add_paragraph('✓ EU-Rechte = Krankenversicherung, Rente bleiben einfach')
doc.add_paragraph('⚠ Spanisch lernen nötig (12-18 Monate)')
doc.add_paragraph('⚠ Arbeitsmarkt schwieriger für Familie 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('NÄCHSTER SCHRITT:')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph('4-6 Wochen Scouting-Reise nach Neuseeland (Nordinsel + Südinsel) für Sommer/Herbst 2025 planen. Beide Familien zusammen. Regionen erkunden, Versammlungen besuchen, mit Arbeitgebern sprechen, Grundstücke anschauen. Danach habt ihr eine fundierte Basis für die finale Entscheidung.')

doc.add_paragraph()
doc.add_paragraph('─' * 70)

# Footer
footer = doc.add_paragraph()
footer.add_run('Erstellt: Januar 2025').font.size = Pt(9)
doc.add_paragraph().add_run('Basierend auf: Auswanderungsanalyse 2025 + erweiterte Kriterien').font.size = Pt(9)
doc.add_paragraph().add_run('Erweitert um: Zeugen Jehovas | Fließend Englisch | Eigenkapital Familie 2 | Vor-Ort-Jobs').font.size = Pt(9)

# Save document
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_Final.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_Final.docx")

