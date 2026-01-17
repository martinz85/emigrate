from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

def create_matrix_cell(cell, symbol, points, explanation, is_header=False):
    """Create a matrix cell with symbol, points, and explanation"""
    cell.text = ""
    
    if is_header:
        p = cell.paragraphs[0]
        run = p.add_run(str(symbol))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1F4E79')
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        run.font.bold = True
        run.font.size = Pt(11)
        
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"({points})")
        run2.font.size = Pt(8)
        
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(explanation)
        run3.font.size = Pt(7)
        run3.font.italic = True
        
        if symbol == "++":
            set_cell_shading(cell, 'C6EFCE')
        elif symbol == "o":
            set_cell_shading(cell, 'FFEB9C')
        elif symbol == "--":
            set_cell_shading(cell, 'FFC7CE')
        
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

# Create document
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025 - FINALE VERSION')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT | Zwei-Familien-Projekt | Fließend Englisch | Zeugen Jehovas')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR): ').bold = True
p1.add_run('Öffnet praktisch alle Türen.')

p2 = doc.add_paragraph()
p2.add_run('2. Remote IT-Job (Familie 1): ').bold = True
p2.add_run('Data Architect - gefragter Beruf, öffnet Visa-Optionen.')

p3 = doc.add_paragraph()
p3.add_run('3. FLIESSEND ENGLISCH: ').bold = True
p3.add_run('Öffnet englischsprachige Länder! Keine Sprachbarriere.')

p4 = doc.add_paragraph()
p4.add_run('4. Zwei-Familien-Projekt: ').bold = True
p4.add_run('Geteilte Kosten, soziales Netz von Tag 1.')

p5 = doc.add_paragraph()
p5.add_run('5. Familie 2 - Eigenkapital + Vor-Ort-Jobs: ').bold = True
p5.add_run('Flexibel für lokalen Arbeitsmarkt.')

p6 = doc.add_paragraph()
p6.add_run('6. Zeugen Jehovas: ').bold = True
p6.add_run('Weltweites Netzwerk, sofortige Gemeinschaft.')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING - KORREKT SORTIERT NACH PUNKTEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('GESAMTRANKING - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# Ranking nach Gesamtpunkten sortiert (Corona ist EIN Kriterium mit x1)
ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Für euer Profil', 'Corona-Hinweis'],
    ['1', 'Neuseeland', '47/52', '90%', 'TOP - Englisch + Jobs + Krise + Pflege', '⚠️ Sehr restriktiv'],
    ['2', 'Uruguay', '41/52', '79%', 'Neutral + Krisensicher + Liberal!', '✅ Liberal'],
    ['3', 'Costa Rica', '40/52', '77%', 'Kein Militär + krisensicher', '✅ Moderat'],
    ['4', 'Spanien (Süden)', '39/52', '75%', 'EU-Rechte + Sonne + Pflege', '⚠️ Streng'],
    ['5', 'Australien', '38/52', '73%', 'Englisch + Jobs + Pflege', '⚠️ Extrem restriktiv!'],
    ['6', 'Zypern', '38/52', '73%', 'EU + Englisch + Sonne', '⚠️ Streng (EU-Niveau)'],
    ['7', 'Kanarische Inseln', '38/52', '73%', 'EU + bestes Klima', '⚠️ Streng (wie ES)'],
    ['8', 'Chile (Mitte)', '36/52', '69%', 'Isoliert + selbstversorgend', '⚠️ Streng'],
    ['9', 'Schweden', '26/52', '50%', 'Liberal bei Corona, aber NATO + kalt', '✅ Sehr liberal'],
    ['10', 'Namibia', '27/52', '52%', 'Englisch + günstig', 'o Moderat'],
    ['11', 'Deutschland', '22/52', '42%', 'NATO-Front + restriktiv', '⚠️ Sehr restriktiv'],
    ['12', 'Nicaragua', '20/52', '38%', 'Instabil, Diktatur', '✅ Liberal (Diktatur)'],
    ['13', 'Nigeria', '14/52', '27%', 'Nicht empfohlen', 'o Kaum durchgesetzt'],
]
create_table_original_style(doc, ranking_data)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Hinweis: ').bold = True
p.add_run('Corona-Restriktivität ist EIN Kriterium unter 18 (Gewicht x1). Die Spalte zeigt, wie das Land bei COVID reagiert hat - ein Indikator für zukünftige Krisen.')

doc.add_page_break()

# =============================================================================
# DETAILMATRIX
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX - 18 KRITERIEN')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('Symbole: ').bold = True
legend.add_run('++ = Sehr gut (2 Pkt) | o = Mittel (1 Pkt) | -- = Schlecht (0 Pkt)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
run = abbrev.add_run('UY=Uruguay, NZ=Neuseeland, ES=Spanien, AU=Australien, CY=Zypern, KAN=Kanaren, CR=Costa Rica, CL=Chile, NAM=Namibia, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix mit allen 18 Kriterien inkl. Corona (Gewicht x1)
matrix_rows = [
    ["Kriterium", "UY", "NZ", "ES", "AU", "CY", "KAN", "CR", "CL", "NAM", "DE", "SE", "NI", "NG"],
    
    ["Geopolitik x2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "NATO"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Geteilt"),
     ("++", "2", "Weit"),
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Stabil"),
     ("--", "0", "Front!"),
     ("--", "0", "Front!"),
     ("o", "1", "Instabil"),
     ("--", "0", "Unsicher")],
    
    ["Einwanderung x1.5",
     ("++", "2", "Einfach"),
     ("++", "2", "IT-Liste"),
     ("++", "2", "EU"),
     ("++", "2", "IT-Liste"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("++", "2", "Investor"),
     ("++", "2", "Einfach"),
     ("++", "2", "Einfach"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "OK"),
     ("o", "1", "Schwer")],
    
    ["Englisch x1.5",
     ("o", "1", "Spanisch"),
     ("++", "2", "Mutter!"),
     ("o", "1", "Spanisch"),
     ("++", "2", "Mutter!"),
     ("++", "2", "Weit verb."),
     ("o", "1", "Spanisch"),
     ("o", "1", "OK"),
     ("--", "0", "Spanisch"),
     ("++", "2", "Amtsspr."),
     ("o", "1", "Deutsch"),
     ("o", "1", "Schwed."),
     ("o", "1", "Spanisch"),
     ("o", "1", "Amtsspr.")],
    
    ["Grundstück x1.5",
     ("++", "2", "Günstig"),
     ("o", "1", "Teuer"),
     ("++", "2", "OK"),
     ("o", "1", "Teuer"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "OK"),
     ("++", "2", "Günstig"),
     ("++", "2", "Sehr günstig"),
     ("--", "0", "Teuer"),
     ("--", "0", "OK"),
     ("++", "2", "Billig"),
     ("++", "2", "Risiko")],
    
    ["Zeitzone x1",
     ("++", "2", "-4h"),
     ("--", "0", "+11h"),
     ("++", "2", "Gleich"),
     ("--", "0", "+9h"),
     ("++", "2", "+1h"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h"),
     ("o", "1", "-5h"),
     ("++", "2", "+1h"),
     ("++", "2", "Basis"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h"),
     ("o", "1", "+1h")],
    
    ["Selbstversorg. x1",
     ("++", "2", "Ideal"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("o", "1", "Wasser"),
     ("o", "1", "Wasser"),
     ("o", "1", "Wasser"),
     ("o", "1", "Tropisch"),
     ("o", "1", "OK"),
     ("o", "1", "Trocken"),
     ("o", "1", "Bürokratie"),
     ("--", "0", "Kurz"),
     ("o", "1", "OK"),
     ("--", "0", "Unsicher")],
    
    ["Klima x1",
     ("++", "2", "2400h"),
     ("++", "2", "2200h"),
     ("++", "2", "3000h"),
     ("o", "1", "Extreme"),
     ("++", "2", "3300h!"),
     ("++", "2", "2800h"),
     ("o", "1", "Tropisch"),
     ("o", "1", "1700h"),
     ("++", "2", "3000h"),
     ("o", "1", "1600h"),
     ("--", "0", "Kalt"),
     ("--", "0", "Heiss"),
     ("--", "0", "Heiss")],
    
    ["Rechtssich. x1.5",
     ("++", "2", "Stabil"),
     ("++", "2", "Top"),
     ("++", "2", "EU"),
     ("++", "2", "Stark"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("++", "2", "Stark"),
     ("++", "2", "Stark"),
     ("--", "0", "Schwach"),
     ("--", "0", "Korrupt")],
    
    ["Gesundheit x1",
     ("o", "1", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Gut"),
     ("++", "2", "Top"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Top"),
     ("++", "2", "Gut"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    ["Lebenserwart. x1",
     ("++", "2", "78J"),
     ("++", "2", "83J"),
     ("++", "2", "84J"),
     ("++", "2", "84J"),
     ("++", "2", "81J"),
     ("++", "2", "84J"),
     ("++", "2", "80J"),
     ("++", "2", "80J"),
     ("o", "1", "66J"),
     ("++", "2", "81J"),
     ("++", "2", "83J"),
     ("o", "1", "75J"),
     ("--", "0", "55J")],
    
    ["Westl. Leben x0.5",
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Gut"),
     ("o", "1", "Meist"),
     ("o", "1", "Teils"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt")],
    
    ["ZJ-Gemeinde x2",
     ("o", "1", "13k"),
     ("++", "2", "14k"),
     ("++", "2", "113k"),
     ("++", "2", "68k"),
     ("o", "1", "2k"),
     ("++", "2", "8k+dt"),
     ("++", "2", "28k!"),
     ("++", "2", "79k"),
     ("o", "1", "3k"),
     ("++", "2", "176k"),
     ("o", "1", "Klein"),
     ("o", "1", "Klein"),
     ("++", "2", "407k!")],
    
    ["ZJ Sprache x1",
     ("o", "1", "Spanisch"),
     ("++", "2", "Englisch!"),
     ("++", "2", "Dt. Vers."),
     ("++", "2", "Englisch!"),
     ("++", "2", "Englisch!"),
     ("++", "2", "Dt. Vers."),
     ("o", "1", "Spanisch"),
     ("o", "1", "Spanisch"),
     ("++", "2", "Englisch!"),
     ("++", "2", "Deutsch!"),
     ("o", "1", "Schwed."),
     ("o", "1", "Spanisch"),
     ("++", "2", "Englisch!")],
    
    ["Jobs F2 x2",
     ("--", "0", "Wenig"),
     ("++", "2", "Mangel!"),
     ("o", "1", "Schwer"),
     ("++", "2", "Mangel!"),
     ("o", "1", "Tourism."),
     ("o", "1", "Tourism."),
     ("o", "1", "Begrenzt"),
     ("o", "1", "OK"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("--", "0", "Riskant"),
     ("--", "0", "Gefährl.")],
    
    ["Nähe EU x0.5",
     ("o", "1", "12-14h"),
     ("--", "0", "24h"),
     ("++", "2", "2-3h"),
     ("--", "0", "22h"),
     ("++", "2", "3-4h"),
     ("++", "2", "4h"),
     ("o", "1", "12h"),
     ("o", "1", "14h"),
     ("o", "1", "10h"),
     ("++", "2", "Basis"),
     ("++", "2", "1-2h"),
     ("o", "1", "12h"),
     ("o", "1", "6h")],
    
    ["Krise NATO x2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "NATO"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Nahost"),
     ("o", "1", "NATO"),
     ("++", "2", "Neutral!"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Import"),
     ("--", "0", "Front!"),
     ("--", "0", "Front!"),
     ("o", "1", "RU-nah"),
     ("--", "0", "Unruhen")],
    
    ["Pflege x1",
     ("o", "1", "Gut"),
     ("++", "2", "Gut"),
     ("++", "2", "EU"),
     ("++", "2", "Top"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("--", "0", "Begrenzt"),
     ("++", "2", "Top"),
     ("++", "2", "Top"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    # Corona-Restriktivität (Gewicht x1)
    ["Corona-Restr. x1",
     ("++", "2", "Liberal"),
     ("--", "0", "Extrem"),
     ("o", "1", "Streng"),
     ("--", "0", "Extrem!"),
     ("o", "1", "Streng"),
     ("o", "1", "Streng"),
     ("o", "1", "Moderat"),
     ("o", "1", "Streng"),
     ("o", "1", "Moderat"),
     ("--", "0", "Extrem"),
     ("++", "2", "Liberal!"),
     ("++", "2", "Liberal"),
     ("++", "2", "Liberal")],
]

# Create matrix table
num_cols = len(matrix_rows[0])
num_rows = len(matrix_rows)
table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

for j, header in enumerate(matrix_rows[0]):
    cell = table.rows[0].cells[j]
    cell.text = header
    set_cell_shading(cell, '1F4E79')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, row_data in enumerate(matrix_rows[1:], start=1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        
        if j == 0:
            cell.text = cell_data
            set_cell_shading(cell, 'D9E2F3')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        else:
            symbol, points, explanation = cell_data
            create_matrix_cell(cell, symbol, points, explanation)

doc.add_page_break()

# =============================================================================
# CORONA HINWEIS SEKTION (nicht zentral, aber informativ)
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('HINWEIS: CORONA-RESTRIKTIVITÄT ALS INDIKATOR')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Warum ist dieses Kriterium relevant?').bold = True

doc.add_paragraph()

doc.add_paragraph('Die COVID-19-Pandemie zeigte, wie unterschiedlich Regierungen auf Krisen reagieren. Das Verhalten während Corona ist ein Indikator für zukünftige Krisen.')

doc.add_paragraph()

corona_summary = [
    ['Kategorie', 'Länder', 'Was bedeutet das?'],
    ['EXTREM restriktiv', 'Neuseeland, Australien, Deutschland', 'Grenzen zu, harte Lockdowns, Impfdruck. Bei nächster Krise wahrscheinlich wieder so.'],
    ['STRENG', 'Spanien, Kanaren, Zypern, Chile', 'Ausgangssperren, Kontrollen. EU-typisches Niveau.'],
    ['MODERAT', 'Costa Rica, Namibia', 'Maßnahmen vorhanden, aber weniger hart durchgesetzt.'],
    ['LIBERAL', 'Uruguay, Schweden, Nicaragua', 'Keine harten Lockdowns, Freiwilligkeit, Vertrauen in Bürger.'],
]
create_table_original_style(doc, corona_summary)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Wichtig: ').bold = True
p.add_run('Corona ist EIN Faktor unter 18. Neuseeland bleibt Platz 1, weil es in ALLEN anderen Kriterien hervorragend abschneidet. Aber ihr solltet wissen, dass bei einer neuen Krise dort wahrscheinlich wieder strikte Maßnahmen kommen.')

doc.add_paragraph()

# =============================================================================
# DETAILANALYSEN TOP 5
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILANALYSEN - TOP 5')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# PLATZ 1: NEUSEELAND
h2 = doc.add_paragraph()
run = h2.add_run('🥇 PLATZ 1: NEUSEELAND (90%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

doc.add_paragraph('✅ Englisch als Muttersprache - sofortige Integration')
doc.add_paragraph('✅ IT auf Skilled Occupation List - Visa realistisch')
doc.add_paragraph('✅ Arbeitskräftemangel - Jobs für Familie 2')
doc.add_paragraph('✅ ZJ auf Englisch (~14.000 Verkündiger)')
doc.add_paragraph('✅ Maximal isoliert bei NATO-Krise')
doc.add_paragraph('✅ Exzellente Pflege im Alter')
doc.add_paragraph('✅ 2.200 Sonnenstunden, mildes Klima')
doc.add_paragraph('❌ Entfernung zu Europa (24h Flug)')
doc.add_paragraph('❌ Zeitzone +11h schwierig für Remote-Work DE')
doc.add_paragraph('⚠️ Corona: EXTREM restriktiv (Zero-COVID, Grenzen 2 Jahre zu)')

doc.add_paragraph()

# PLATZ 2: URUGUAY
h2 = doc.add_paragraph()
run = h2.add_run('🥈 PLATZ 2: URUGUAY (79%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

doc.add_paragraph('✅ NEUTRAL - keine Militärbündnisse')
doc.add_paragraph('✅ SELBSTVERSORGEND - Agrar-Exporteur')
doc.add_paragraph('✅ "Schweiz Südamerikas" - stabile Demokratie')
doc.add_paragraph('✅ Perfekte Zeitzone (-4h) für Remote-Work')
doc.add_paragraph('✅ Günstige Grundstücke (10ha: 80-150k USD)')
doc.add_paragraph('✅ Corona: LIBERAL ("Libertad Responsable")')
doc.add_paragraph('❌ Spanisch nötig')
doc.add_paragraph('❌ Jobs für Familie 2 begrenzt')
doc.add_paragraph('❌ ZJ-Gemeinschaft kleiner (13k)')

doc.add_paragraph()

# PLATZ 3: COSTA RICA
h2 = doc.add_paragraph()
run = h2.add_run('🥉 PLATZ 3: COSTA RICA (77%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

doc.add_paragraph('✅ KEIN MILITÄR seit 1948!')
doc.add_paragraph('✅ Kann sich nicht an Kriegen beteiligen')
doc.add_paragraph('✅ Höchster ZJ-Bevölkerungsanteil (0,54%)')
doc.add_paragraph('✅ Corona: Moderate Maßnahmen')
doc.add_paragraph('❌ Tropisches Klima')
doc.add_paragraph('❌ Spanisch nötig')
doc.add_paragraph('❌ Zeitzone -7h grenzwertig für Remote-Work')

doc.add_paragraph()

# PLATZ 4: SPANIEN
h2 = doc.add_paragraph()
run = h2.add_run('PLATZ 4: SPANIEN SÜDEN (75%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

doc.add_paragraph('✅ EU-Rechte - einfache Auswanderung')
doc.add_paragraph('✅ 2-3h Flug nach Deutschland')
doc.add_paragraph('✅ Gleiche Zeitzone - perfekt für Remote-Work')
doc.add_paragraph('✅ ZJ: 113k mit deutschsprachigen Versammlungen')
doc.add_paragraph('✅ 3.000+ Sonnenstunden')
doc.add_paragraph('✅ Große Grundstücke möglich')
doc.add_paragraph('❌ Spanisch lernen nötig (12-18 Monate)')
doc.add_paragraph('❌ NATO-Mitglied (wenn auch am Rand)')
doc.add_paragraph('⚠️ Corona: Strenge Ausgangssperren')

doc.add_paragraph()

# PLATZ 5+6: AUSTRALIEN & ZYPERN
h2 = doc.add_paragraph()
run = h2.add_run('PLATZ 5-6: AUSTRALIEN & ZYPERN (73%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('AUSTRALIEN: ').bold = True
p.add_run('Englisch, Jobs für F2, exzellente Pflege. ABER: AUKUS-Mitglied bei China-Konflikt exponiert. ⚠️ Corona: EXTREM (Melbourne 262 Tage Lockdown - Weltrekord!).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('ZYPERN: ').bold = True
p.add_run('EU + Englisch weit verbreitet + Sonne + Steuervorteil (Non-Dom). Kein NATO-Mitglied! ⚠️ Corona: Streng (EU-Niveau).')

doc.add_paragraph()

# =============================================================================
# FINALE EMPFEHLUNG
# =============================================================================
doc.add_page_break()

h = doc.add_paragraph()
run = h.add_run('FINALE EMPFEHLUNG')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ERSTE WAHL: NEUSEELAND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Mit eurem Profil (Englisch, IT, Kapital) ist Neuseeland ideal:')
doc.add_paragraph('• Sofortige Integration ohne Sprachbarriere')
doc.add_paragraph('• Jobs für beide Familien')
doc.add_paragraph('• ZJ-Gemeinschaft auf Englisch')
doc.add_paragraph('• Maximal sicher bei NATO-Krise')
doc.add_paragraph('• Exzellente Infrastruktur und Pflege')
doc.add_paragraph()
doc.add_paragraph('⚠️ Seid euch bewusst: Bei der nächsten Pandemie/Krise werden sie wahrscheinlich wieder sehr restriktiv reagieren. Wenn euch das stört, überlegt Alternative.')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE A: URUGUAY (Freiheit + Krise)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn Freiheit bei Krisen + maximale Neutralität wichtig sind:')
doc.add_paragraph('• Liberale Corona-Politik bewiesen')
doc.add_paragraph('• Neutral, selbstversorgend')
doc.add_paragraph('• Perfekte Zeitzone')
doc.add_paragraph('• Spanisch lernen nötig')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE B: ZYPERN (EU + Englisch)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn EU-Nähe + Englisch wichtig sind:')
doc.add_paragraph('• Englisch weit verbreitet')
doc.add_paragraph('• EU-Rechte + Steuervorteil')
doc.add_paragraph('• Kein NATO-Mitglied!')
doc.add_paragraph('• 3-4h Flug nach DE')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE C: SPANIEN SÜDEN (EU + Nähe)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn maximale Nähe zu Deutschland + EU wichtig sind:')
doc.add_paragraph('• 2-3h Flug')
doc.add_paragraph('• Gleiche Zeitzone')
doc.add_paragraph('• Große Grundstücke möglich')
doc.add_paragraph('• Spanisch lernen nötig')

doc.add_paragraph()

doc.add_paragraph('─' * 70)

p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025 | FINALE VERSION').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('18 Kriterien inkl.: ZJ | Englisch | Krise | Pflege | Corona-Restriktivität | Zypern').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Familie 1: 40J, 2 Kinder (12+14) | Familie 2: 50J, keine Kinder').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_FINALE_VERSION.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_FINALE_VERSION.docx")

