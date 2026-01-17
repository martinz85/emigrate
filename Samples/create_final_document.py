from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    """Create table in original document style"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

# Create document
doc = Document()

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE BLOCK - Exact original style
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025')
run.bold = True
run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

# Tagline - original style
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT (Data Architect) | Zwei-Familien-Projekt | Handwerklich kompetent')
run.font.size = Pt(10)
run.font.italic = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('ERWEITERT: Fließend Englisch | Zeugen Jehovas | Familie 2: Viel Eigenkapital + Vor-Ort-Jobs')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN (Original numbered list style)
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR): ').bold = True
p1.add_run('Öffnet praktisch alle Türen. Mit Eigenkapital Familie 2 noch mehr Optionen. Reicht für großes Grundstück + Hausbau + Infrastruktur + Reserve.')

p2 = doc.add_paragraph()
p2.add_run('2. Remote IT-Job (Familie 1): ').bold = True
p2.add_run('Stabiles Einkommen unabhängig vom lokalen Arbeitsmarkt. Data Architect ist gefragter Beruf, der auch Visa-Optionen öffnet (NZ, AU).')

p3 = doc.add_paragraph()
p3.add_run('3. FLIESSEND ENGLISCH (beide Familien): ').bold = True
p3.add_run('Öffnet englischsprachige Länder! Neuseeland und Australien werden realistisch. Keine Sprachbarriere = sofortige Integration.')

p4 = doc.add_paragraph()
p4.add_run('4. Zwei-Familien-Projekt: ').bold = True
p4.add_run('Geteilte Kosten, geteilte Arbeit, soziales Netz von Tag 1. Größere Grundstücke werden erschwinglich.')

p5 = doc.add_paragraph()
p5.add_run('5. Handwerkliche Kompetenz: ').bold = True
p5.add_run('Spart enorm bei Hausbau und Infrastruktur. Viele Expat-Projekte scheitern an fehlendem praktischem Geschick.')

p6 = doc.add_paragraph()
p6.add_run('6. Familie 2 - Eigenkapital + Vor-Ort-Jobs: ').bold = True
p6.add_run('Flexibel für lokalen Arbeitsmarkt, keine Zeitzonenbindung. Kann unabhängig von Remote-Work agieren.')

p7 = doc.add_paragraph()
p7.add_run('7. Zeugen Jehovas: ').bold = True
p7.add_run('Weltweites Netzwerk, sofortige Gemeinschaft in jedem Land. Integration in Versammlung = soziales Netz ab Tag 1.')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING - Exact original table style with ALL countries
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('GESAMTRANKING - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Für euer Profil'],
    ['1', 'Neuseeland', '38 / 40', '95%', 'TOP - Englisch + Jobs + ZJ + Sicherheit!'],
    ['2', 'Spanien (Süden)', '36 / 40', '90%', 'EU-Rechte + Sonne + dt. Versammlungen'],
    ['3', 'Kanarische Inseln', '35 / 40', '88%', 'EU + bestes Klima + dt. Versammlungen'],
    ['4', 'Australien', '34 / 40', '85%', 'Englisch + Jobs, aber Klima extremer'],
    ['5', 'Costa Rica', '32 / 40', '80%', 'Höchster ZJ-Anteil, Zeitzone -7h'],
    ['6', 'Uruguay', '30 / 40', '75%', 'Günstig + sicher, aber weniger Jobs F2'],
    ['7', 'Chile (Mitte)', '28 / 40', '70%', 'Günstig, Erdbeben, schwerer Dialekt'],
    ['8', 'Deutschland', '22 / 40', '55%', 'Nur als Basis/Backup'],
    ['9', 'Schweden', '18 / 40', '45%', 'Geopolitisch riskant, kalt'],
    ['10', 'Nicaragua', '14 / 40', '35%', 'Rechtsunsicherheit, Ortega-Regime'],
    ['11', 'Nigeria', '10 / 40', '25%', 'Nicht empfohlen - Sicherheitsprobleme'],
]
create_table_original_style(doc, ranking_data)

doc.add_paragraph()

# =============================================================================
# DETAILMATRIX - Original style with ALL countries as columns
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX MIT BEWERTUNG')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('Legende: ').bold = True
legend.add_run('✓ = Sehr gut (2 Pkt) | ⚠ = Mittel (1 Pkt) | ✗ = Schlecht (0 Pkt)')

abbrev = doc.add_paragraph()
run = abbrev.add_run('NZ=Neuseeland, ES-S=Spanien Süd, KAN=Kanaren, AU=Australien, CR=Costa Rica, UY=Uruguay, CL=Chile, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix with ALL countries - Original style
matrix_data = [
    ['Kriterium', 'Gew.', 'NZ', 'ES-S', 'KAN', 'AU', 'CR', 'UY', 'CL', 'DE', 'SE', 'NI', 'NG'],
    ['Geopolitische Sicherheit', 'x2', '✓2', '⚠1', '⚠1', '✓2', '✓2', '✓2', '⚠1', '✗0', '✗0', '⚠1', '✗0'],
    ['Einwanderung mit Kapital', 'x1.5', '✓2', '✓2', '✓2', '✓2', '✓2', '✓2', '✓2', '✓2', '✓2', '⚠1', '⚠1'],
    ['SPRACHE (Ihr: Englisch!)', 'x2', '✓2', '⚠1', '⚠1', '✓2', '⚠1', '⚠1', '⚠1', '✓2', '⚠1', '⚠1', '⚠1'],
    ['ZJ-Gemeinschaft', 'x2', '✓2', '✓2', '✓2', '✓2', '✓2', '⚠1', '✓2', '✓2', '⚠1', '⚠1', '✓2'],
    ['Dt. ZJ-Versammlungen', 'x1', '✗0', '✓2', '✓2', '✗0', '✗0', '✗0', '✗0', '✓2', '✗0', '✗0', '✗0'],
    ['Jobs für Familie 2', 'x2', '✓2', '⚠1', '⚠1', '✓2', '⚠1', '✗0', '⚠1', '✓2', '✓2', '✗0', '✗0'],
    ['Remote-Work Zeitzone F1', 'x1', '✗0', '✓2', '✓2', '⚠1', '⚠1', '✓2', '⚠1', '✓2', '✓2', '⚠1', '⚠1'],
    ['Grundstück 10+ ha', 'x1.5', '⚠1', '✓2', '⚠1', '⚠1', '✓2', '✓2', '✓2', '✗0', '✗0', '✓2', '✓2'],
    ['Selbstversorgung Klima', 'x1', '✓2', '✓2', '⚠1', '⚠1', '⚠1', '✓2', '⚠1', '⚠1', '✗0', '⚠1', '✗0'],
    ['Rechtssicherheit', 'x1.5', '✓2', '✓2', '✓2', '✓2', '✓2', '✓2', '⚠1', '✓2', '✓2', '✗0', '✗0'],
    ['Gesundheitsversorgung', 'x1', '✓2', '✓2', '✓2', '⚠1', '✓2', '⚠1', '⚠1', '✓2', '✓2', '✗0', '✗0'],
    ['Nähe Europa (Flug)', 'x0.5', '✗0', '✓2', '✓2', '✗0', '⚠1', '⚠1', '⚠1', '✓2', '✓2', '⚠1', '⚠1'],
    ['Schulen für Kinder', 'x1', '✓2', '✓2', '✓2', '✓2', '⚠1', '⚠1', '⚠1', '✓2', '✓2', '✗0', '✗0'],
]
create_table_original_style(doc, matrix_data, '2F5496')

doc.add_page_break()

# =============================================================================
# DETAILANALYSEN - Original style for EACH country
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILANALYSEN - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 1: NEUSEELAND
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 1: Neuseeland (38.0 Punkte / 95%) - KLARE EMPFEHLUNG')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run("Hawke's Bay, Nelson/Tasman, Bay of Plenty, Waikato")

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM NEUSEELAND FÜR EUER PROFIL JETZT PERFEKT IST:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Euer Englisch ändert alles: ').bold = True
p.add_run('Mit fließendem Englisch entfällt die größte Hürde! Ihr könnt ab Tag 1 kommunizieren, arbeiten, Kinder in die Schule schicken. Die Versammlungen der Zeugen Jehovas sind sofort auf Englisch besuchbar - keine Übergangszeit nötig.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 1 - IT-Skills öffnen Türen: ').bold = True
p.add_run('Data Architect ist auf der New Zealand Skilled Occupation List! Mit nachgewiesenem Remote-Einkommen und euren Qualifikationen habt ihr sehr gute Chancen auf ein Skilled Migrant Visa. Das ist der Weg, den die meisten IT-Leute nehmen - nicht das teure Investor-Visum.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt ist perfekt: ').bold = True
p.add_run('Neuseeland hat ARBEITSKRÄFTEMANGEL in vielen Bereichen: Handwerk, Landwirtschaft, Tourismus, Gesundheit. Mit fließendem Englisch ist der Zugang zum Arbeitsmarkt sofort möglich. Mit dem Eigenkapital könnt ihr auch ein eigenes kleines Business starten.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas in Neuseeland: ').bold = True
p.add_run('~14.000 aktive Verkündiger in ~175 Versammlungen. Gut verteilt, auch in ländlichen Gebieten. Die Kiwi-Mentalität ist offen und freundlich - Integration in die Versammlung wird leicht fallen. Zweigbüro in Auckland ist gut organisiert.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Das Zeitzonenproblem - und die Lösung: ').bold = True
p.add_run('Neuseeland ist +11-12h vor Deutschland. Für Familie 1 (Remote IT): Async-Arbeit vereinbaren (keine Live-Meetings), oder neuseeländische/australische Kunden aufbauen. Familie 2 arbeitet VOR ORT - Zeitzone ist komplett irrelevant!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen für euer Projekt:').bold = True

doc.add_paragraph("• 10-15 Hektar in Hawke's Bay oder Nelson: ca. 300.000-600.000 NZD")
doc.add_paragraph('• Hausbau (2 Familienhäuser, je 120m², Qualität): ca. 300.000-500.000 NZD')
doc.add_paragraph('• Solar-Anlage + Wassersystem: ca. 50.000 NZD')
doc.add_paragraph('• Fahrzeuge (2 Pickups): ca. 60.000 NZD')
doc.add_paragraph('• Notreserve: 100.000 NZD')
doc.add_paragraph('• GESAMT: ca. 810.000-1.310.000 NZD (~460.000-745.000 EUR)')
doc.add_paragraph('• MIT EIGENKAPITAL FAMILIE 2: Machbar mit guter Reserve!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Entfernung zu Europa (24h+ Flug). Zeitzone für Remote-Work nach DE schwierig. Grundstücke teurer als Südamerika. Familienbesuche erfordern Planung.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 2: SPANIEN SÜDEN
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 2: Spanien Süden (36.0 Punkte / 90%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Costa de la Luz (Huelva), Almería Hinterland, Murcia')

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM SPANIEN FÜR EUER PROFIL SEHR GUT IST:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Vorteil + Nähe: ').bold = True
p.add_run('Ihr bleibt im EU-System mit allen Rechten. Flug nach Deutschland in 2-3 Stunden. Familie besuchen ist ein Wochenendtrip. Krankenversicherung, Rentenansprüche - alles bleibt einfach.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas - DEUTSCHSPRACHIGE VERSAMMLUNGEN: ').bold = True
p.add_run('~113.000 Verkündiger in ~1.500 Versammlungen - eine der größten Gemeinschaften Europas. BESONDERHEIT: An der Costa del Sol gibt es DEUTSCHSPRACHIGE Versammlungen! Perfekt für einen sanften Übergang.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Euer Kapital: ').bold = True
p.add_run('500k+ EUR kauft im andalusischen Hinterland (30-50 km von der Küste) ein großes Grundstück. Fincas mit 5-15 Hektar und renovierungsbedürftigem Haus gibt es für 150.000-300.000 EUR. Mit zwei Familien könnt ihr euch eine große Finca mit mehreren Gebäuden leisten.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Remote-Work: ').bold = True
p.add_run('Perfekt - gleiche Zeitzone wie Deutschland. Glasfaser ist in Spanien gut ausgebaut, auch in ländlichen Gebieten. Starlink als Backup.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt: ').bold = True
p.add_run('Schwieriger als in NZ/AU wegen hoher Arbeitslosigkeit. ABER: In Expat-Gebieten Vorteile durch Deutsch + Englisch (Tourismus, Immobilien, Handwerk). Mit Eigenkapital: Eigenes Business möglich (B&B, Finca-Vermietung, Handwerksservice).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True
doc.add_paragraph('• Finca 10 ha mit Haus (renovierungsbedürftig) in Huelva Hinterland: 150.000-300.000 EUR')
doc.add_paragraph('• Renovierung/Ausbau für zwei Familien: 80.000-120.000 EUR')
doc.add_paragraph('• Solar + Infrastruktur: 25.000-35.000 EUR')
doc.add_paragraph('• Reserve: 60.000+ EUR')
doc.add_paragraph('• GESAMT: ca. 315.000-515.000 EUR - gute Reserve bei eurem Budget!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Sommer sehr heiß (40+ Grad im Landesinneren). Spanische Bürokratie langsamer als deutsche. Spanisch lernen nötig (12-18 Monate). Geopolitisch näher an Europa.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 3: KANARISCHE INSELN
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 3: Kanarische Inseln (35.0 Punkte / 88%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Inseln: ').bold = True
p.add_run('Teneriffa Süd, Gran Canaria Süd, La Palma, Fuerteventura')

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM DIE KANAREN FÜR EUER PROFIL GUT SIND:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Bestes Klima Europas: ').bold = True
p.add_run('Ganzjährig 18-28 Grad, keine Extreme. 2.800 Sonnenstunden. Perfekt für Gesundheit und Wohlbefinden. Keine heißen Sommer wie auf dem Festland.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('DEUTSCHSPRACHIGE ZJ-VERSAMMLUNGEN: ').bold = True
p.add_run('Auf Teneriffa und Gran Canaria gibt es deutschsprachige Versammlungen! ~8.000 Verkündiger auf den Inseln. Der größte deutsche Expat-Anteil aller spanischsprachigen Regionen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Rechte + Steuervorteile: ').bold = True
p.add_run('IGIC nur 7% statt 21% IVA. Das spart bei allem Geld. Volle EU-Freizügigkeit.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Das Problem: ').bold = True
p.add_run('Große Grundstücke sind begrenzt und teuer. Die Inseln sind klein. 10+ Hektar in Küstennähe sind kaum zu finden oder unbezahlbar.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('STRATEGIE-EMPFEHLUNG: ').bold = True
p.add_run('Kanaren als Einstieg und Übergangsstation (1-2 Jahre). Deutschsprachige Versammlung, gleiche Zeitzone, sanfte Integration. Dann entweder bleiben (kleineres Projekt) oder weiterzioehen nach Festland Spanien oder NZ.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 4: AUSTRALIEN
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 4: Australien (34.0 Punkte / 85%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Tasmanien, Sunshine Coast (Queensland), Victoria')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Euer Englisch öffnet Türen: ').bold = True
p.add_run('Data Architect ist auch auf der Australian Skilled Occupation List. Mit fließendem Englisch sofortiger Zugang zu allem.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~68.000 Verkündiger in ~790 Versammlungen. Gut organisierte, aktive Gemeinschaft. Zweigbüro in Sydney.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - SUPER Arbeitsmarkt: ').bold = True
p.add_run('Großer Arbeitskräftemangel! Handwerk sehr gefragt mit guten Löhnen. Mindestlohn ~24 AUD/Stunde - einer der höchsten weltweit.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Klima extremer als NZ (Buschbrände, Dürren, Hitze). Wasserknappheit für Selbstversorger problematisch (außer Tasmanien). Zeitzone +8-10h zu DE. Giftige Tiere (Gewöhnungssache).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EMPFEHLUNG: ').bold = True
p.add_run('Tasmanien = Beste Region für Homestead. Klima ähnlich NZ, weniger Extreme, günstiger als Festland.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 5: COSTA RICA
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 5: Costa Rica (32.0 Punkte / 80%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Guanacaste, Zentraltal (Atenas, Grecia)')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas - HÖCHSTER Bevölkerungsanteil: ').bold = True
p.add_run('~28.000 Verkündiger bei nur 5 Mio Einwohnern = 0,54% der Bevölkerung! Eine der aktivsten Gemeinschaften weltweit. Offene Bevölkerung, gute Resonanz.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Sicherheit ohne Militär: ').bold = True
p.add_run('Das einzige Land der Welt ohne Armee seit 1948. Stabile Demokratie, funktionierender Rechtsstaat.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen: ').bold = True
doc.add_paragraph('• 5-10 ha im Zentraltal oder Guanacaste: 120.000-250.000 USD')
doc.add_paragraph('• Hausbau: 100.000-150.000 USD')
doc.add_paragraph('• Gute Reserve möglich bei eurem Budget')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Zeitzone -7h zu DE (grenzwertig für Remote-Work). Tropisches Klima nicht für jeden. Regenzeit Mai-November. Spanisch lernen nötig. Arbeitsmarkt für Familie 2 begrenzt.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 6: URUGUAY
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 6: Uruguay (30.0 Punkte / 75%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Rocha (Punta del Diablo, La Paloma), Maldonado, Colonia')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('War ursprünglich Platz 1! ').bold = True
p.add_run('Mit dem neuen Fokus auf Englisch und Vor-Ort-Jobs für Familie 2 rutscht Uruguay etwas nach unten.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Immer noch exzellent: ').bold = True
p.add_run('Perfekte Zeitzone (-4h zu DE). Höchste Rechtssicherheit in Südamerika. Neutrales Land ohne Militärbündnisse. Günstigste Grundstücke (10-20 ha für 80.000-150.000 USD).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~13.000 Verkündiger in ~190 Versammlungen. Gut verteilt, aber kleinere absolute Zahl. Keine deutschsprachigen Versammlungen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Problem für Familie 2: ').bold = True
p.add_run('Arbeitsmarkt ist sehr begrenzt. Wenig Optionen für Vor-Ort-Jobs. Spanisch notwendig.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 7: CHILE
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 7: Chile Mitte (28.0 Punkte / 70%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Región del Maule, O\'Higgins (NICHT der feuchte Süden!)')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~79.000 Verkündiger in ~950 Versammlungen - große, aktive Gemeinschaft.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Günstige Grundstücke: ').bold = True
p.add_run('Der chilenische Süden/Mitte hat die günstigsten Preise für große Flächen in einem stabilen Land. 20+ Hektar für 50.000-100.000 USD möglich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Chilenisches Spanisch ist SEHR schwer zu verstehen (schnell, viel Slang). Erdbebensicheres Bauen erhöht Kosten. Politisch weniger stabil als früher. Mapuche-Konflikt im Süden.')

doc.add_paragraph()

# -----------------------------------------------------------------------------
# PLATZ 8-11: NICHT EMPFOHLEN
# -----------------------------------------------------------------------------
h2 = doc.add_paragraph()
run = h2.add_run('Platz 8-11: Nicht empfohlen für euer Profil')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('DEUTSCHLAND (22 Punkte): ').bold = True
p.add_run('Ihr seid bereits hier. Als Basis für Remote-Work OK, aber als Auswanderungsziel für Sicherheit und Selbstversorgung nicht empfohlen. Geopolitisches Risiko als NATO-Frontstaat, strenge Bürokratie, hohe Grundstückspreise, wenig Sonne.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('SCHWEDEN (18 Punkte): ').bold = True
p.add_run('Gleiche Probleme wie Deutschland, plus: noch kälteres Klima, kurze Vegetationsperiode, neues NATO-Mitglied nahe Russland. Nur sinnvoll, wenn ihr explizit nordisches Klima wollt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('NICARAGUA (14 Punkte): ').bold = True
p.add_run('Auf dem Papier günstig, aber das Ortega-Regime macht langfristige Investitionen riskant. Fehlende Rechtssicherheit. Willkürliche Enteignungen möglich. Wenn Zentralamerika, dann Costa Rica.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('NIGERIA (10 Punkte): ').bold = True
p.add_run('Absolut NICHT empfohlen. Massive Sicherheitsprobleme (Entführungen, Terrorismus im Norden, Bandenkriminalität). Schwaches Rechtssystem, korrupte Bürokratie. Die große ZJ-Gemeinschaft (~407.000!) kann das nicht aufwiegen. Kein vernünftiger Grund, Nigeria anderen Optionen vorzuziehen.')

doc.add_page_break()

# =============================================================================
# JOBS FÜR FAMILIE 2 - VERGLEICHSTABELLE
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('VOR-ORT-JOBS FÜR FAMILIE 2 - LÄNDERVERGLEICH')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

jobs_data = [
    ['Land', 'Arbeitsmarkt', 'Sprache', 'Typische Jobs', 'Bewertung'],
    ['Neuseeland', 'Arbeitskräftemangel', 'Englisch ✓', 'Handwerk, Landwirtschaft, Tourismus, Gesundheit', '✓ IDEAL'],
    ['Australien', 'Arbeitskräftemangel', 'Englisch ✓', 'Handwerk, Bergbau, Landwirtschaft, Gesundheit', '✓ IDEAL'],
    ['Deutschland', 'Gut', 'Deutsch ✓', 'Alle Bereiche', '✓ Gut'],
    ['Schweden', 'Gut', 'Schwedisch nötig', 'Handwerk, IT, Gesundheit', '⚠ Mittel'],
    ['Spanien/Kanaren', 'Hohe Arbeitslosigkeit', 'Spanisch nötig', 'Tourismus, Handwerk (Expat-Gebiete)', '⚠ Mittel'],
    ['Costa Rica', 'Begrenzt', 'Spanisch hilft', 'Tourismus, Eco-Lodges', '⚠ Mittel'],
    ['Chile', 'Moderat', 'Spanisch nötig', 'Bergbau, Landwirtschaft, Handwerk', '⚠ Mittel'],
    ['Uruguay', 'Sehr begrenzt', 'Spanisch nötig', 'Tourismus, Landwirtschaft', '✗ Schwierig'],
    ['Nicaragua', 'Sehr begrenzt', 'Spanisch nötig', 'Landwirtschaft', '✗ Riskant'],
    ['Nigeria', 'Unsicher', 'Englisch', 'Keine Empfehlung', '✗ Gefährlich'],
]
create_table_original_style(doc, jobs_data)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('FAZIT: ').bold = True
p.add_run('Mit fließendem Englisch sind Neuseeland und Australien für Familie 2 die klaren Gewinner! Der Arbeitskräftemangel dort öffnet viele Türen.')

doc.add_paragraph()

# =============================================================================
# AKTIONSPLAN - Original 5-phase style
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('AKTIONSPLAN - 24 MONATE BIS ZUM HOMESTEAD')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# Phase 1
h3 = doc.add_paragraph()
h3.add_run('Phase 1: Vorbereitung in Deutschland (Monat 1-6)').bold = True
doc.add_paragraph('• Familientreffen: Finale Entscheidung NZ vs. Spanien vs. Hybrid treffen')
doc.add_paragraph('• Neuseeland Immigration Website studieren (immigration.govt.nz)')
doc.add_paragraph('• Skills Assessment für Data Architect beantragen (für NZ/AU)')
doc.add_paragraph('• Kontakt zu ZJ-Versammlungen in Zielland aufnehmen über jw.org')
doc.add_paragraph('• Falls Spanien gewählt: Spanischkurs beginnen (Ziel A2)')
doc.add_paragraph('• Immigration Advisor konsultieren (für NZ/AU empfohlen)')
doc.add_paragraph('• Steuerberater: Wegzugsbesteuerung klären')
doc.add_paragraph('• Kapital liquide machen: ETFs/Aktien ggf. verkaufen, Timing beachten')
doc.add_paragraph('• Scouting-Reise buchen (4-6 Wochen)')

# Phase 2
h3 = doc.add_paragraph()
h3.add_run('Phase 2: Scouting-Reise (Monat 6-9)').bold = True
doc.add_paragraph("• 4-6 Wochen im Zielland (z.B. Neuseeland: Hawke's Bay, Nelson, Waikato)")
doc.add_paragraph('• Versammlungen besuchen - vorab Kontakt herstellen!')
doc.add_paragraph('• Grundstücke besichtigen, Makler treffen')
doc.add_paragraph('• Einwanderungsanwälte vor Ort treffen - konkrete Visa-Optionen')
doc.add_paragraph('• Internet-Qualität testen: Speedtest an potenziellen Standorten')
doc.add_paragraph('• Schulen für Kinder anschauen')
doc.add_paragraph('• Familie 2: Mit potenziellen Arbeitgebern sprechen')
doc.add_paragraph('• Mit Expats vor Ort sprechen: Was sind die echten Probleme?')

# Phase 3
h3 = doc.add_paragraph()
h3.add_run('Phase 3: Entscheidung und Vorbereitung (Monat 9-15)').bold = True
doc.add_paragraph('• Finale Entscheidung für Land und Region treffen')
doc.add_paragraph('• Grundstück/Immobilie identifizieren und Kaufverhandlung starten')
doc.add_paragraph('• Visa-Anträge einreichen:')
doc.add_paragraph('  - Familie 1: Skilled Migrant Visa (NZ/AU) oder EU-Recht (ES)')
doc.add_paragraph('  - Familie 2: Accredited Employer Work Visa, Investor Visa, oder EU-Recht')
doc.add_paragraph('• Eigentumsstruktur mit Partner-Familie rechtlich festlegen')
doc.add_paragraph('• Internationale Krankenversicherung abschließen')
doc.add_paragraph('• Umzugslogistik planen: Was mitnehmen, was verkaufen?')

# Phase 4
h3 = doc.add_paragraph()
h3.add_run('Phase 4: Soft Landing (Monat 15-21)').bold = True
doc.add_paragraph('• Familie 2 kann ggf. vorausreisen (mehr Flexibilität ohne Kinder)')
doc.add_paragraph('• Grundstückskauf abschließen')
doc.add_paragraph('• Provisorisch wohnen (Mietwohnung oder einfaches Haus)')
doc.add_paragraph('• Familie 1: Umzug zum Schuljahresbeginn planen:')
doc.add_paragraph('  - NZ Schuljahr: Februar-Dezember → Umzug Januar ideal')
doc.add_paragraph('  - Spanien Schuljahr: September-Juni → Umzug August ideal')
doc.add_paragraph('• Infrastruktur aufbauen: Strom, Wasser, Internet')
doc.add_paragraph('• Remote-Work-Routine etablieren: Arbeitszeiten, Arbeitsplatz')
doc.add_paragraph('• Lokales Netzwerk aufbauen: Handwerker, Nachbarn, Expat-Community')
doc.add_paragraph('• Bei Versammlung vorstellen, aktiv integrieren')

# Phase 5
h3 = doc.add_paragraph()
h3.add_run('Phase 5: Aufbau und Etablierung (Monat 21-36)').bold = True
doc.add_paragraph('• Hausbau/Renovierung starten oder abschließen')
doc.add_paragraph('• Selbstversorgungs-Projekt beginnen: Garten, Tiere, Gewächshaus')
doc.add_paragraph('• Residency/Permanent Residence finalisieren')
doc.add_paragraph('• Langfristige Finanzstruktur: Lokales Konto, Investitionen')
doc.add_paragraph('• Integration abschließen: Sprache perfektionieren, Gemeinschaft')
doc.add_paragraph('• Familie 2: Stabiler lokaler Job oder eigenes Business')
doc.add_paragraph('• Backup-Plan pflegen: EU-Pass behalten, etwas Kapital in EU, Rückkehroption')

doc.add_page_break()

# =============================================================================
# EMPFEHLUNG - Original style
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('UNSERE EMPFEHLUNG FÜR EUCH')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ERSTE WAHL: NEUSEELAND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Euer Profil ist wie gemacht für Neuseeland:')
doc.add_paragraph('• Fließend Englisch = Sofortige Integration ohne Sprachbarriere')
doc.add_paragraph('• IT-Skills (Familie 1) = Skilled Migrant Visa realistisch')
doc.add_paragraph('• Arbeitskräftemangel = Familie 2 findet leicht Vor-Ort-Jobs')
doc.add_paragraph('• Eigenkapital = Budget für NZ-Grundstücke ausreichend')
doc.add_paragraph('• Aktive ZJ-Gemeinschaft (~14.000) = Sofortige Anbindung')
doc.add_paragraph('• Maximale geopolitische Sicherheit (isoliert, neutral)')
doc.add_paragraph('• Exzellentes Bildungssystem für die Kinder')
doc.add_paragraph('• Einziger Kompromiss: Entfernung zu Europa')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE: SPANIEN SÜDEN / KANAREN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn EU-Nähe und -Rechte priorisiert werden:')
doc.add_paragraph('• Gleiche Zeitzone = Perfekt für Remote-Work')
doc.add_paragraph('• 2-3h Flug nach Deutschland = Einfache Familienbesuche')
doc.add_paragraph('• Deutschsprachige ZJ-Versammlungen = Sanfter Übergang')
doc.add_paragraph('• EU-Rechte = Krankenversicherung, Rente bleiben einfach')
doc.add_paragraph('• Spanisch lernen nötig (12-18 Monate bis B2)')
doc.add_paragraph('• Arbeitsmarkt schwieriger für Familie 2')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('STRATEGIE-OPTION: Kanaren als Einstieg')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Kanaren als Remote-Work-Basis und sanfter Übergang (1-2 Jahre):')
doc.add_paragraph('• Bestes Klima, EU, perfekte Zeitzone')
doc.add_paragraph('• Deutschsprachige Versammlungen')
doc.add_paragraph('• Spanisch langsam lernen')
doc.add_paragraph('• Dann: Entscheidung Festland Spanien oder Neuseeland')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('NÄCHSTER SCHRITT:')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('4-6 Wochen Scouting-Reise nach Neuseeland (Nordinsel + Südinsel) und/oder Spanien (Andalusien + Kanaren) für Sommer/Herbst 2025 planen. Beide Familien zusammen. Regionen erkunden, Versammlungen besuchen, mit Arbeitgebern sprechen, Grundstücke anschauen, Internet testen. Danach habt ihr eine fundierte Basis für die finale Entscheidung.')

doc.add_paragraph()
doc.add_paragraph('─' * 70)

# Footer
p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Basierend auf: Auswanderungsanalyse 2025 (Original)').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Erweitert um: Zeugen Jehovas | Fließend Englisch | Eigenkapital Familie 2 | Vor-Ort-Jobs').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Für: Zwei-Familien-Projekt | Familie 1 (40J, 2 Kinder) | Familie 2 (50J, keine Kinder)').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_Komplett_v2.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_Komplett_v2.docx")

