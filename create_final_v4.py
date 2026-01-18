from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

def create_matrix_cell(cell, symbol, points, explanation, is_header=False):
    """Create a matrix cell with symbol, points, and explanation"""
    cell.text = ""
    
    if is_header:
        p = cell.paragraphs[0]
        run = p.add_run(str(symbol))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1F4E79')
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        run.font.bold = True
        run.font.size = Pt(11)
        
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"({points})")
        run2.font.size = Pt(8)
        
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(explanation)
        run3.font.size = Pt(7)
        run3.font.italic = True
        
        if symbol == "++":
            set_cell_shading(cell, 'C6EFCE')
        elif symbol == "o":
            set_cell_shading(cell, 'FFEB9C')
        elif symbol == "--":
            set_cell_shading(cell, 'FFC7CE')
        
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

# Create document
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025')
run.bold = True
run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT (Data Architect) | Zwei-Familien-Projekt | Handwerklich kompetent')
run.font.size = Pt(10)
run.font.italic = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('ERWEITERT: Fließend Englisch | Zeugen Jehovas | Familie 2: Viel Eigenkapital + Vor-Ort-Jobs')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR): ').bold = True
p1.add_run('Öffnet praktisch alle Türen. Reicht für großes Grundstück + Hausbau + Infrastruktur + Reserve.')

p2 = doc.add_paragraph()
p2.add_run('2. Remote IT-Job (Familie 1): ').bold = True
p2.add_run('Stabiles Einkommen unabhängig vom lokalen Arbeitsmarkt. Data Architect ist gefragter Beruf, der auch Visa-Optionen öffnet (NZ, AU).')

p3 = doc.add_paragraph()
p3.add_run('3. FLIESSEND ENGLISCH (beide Familien): ').bold = True
p3.add_run('Öffnet englischsprachige Länder! NZ und AU werden realistisch. Keine Sprachbarriere = sofortige Integration.')

p4 = doc.add_paragraph()
p4.add_run('4. Zwei-Familien-Projekt: ').bold = True
p4.add_run('Geteilte Kosten, geteilte Arbeit, soziales Netz von Tag 1. Größere Grundstücke werden erschwinglich.')

p5 = doc.add_paragraph()
p5.add_run('5. Handwerkliche Kompetenz: ').bold = True
p5.add_run('Spart enorm bei Hausbau und Infrastruktur.')

p6 = doc.add_paragraph()
p6.add_run('6. Familie 2 - Eigenkapital + Vor-Ort-Jobs: ').bold = True
p6.add_run('Flexibel für lokalen Arbeitsmarkt, keine Zeitzonenbindung.')

p7 = doc.add_paragraph()
p7.add_run('7. Zeugen Jehovas: ').bold = True
p7.add_run('Weltweites Netzwerk, sofortige Gemeinschaft in jedem Land.')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('GESAMTRANKING - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Für euer Profil'],
    ['1', 'Neuseeland', '46 / 50', '92%', 'TOP - Englisch + Jobs + Krise + Pflege!'],
    ['2', 'Spanien (Süden)', '38 / 50', '76%', 'EU-Rechte + Sonne + gute Pflege'],
    ['3', 'Australien', '38 / 50', '76%', 'Englisch + Jobs + exzellente Pflege'],
    ['4', 'Kanarische Inseln', '37 / 50', '74%', 'EU + bestes Klima + Pflege'],
    ['5', 'Uruguay', '39 / 50', '78%', 'Krisensicher! Pflege OK mit Kapital'],
    ['6', 'Costa Rica', '39 / 50', '78%', 'Kein Militär + krisensicher!'],
    ['7', 'Chile (Mitte)', '37 / 50', '74%', 'Isoliert + selbstversorgend'],
    ['8', 'Deutschland', '24 / 50', '48%', 'Exzellente Pflege, aber NATO-FRONT!'],
    ['9', 'Schweden', '22 / 50', '44%', 'Beste Pflege, aber NATO-Ostsee!'],
    ['10', 'Namibia', '26 / 50', '52%', 'Englisch + günstig, Pflege begrenzt'],
    ['11', 'Nicaragua', '16 / 50', '32%', 'Instabil, keine Pflege-Infrastruktur'],
    ['12', 'Nigeria', '10 / 50', '20%', 'Nicht empfohlen - keine Pflege'],
]
create_table_original_style(doc, ranking_data)

doc.add_page_break()

# =============================================================================
# DETAILMATRIX
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('Symbole: ').bold = True
legend.add_run('++ = Sehr gut (2 Pkt) | o = Mittel (1 Pkt) | -- = Schlecht (0 Pkt)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
run = abbrev.add_run('Länder: UY=Uruguay, NZ=Neuseeland, ES-S=Spanien Süd, AU=Australien, KAN=Kanaren, CR=Costa Rica, CL=Chile, NAM=Namibia, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix rows - now includes Namibia (NAM) and Life Expectancy
matrix_rows = [
    ["Kriterium (Gewicht)", "UY", "NZ", "ES-S", "AU", "KAN", "CR", "CL", "NAM", "DE", "SE", "NI", "NG"],
    
    ["Geopolitische Sicherheit\nx2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "EU-Rand"),
     ("o", "1", "AUKUS"),
     ("++", "2", "Weit weg"),
     ("++", "2", "Kein Militär"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Stabil"),
     ("--", "0", "NATO-Front"),
     ("--", "0", "NATO-Ostsee"),
     ("o", "1", "Instabil"),
     ("--", "0", "Unsicher")],
    
    ["Einwanderung (500k + IT)\nx1.5",
     ("++", "2", "Sehr einfach"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "Investor OK"),
     ("++", "2", "Einfach"),
     ("++", "2", "Einfach"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "Einfach"),
     ("o", "1", "Kompliziert")],
    
    ["Arbeiten mit Englisch\nx1.5",
     ("o", "1", "Spanisch\nnötig"),
     ("++", "2", "Mutter-\nsprache"),
     ("o", "1", "Spanisch\nbesser"),
     ("++", "2", "Mutter-\nsprache"),
     ("o", "1", "Spanisch\nbesser"),
     ("o", "1", "Viel Englisch"),
     ("--", "0", "Nur Spanisch"),
     ("++", "2", "Amts-\nsprache!"),
     ("++", "2", "Deutsch\nnötig"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Etwas\nEnglisch"),
     ("o", "1", "Amtssprache")],
    
    ["Grundstück 10+ ha\nx1.5",
     ("++", "2", "50-150k USD"),
     ("o", "1", "Teuer 500k+"),
     ("++", "2", "Hinterland\nOK"),
     ("o", "1", "Teuer"),
     ("o", "1", "Sehr\nbegrenzt"),
     ("o", "1", "Gestiegen"),
     ("++", "2", "Sehr günstig"),
     ("++", "2", "Sehr\ngünstig!"),
     ("--", "0", "Sehr teuer"),
     ("--", "0", "Norden OK"),
     ("++", "2", "Billig"),
     ("++", "2", "Risiko")],
    
    ["Remote-Work Zeitzone\nx1",
     ("++", "2", "-4h ideal"),
     ("--", "0", "+11h schwer"),
     ("++", "2", "Gleich"),
     ("--", "0", "+9h schwer"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "-5h OK"),
     ("++", "2", "+1h gut"),
     ("++", "2", "Basis"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "+1h gut")],
    
    ["Selbstversorgung\nx1",
     ("++", "2", "Ideal"),
     ("++", "2", "Nordinsel top"),
     ("++", "2", "Sehr gut"),
     ("o", "1", "Wasser knapp"),
     ("o", "1", "Wasser\nknapp"),
     ("o", "1", "Tropisch"),
     ("o", "1", "Feucht/kühl"),
     ("o", "1", "Trocken"),
     ("o", "1", "Bürokratie"),
     ("--", "0", "Kurze Saison"),
     ("o", "1", "Tropisch"),
     ("--", "0", "Unsicher")],
    
    ["Klima (Sonne, mild)\nx1",
     ("++", "2", "2400h mild"),
     ("++", "2", "2200h mild"),
     ("++", "2", "3000h"),
     ("o", "1", "Extreme"),
     ("++", "2", "2800h\nperfekt"),
     ("o", "1", "Tropisch"),
     ("o", "1", "1700h feucht"),
     ("++", "2", "3000h\ntrocken"),
     ("o", "1", "1600h"),
     ("--", "0", "Kalt dunkel"),
     ("--", "0", "Heiss feucht"),
     ("--", "0", "Heiss feucht")],
    
    ["Rechtssicherheit\nx1.5",
     ("++", "2", "Stabil"),
     ("++", "2", "Exzellent"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "Stark"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("++", "2", "Stark"),
     ("++", "2", "Stark"),
     ("--", "0", "Schwach"),
     ("--", "0", "Korrupt")],
    
    ["Gesundheitsversorgung\nx1",
     ("o", "1", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Sehr gut"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    ["Lebenserwartung\nx1",
     ("++", "2", "78 Jahre"),
     ("++", "2", "83 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "80 Jahre"),
     ("++", "2", "80 Jahre"),
     ("o", "1", "66 Jahre"),
     ("++", "2", "81 Jahre"),
     ("++", "2", "83 Jahre"),
     ("o", "1", "75 Jahre"),
     ("--", "0", "55 Jahre")],
    
    ["Westlicher Lebensstil\nx0.5",
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Gut"),
     ("o", "1", "Meist"),
     ("o", "1", "Teilweise"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt")],
    
    ["ZJ-Gemeinschaft\nx2",
     ("o", "1", "13k klein"),
     ("++", "2", "14k aktiv"),
     ("++", "2", "113k groß"),
     ("++", "2", "68k aktiv"),
     ("++", "2", "8k + dt.\nVersamml."),
     ("++", "2", "28k 0,54%!"),
     ("++", "2", "79k groß"),
     ("o", "1", "3k klein"),
     ("++", "2", "176k"),
     ("o", "1", "Klein"),
     ("o", "1", "Klein"),
     ("++", "2", "407k!")],
    
    ["ZJ in Eurer Sprache (EN/DE)\nx1",
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "Dt. Versam.\nvorhanden"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "Dt. Versam.\nvorhanden"),
     ("o", "1", "Spanisch,\nEN teils"),
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "DEUTSCH!"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!")],
    
    ["Jobs für Familie 2\nx2",
     ("--", "0", "Sehr\nbegrenzt"),
     ("++", "2", "Arbeits-\nkräftemangel"),
     ("o", "1", "Hohe\nArbeitslos."),
     ("++", "2", "Arbeits-\nkräftemangel"),
     ("o", "1", "Tourismus"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Moderat"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("--", "0", "Riskant"),
     ("--", "0", "Gefährlich")],
    
    ["Nähe Europa (Flug)\nx0.5",
     ("o", "1", "12-14h"),
     ("--", "0", "24h"),
     ("++", "2", "2-3h"),
     ("--", "0", "22h"),
     ("++", "2", "4h"),
     ("o", "1", "12h"),
     ("o", "1", "14h"),
     ("o", "1", "10h"),
     ("++", "2", "Basis"),
     ("++", "2", "1-2h"),
     ("o", "1", "12h"),
     ("o", "1", "6h")],
    
    ["KRISE: Russland-NATO\nx2",
     ("++", "2", "Neutral,\nselbstvers."),
     ("++", "2", "Isoliert,\nselbstvers."),
     ("o", "1", "NATO, aber\nRand"),
     ("o", "1", "AUKUS,\naber weit"),
     ("o", "1", "NATO, aber\nsehr weit"),
     ("++", "2", "Neutral,\nkein Militär"),
     ("++", "2", "Isoliert,\nselbstvers."),
     ("o", "1", "Neutral,\naber Import"),
     ("--", "0", "NATO-Front!\nSehr betroffen"),
     ("--", "0", "NATO-Front!\nOstsee"),
     ("o", "1", "Russland-\nfreundlich"),
     ("--", "0", "Import-abh.\nUnruhen")],
    
    ["Pflege im Alter\n(mit Kapital) x1",
     ("o", "1", "Gut für\nLateinam."),
     ("++", "2", "Sehr gut,\nwestl. Std."),
     ("++", "2", "EU-Standard,\nPflegeheime"),
     ("++", "2", "Exzellent,\ntop System"),
     ("++", "2", "EU-Standard,\nExpat-Pflege"),
     ("o", "1", "OK mit\nKapital"),
     ("o", "1", "Moderat,\nPrivatkliniken"),
     ("--", "0", "Begrenzt,\nFamilie nötig"),
     ("++", "2", "Exzellent,\nPflegevers."),
     ("++", "2", "Exzellent,\nbekannt gut"),
     ("--", "0", "Schwach,\nkeine Infra"),
     ("--", "0", "Sehr schwach,\nnur Familie")],
]

# Create matrix table
num_cols = len(matrix_rows[0])
num_rows = len(matrix_rows)
table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

for j, header in enumerate(matrix_rows[0]):
    cell = table.rows[0].cells[j]
    cell.text = header
    set_cell_shading(cell, '1F4E79')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, row_data in enumerate(matrix_rows[1:], start=1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        
        if j == 0:
            cell.text = cell_data
            set_cell_shading(cell, 'D9E2F3')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        else:
            symbol, points, explanation = cell_data
            create_matrix_cell(cell, symbol, points, explanation)

doc.add_page_break()

# =============================================================================
# DETAILANALYSEN - VOLLSTÄNDIG MIT ALLEN KRITERIEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILANALYSEN - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# =============================================================================
# PLATZ 1: NEUSEELAND - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 1: Neuseeland (38.0 Punkte / 95%) - KLARE EMPFEHLUNG')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run("Hawke's Bay, Nelson/Tasman, Bay of Plenty, Waikato")

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM NEUSEELAND FÜR EUER PROFIL JETZT PERFEKT IST:').bold = True

doc.add_paragraph()

# Sprache/Englisch
p = doc.add_paragraph()
p.add_run('Euer Englisch ändert alles: ').bold = True
p.add_run('Mit fließendem Englisch entfällt die größte Hürde! Ihr könnt ab Tag 1 kommunizieren, arbeiten, Kinder in die Schule schicken. Die Versammlungen der Zeugen Jehovas sind sofort auf Englisch besuchbar - keine Übergangszeit nötig.')

doc.add_paragraph()

# Familie 1 IT
p = doc.add_paragraph()
p.add_run('Familie 1 - IT-Skills öffnen Türen: ').bold = True
p.add_run('Data Architect ist auf der New Zealand Skilled Occupation List! Mit nachgewiesenem Remote-Einkommen und euren Qualifikationen habt ihr sehr gute Chancen auf ein Skilled Migrant Visa.')

doc.add_paragraph()

# Familie 2 Jobs
p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt ist perfekt: ').bold = True
p.add_run('Neuseeland hat ARBEITSKRÄFTEMANGEL in vielen Bereichen: Handwerk, Landwirtschaft, Tourismus, Gesundheit. Mit fließendem Englisch ist der Zugang zum Arbeitsmarkt sofort möglich. Mit dem Eigenkapital könnt ihr auch ein eigenes kleines Business starten.')

doc.add_paragraph()

# ZJ
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas in Neuseeland: ').bold = True
p.add_run('~14.000 aktive Verkündiger in ~175 Versammlungen. Gut verteilt, auch in ländlichen Gebieten. Die Kiwi-Mentalität ist offen und freundlich - Integration in die Versammlung wird leicht fallen. Zweigbüro in Auckland ist gut organisiert.')

doc.add_paragraph()

# GEOPOLITIK - NEU
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('MAXIMAL ISOLIERT - Neuseeland liegt am Ende der Welt, weit weg von allen Konfliktzonen. Kein NATO-Mitglied, keine Militärbündnisse die zu Konflikten führen. Stabile Demokratie seit über 150 Jahren, friedliche Gesellschaft.')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('IDEAL - Neuseeland ist maximal weit weg vom Konfliktgebiet. Kein NATO-Mitglied, keine Verpflichtung zur Beteiligung. Das Land ist zu 80%+ selbstversorgend bei Lebensmitteln (Fleisch, Milch, Gemüse). Keine Abhängigkeit von russischem Gas/Öl. Innere Unruhen sehr unwahrscheinlich - stabile, homogene Gesellschaft. Lebensmittelpreise würden weniger steigen als in Europa.')

doc.add_paragraph()

# KLIMA - NEU
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('2.200 Sonnenstunden pro Jahr. Mildes, gemäßigtes Klima ohne Extreme. Nordinsel: Subtropisch im Norden, gemäßigt im Süden. Südinsel: Kühler, alpiner. Keine Hitzewellen wie Australien, keine Kälte wie Nordeuropa. Viel Grün, ausreichend Regen - ideal für Selbstversorgung.')

doc.add_paragraph()

# SELBSTVERSORGUNG - NEU
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('IDEAL - Besonders die Nordinsel. Ganzjährige Anbausaison möglich. Gute Böden, ausreichend Niederschlag. Keine Wasserknappheit wie Australien. Tierhaltung unproblematisch. Das Land ist dünn besiedelt - echte Selbstversorgung ist hier realistisch.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL - NEU
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. Englischsprachig, britische Wurzeln. Supermärkte, Internet, Infrastruktur wie in Europa. Ihr vermisst nichts vom gewohnten Lebensstandard. Gleichzeitig entspannter und naturverbundener als Europa.')

doc.add_paragraph()

# RECHTSSICHERHEIT - NEU
p = doc.add_paragraph()
p.add_run('Rechtssicherheit/Eigentum: ').bold = True
p.add_run('EXZELLENT - Eines der besten Rechtssysteme weltweit. Transparente Bürokratie, wenig Korruption. Eigentumsrechte werden vollständig respektiert. Kaufvertrag = sicher. Keine Gefahr willkürlicher Enteignungen.')

doc.add_paragraph()

# GESUNDHEIT - NEU
p = doc.add_paragraph()
p.add_run('Gesundheitsversorgung: ').bold = True
p.add_run('SEHR GUT - Öffentliches Gesundheitssystem (ACC) für Unfälle kostenlos. Gute Krankenhäuser in allen größeren Städten. Private Zusatzversicherung empfohlen für schnellere Behandlung.')

doc.add_paragraph()

# PFLEGE IM ALTER
p = doc.add_paragraph()
p.add_run('Pflege im Alter: ').bold = True
p.add_run('SEHR GUT - Westlicher Standard. Pflegeheime (Rest Homes, Retirement Villages) in guter Qualität verfügbar. Mit Kapital: Private Pflegeeinrichtungen mit hohem Standard. Kosten: ca. 1.000-2.000 NZD/Woche für Vollzeitpflege. Häusliche Pflege auch möglich. Keine kulturelle Erwartung, dass Kinder pflegen müssen - institutionelle Pflege ist normal und akzeptiert.')

doc.add_paragraph()

# Zeitzone
p = doc.add_paragraph()
p.add_run('Das Zeitzonenproblem - und die Lösung: ').bold = True
p.add_run('Neuseeland ist +11-12h vor Deutschland. Für Familie 1 (Remote IT): Async-Arbeit vereinbaren (keine Live-Meetings), oder neuseeländische/australische Kunden aufbauen. Familie 2 arbeitet VOR ORT - Zeitzone ist komplett irrelevant!')

doc.add_paragraph()

# Zahlen
p = doc.add_paragraph()
p.add_run('Konkrete Zahlen für euer Projekt:').bold = True
doc.add_paragraph("• 10-15 Hektar in Hawke's Bay oder Nelson: ca. 300.000-600.000 NZD")
doc.add_paragraph('• Hausbau (2 Familienhäuser, je 120m², Qualität): ca. 300.000-500.000 NZD')
doc.add_paragraph('• Solar-Anlage + Wassersystem: ca. 50.000 NZD')
doc.add_paragraph('• Fahrzeuge (2 Pickups): ca. 60.000 NZD')
doc.add_paragraph('• Notreserve: 100.000 NZD')
doc.add_paragraph('• GESAMT: ca. 810.000-1.310.000 NZD (~460.000-745.000 EUR)')
doc.add_paragraph('• MIT EIGENKAPITAL FAMILIE 2: Machbar mit guter Reserve!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Entfernung zu Europa (24h+ Flug). Zeitzone für Remote-Work nach DE schwierig. Grundstücke teurer als Südamerika. Familienbesuche erfordern Planung.')

doc.add_paragraph()

# =============================================================================
# PLATZ 2: SPANIEN - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 2: Spanien Süden (36.0 Punkte / 90%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Costa de la Luz (Huelva), Almería Hinterland, Murcia')

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM SPANIEN FÜR EUER PROFIL SEHR GUT IST:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Vorteil + Nähe: ').bold = True
p.add_run('Ihr bleibt im EU-System mit allen Rechten. Flug nach Deutschland in 2-3 Stunden. Familie besuchen ist ein Wochenendtrip. Krankenversicherung, Rentenansprüche - alles bleibt einfach.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~113.000 Verkündiger in ~1.500 Versammlungen - eine der größten Gemeinschaften Europas. BONUS: An der Costa del Sol gibt es auch deutschsprachige Versammlungen. Aber ihr könnt mit eurem Englisch auch spanische Versammlungen besuchen und euch schnell integrieren.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('AM RAND EUROPAS - Spanien liegt weit weg von der NATO-Ostflanke. Bei einem europäischen Konflikt weniger exponiert als Deutschland. Aber: Immer noch EU und NATO-Mitglied.')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('MITTEL - Als NATO-Mitglied theoretisch betroffen, aber geographisch am Rand. Keine direkte Grenze zu Russland. Lebensmittel-Selbstversorgung gut (Obst, Gemüse, Olivenöl). ABER: Abhängig von Energieimporten, Preise würden steigen. Innere Unruhen möglich bei Wirtschaftskrise. Besser als Deutschland, schlechter als Südamerika.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('3.000+ Sonnenstunden im Süden - mehr als fast überall in Europa! Milde Winter (10-15°C), heiße Sommer (bis 40°C im Landesinneren). Perfekt für Solar. Anbausaison fast ganzjährig. Mittelmeer oder Atlantik erreichbar.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('SEHR GUT - Lange Anbausaison, viel Sonne. Oliven, Zitrusfrüchte, Gemüse wachsen hervorragend. ABER: Wasser kann in manchen Regionen knapp sein - eigener Brunnen und Zisterne wichtig. Mit Planung aber machbar.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. EU-Standard in allem. Supermärkte, Internet, Infrastruktur. Deutsche Produkte erhältlich. Große deutsche Expat-Community, besonders an der Küste.')

doc.add_paragraph()

# RECHTSSICHERHEIT
p = doc.add_paragraph()
p.add_run('Rechtssicherheit/Eigentum: ').bold = True
p.add_run('EU-STANDARD - Solides Rechtssystem. Eigentumsrechte geschützt. Notare für Kaufverträge. Bürokratie langsamer als in Deutschland, aber funktioniert.')

doc.add_paragraph()

# GESUNDHEIT
p = doc.add_paragraph()
p.add_run('Gesundheitsversorgung: ').bold = True
p.add_run('SEHR GUT - Öffentliches Gesundheitssystem mit EU-Karte nutzbar. Gute Krankenhäuser auch in ländlichen Gebieten. Private Zusatzversicherung für schnellere Termine möglich.')

doc.add_paragraph()

# PFLEGE IM ALTER
p = doc.add_paragraph()
p.add_run('Pflege im Alter: ').bold = True
p.add_run('EU-STANDARD - Gute Pflegeheime (Residencias) verfügbar. Mit Kapital: Private Seniorenresidenzen mit Pool, Garten, deutschsprachigem Personal in Expat-Gebieten! Kosten: 2.000-4.000 EUR/Monat für Vollzeitpflege. Viele deutsche Rentner nutzen Spanien für den Ruhestand. Häusliche Pflege auch günstig (Pflegekräfte aus Osteuropa/Lateinamerika).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Remote-Work: ').bold = True
p.add_run('PERFEKT - Gleiche Zeitzone wie Deutschland. Glasfaser gut ausgebaut, auch in ländlichen Gebieten. Starlink als Backup.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt: ').bold = True
p.add_run('Schwieriger als in NZ/AU wegen hoher Arbeitslosigkeit in Spanien. ABER: In Expat-Gebieten Vorteile durch Deutsch + Englisch (Tourismus, Immobilien, Handwerk). Mit Eigenkapital: Eigenes Business möglich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True
doc.add_paragraph('• Finca 10 ha mit Haus (renovierungsbedürftig) im Hinterland: 150.000-300.000 EUR')
doc.add_paragraph('• Renovierung/Ausbau für zwei Familien: 80.000-120.000 EUR')
doc.add_paragraph('• Solar + Infrastruktur: 25.000-35.000 EUR')
doc.add_paragraph('• Reserve: 60.000+ EUR')
doc.add_paragraph('• GESAMT: ca. 315.000-515.000 EUR - gute Reserve bei eurem Budget!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Sommer sehr heiß (40+ Grad im Landesinneren). Spanische Bürokratie langsamer als deutsche. Spanisch lernen nötig (12-18 Monate bis B2).')

doc.add_paragraph()

# =============================================================================
# PLATZ 3: KANAREN - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 3: Kanarische Inseln (35.0 Punkte / 88%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Inseln: ').bold = True
p.add_run('Teneriffa Süd, Gran Canaria Süd, La Palma, Fuerteventura')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Bestes Klima Europas: ').bold = True
p.add_run('Ganzjährig 18-28 Grad, keine Extreme. 2.800 Sonnenstunden - perfekt. Kein heißer Sommer wie auf dem Festland, kein kalter Winter. Ideal für Gesundheit und Wohlbefinden.')

doc.add_paragraph()

# ZJ
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~8.000 Verkündiger auf den Inseln. Auf Teneriffa und Gran Canaria gibt es auch deutschsprachige Versammlungen als Option. Mit eurem Englisch könnt ihr aber auch direkt spanische Versammlungen besuchen - der große deutsche Expat-Anteil hilft bei der allgemeinen Integration.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('WEIT WEG - Die Kanaren liegen vor Afrika im Atlantik, weit weg vom europäischen Festland. Bei einem Konflikt in Europa relativ geschützt. EU-Gebiet aber geographisch isoliert.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. EU-Rechte, deutsche Supermärkte (Lidl, Aldi), deutsches Fernsehen. Große deutsche Community - ihr seid nicht allein.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Rechte + Steuervorteile: ').bold = True
p.add_run('IGIC nur 7% statt 21% IVA. Das spart bei allem Geld. Volle EU-Freizügigkeit.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Das Problem - Grundstücke: ').bold = True
p.add_run('Große Grundstücke sind begrenzt und teuer. Die Inseln sind klein. 10+ Hektar in Küstennähe sind kaum zu finden oder unbezahlbar. Alternative: Kleineres Grundstück (2-5 ha) mit intensiver Permakultur.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('STRATEGIE-EMPFEHLUNG: ').bold = True
p.add_run('Kanaren als Einstieg und Übergangsstation (1-2 Jahre). Deutschsprachige Versammlung besuchen, gleiche Zeitzone für Remote-Work, sanfte Integration. Dann entweder bleiben (kleineres Projekt) oder weiterziehen nach Festland Spanien oder Neuseeland.')

doc.add_paragraph()

# =============================================================================
# PLATZ 4: AUSTRALIEN - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 4: Australien (34.0 Punkte / 85%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Tasmanien (beste Wahl!), Sunshine Coast (Queensland), Victoria')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Euer Englisch öffnet Türen: ').bold = True
p.add_run('Data Architect ist auf der Australian Skilled Occupation List. Mit fließendem Englisch sofortiger Zugang zu allem.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~68.000 Verkündiger in ~790 Versammlungen. Gut organisierte, aktive Gemeinschaft. Zweigbüro in Sydney.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('AUKUS-MITGLIED - Australien ist Teil des Militärbündnisses mit USA und UK. Bei einem Konflikt mit China potentiell exponiert. Weniger neutral als Neuseeland. Aber: Weit weg von Europa.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('EXTREME möglich - Buschbrände, Dürren, Hitzewellen auf dem Festland. Tasmanien ist die Ausnahme: Gemäßigtes Klima ähnlich Neuseeland. Für Selbstversorgung ist Tasmanien die beste Wahl.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('PROBLEMATISCH auf dem Festland wegen Wasserknappheit und Dürren. Tasmanien ist anders: Genug Regen, grüne Landschaft, machbar für Homestead. Festland nur mit aufwändiger Wasserwirtschaft.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. Englischsprachig, britische Wurzeln. Hoher Lebensstandard. Alles verfügbar was ihr braucht.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - SUPER Arbeitsmarkt: ').bold = True
p.add_run('Großer Arbeitskräftemangel! Handwerk sehr gefragt mit guten Löhnen. Mindestlohn ~24 AUD/Stunde - einer der höchsten weltweit.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Klima extremer als NZ. Giftige Tiere (Schlangen, Spinnen - Gewöhnungssache). Zeitzone +8-10h zu DE problematisch für Remote-Work.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EMPFEHLUNG: ').bold = True
p.add_run('Wenn Australien, dann TASMANIEN! Beste Region für Homestead - Klima ähnlich NZ, weniger Extreme, günstiger als Festland, grün und wasserreich.')

doc.add_paragraph()

# =============================================================================
# PLATZ 5: COSTA RICA - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 5: Costa Rica (32.0 Punkte / 80%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Guanacaste, Zentraltal (Atenas, Grecia, San Ramón)')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas - HÖCHSTER Bevölkerungsanteil: ').bold = True
p.add_run('~28.000 Verkündiger bei nur 5 Mio Einwohnern = 0,54% der Bevölkerung! Eine der aktivsten Gemeinschaften weltweit. Offene, gastfreundliche Bevölkerung, sehr gute Resonanz.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('KEIN MILITÄR - Das einzige Land der Welt ohne Armee seit 1948! Kann sich nicht an Kriegen beteiligen. Stabile Demokratie seit über 70 Jahren. Friedliche Gesellschaft - "Pura Vida" Mentalität.')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('SEHR GUT - Kein Militär = kann nicht in Konflikt hineingezogen werden! Politisch neutral. Lebensmittel-Selbstversorgung gut (tropische Früchte, Reis, Bohnen). Wenig Abhängigkeit von Europa/Russland. Innere Unruhen unwahrscheinlich - stabile Demokratie. Lebensmittelpreise würden weniger steigen.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('TROPISCH - Ganzjährig warm (20-30°C je nach Höhenlage). Regenzeit Mai-November (nachmittags/abends Regen). Das Zentraltal (Atenas, Grecia) hat das beste Klima - ganzjährig 24 Grad, weder zu heiß noch zu feucht. Küste heißer und feuchter.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('ANDERS als europäische Landwirtschaft. Tropische Früchte, Gemüse ganzjährig. Tierhaltung möglich. Lernkurve nötig für tropische Methoden. Mit eurem handwerklichen Geschick machbar.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('GUT - Costa Rica ist das entwickeltste Land Zentralamerikas. Moderne Supermärkte, gutes Internet, Infrastruktur OK. Viele Expats, Englisch verbreitet in Touristengebieten. Nicht ganz europäischer Standard, aber komfortabel.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True
doc.add_paragraph('• 5-10 ha im Zentraltal oder Guanacaste: 120.000-250.000 USD')
doc.add_paragraph('• Hausbau: 100.000-150.000 USD')
doc.add_paragraph('• Gute Reserve möglich bei eurem Budget')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Zeitzone -7h zu DE (grenzwertig für Remote-Work, 15-23 Uhr Arbeitszeit). Tropisches Klima nicht für jeden. Spanisch lernen nötig. Arbeitsmarkt für Familie 2 begrenzt.')

doc.add_paragraph()

# =============================================================================
# PLATZ 6: URUGUAY - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 6: Uruguay (30.0 Punkte / 75%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Rocha (Punta del Diablo, La Paloma), Maldonado, Colonia')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('WAR URSPRÜNGLICH PLATZ 1! ').bold = True
p.add_run('Mit dem neuen Fokus auf Englisch und Vor-Ort-Jobs für Familie 2 rutscht Uruguay etwas nach unten.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('NEUTRAL - Uruguay ist politisch neutral, keine Militärbündnisse. Klein und irrelevant auf der Weltbühne - niemand hat Interesse es anzugreifen. Stabile Demokratie, die "Schweiz Südamerikas".')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('IDEAL - Politisch neutral, keine Bündnisse! Uruguay ist Agrar-Exporteur - produziert mehr Lebensmittel als es braucht (Rindfleisch, Soja, Reis). Bei globaler Krise würden Lebensmittelpreise hier WENIGER steigen als woanders. Keine Abhängigkeit von russischer Energie. Innere Unruhen sehr unwahrscheinlich - stabile, homogene Gesellschaft.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('2.400 Sonnenstunden, mild. Vier Jahreszeiten wie in Südeuropa. Sommer 25-30°C, Winter mild (8-15°C). Selten Extreme. Atlantikküste mit schönen Stränden.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('IDEAL - Perfektes Klima für Gemüseanbau und Tierhaltung. Genug Regen, keine Wasserknappheit. Günstige große Grundstücke (10-20 ha für 80.000-150.000 USD). Das beste Preis-Leistungs-Verhältnis aller Länder!')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. Europäische Wurzeln (viele italienische und deutsche Einwanderer). Säkulare Gesellschaft, liberal. Supermärkte, Internet, Infrastruktur gut. Montevideo ist sehr modern.')

doc.add_paragraph()

# RECHTSSICHERHEIT
p = doc.add_paragraph()
p.add_run('Rechtssicherheit/Eigentum: ').bold = True
p.add_run('BESTE IN SÜDAMERIKA - Stabil, transparent, funktioniert. Eigentumsrechte werden respektiert. Keine Gefahr willkürlicher Enteignungen. Korruptionsindex sehr gut für die Region.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~13.000 Verkündiger in ~190 Versammlungen. Kleiner als Spanien oder Costa Rica, aber gut organisiert und aktiv. Keine deutschsprachigen Versammlungen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Remote-Work: ').bold = True
p.add_run('PERFEKTE Zeitzone - nur -4h zu Deutschland. Glasfaser in den meisten Küstenorten, Starlink verfügbar.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Problem für Familie 2: ').bold = True
p.add_run('Arbeitsmarkt ist SEHR BEGRENZT. Wenig Optionen für Vor-Ort-Jobs. Spanisch notwendig.')

doc.add_paragraph()

# PFLEGE IM ALTER
p = doc.add_paragraph()
p.add_run('Pflege im Alter: ').bold = True
p.add_run('GUT FÜR LATEINAMERIKA - Mit Kapital gute Optionen. Private Pflegeheime in Montevideo und Punta del Este. Häusliche Pflege günstig (Pflegekräfte vor Ort). Kosten: 1.500-3.000 USD/Monat. Viele Argentinier und Brasilianer nutzen Uruguay für den Ruhestand. Nicht europäischer Standard, aber mit Kapital komfortabel.')

doc.add_paragraph()

# =============================================================================
# PLATZ 7: CHILE - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 7: Chile Mitte (28.0 Punkte / 70%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run("Región del Maule, O'Higgins (NICHT der feuchte Süden!)")

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~79.000 Verkündiger in ~950 Versammlungen - große, aktive Gemeinschaft.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('ISOLIERT - Chile liegt am Ende der Welt, geschützt durch die Anden. Keine Militärbündnisse mit Großmächten. Stabil, aber politisch zuletzt etwas unruhiger.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('VARIERT stark. Der Süden (Valdivia, Los Lagos) ist NICHT sonnig: 1.600-1.800 Sonnenstunden, viel Regen, kühl wie Irland. Die MITTE (Maule, O\'Higgins) ist besser: Mediterranes Klima, mehr Sonne. Wenn ihr Sonne wollt, nicht in den Süden!')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('Im Süden: Feucht und kühl, aber fruchtbar. In der Mitte: Mediterranes Klima, Weinbau, Obstanbau. Günstige Grundstücke: 20+ ha für 50.000-100.000 USD möglich.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('MEIST westlich. Das modernste Land Südamerikas. Gute Infrastruktur, Internet, Supermärkte. In Santiago sehr europäisch. Auf dem Land einfacher.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Chilenisches Spanisch ist SEHR schwer zu verstehen (schnell, viel Slang). Erdbebensicheres Bauen erhöht Kosten. Politisch weniger stabil als früher. Mapuche-Konflikt im Süden. Arbeitsmarkt für Familie 2 moderat.')

doc.add_paragraph()

# =============================================================================
# PLATZ 8: NAMIBIA
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 8: Namibia (26.0 Punkte / 59%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Windhoek Umgebung, Swakopmund, Omaruru')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Englisch als Amtssprache: ').bold = True
p.add_run('Namibia ist eines der wenigen afrikanischen Länder mit Englisch als Amtssprache! Außerdem gibt es eine deutschsprachige Minderheit (~20.000) mit deutschen Schulen, Zeitungen und Kulturvereinen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~3.100 Verkündiger in ~54 Versammlungen. Kleine, aber aktive Gemeinschaft. Versammlungen auf Englisch.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('STABIL - Namibia ist seit der Unabhängigkeit 1990 eine friedliche Demokratie. Keine Konflikte, keine Militärbündnisse. Dünn besiedelt, ruhig.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('3.000+ Sonnenstunden! Wüsten- und Halbwüstenklima. Sehr trocken, wenig Regen. Küste (Swakopmund) kühler und angenehmer. Hochland (Windhoek) mild.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('SCHWIERIG - Wasserknappheit ist das Hauptproblem. Landwirtschaft nur mit Bewässerung möglich. Viehzucht (Rinder, Ziegen) ist verbreitet. Für echte Selbstversorgung herausfordernd.')

doc.add_paragraph()

# LEBENSERWARTUNG
p = doc.add_paragraph()
p.add_run('Lebenserwartung: ').bold = True
p.add_run('NUR 66 JAHRE - Deutlich niedriger als Europa oder NZ/AU. Ursachen: HIV/AIDS-Prävalenz, begrenzte Gesundheitsversorgung außerhalb der Städte. Dies ist ein wichtiger Nachteil!')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('TEILWEISE - In Windhoek und Swakopmund recht westlich (deutsche Bäckereien, Supermärkte). Auf dem Land deutlich einfacher. Infrastruktur begrenzt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Grundstücke: ').bold = True
p.add_run('SEHR GÜNSTIG - Große Farmen (1.000+ ha!) für 100.000-300.000 EUR möglich. Namibia hat eine der niedrigsten Bevölkerungsdichten der Welt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Niedrige Lebenserwartung (66 Jahre). Wasserknappheit. Begrenzte Gesundheitsversorgung. Kleine ZJ-Gemeinschaft. Infrastruktur auf dem Land eingeschränkt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('FAZIT: ').bold = True
p.add_run('Interessant wegen Englisch + Deutsch + günstigen Preisen. ABER: Die niedrige Lebenserwartung und Wasserprobleme sind ernste Nachteile. Nur empfohlen wenn ihr Abenteuer sucht und mit einfacheren Bedingungen umgehen könnt.')

doc.add_paragraph()

# =============================================================================
# PLATZ 9-12: NICHT EMPFOHLEN
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 9-12: Nicht empfohlen für euer Profil')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# DEUTSCHLAND
p = doc.add_paragraph()
p.add_run('DEUTSCHLAND (22 Punkte / 46%): ').bold = True
p.add_run('Ihr seid bereits hier. Als Basis für Remote-Work OK, aber als Auswanderungsziel für Sicherheit NICHT empfohlen.')

doc.add_paragraph('• Geopolitik: NATO-FRONTSTAAT - Bei einem Konflikt mit Russland direkt exponiert.')
doc.add_paragraph('• BEI RUSSLAND-NATO-KONFLIKT: SEHR SCHLECHT!')
doc.add_paragraph('  - US-Militärbasen = primäres Ziel')
doc.add_paragraph('  - Ramstein, Spangdahlem, Grafenwöhr = strategische Ziele')
doc.add_paragraph('  - Energie: 100% abhängig von Importen (Gas, Öl)')
doc.add_paragraph('  - Lebensmittel: Stark abhängig von Importen, Preise würden explodieren')
doc.add_paragraph('  - Innere Unruhen: Wahrscheinlich bei Versorgungsengpässen')
doc.add_paragraph('  - Selbstversorgung: Bürokratie verhindert schnelle Anpassung')
doc.add_paragraph('• Klima: 1.600 Sonnenstunden, wenig Sonne, lange Winter.')
doc.add_paragraph('• Positiv: ZJ gut organisiert (176k), Jobs vorhanden, Rechtssicherheit.')

doc.add_paragraph()

# SCHWEDEN
p = doc.add_paragraph()
p.add_run('SCHWEDEN (18 Punkte / 45%): ').bold = True
p.add_run('Gleiche Probleme wie Deutschland, plus mehr.')

doc.add_paragraph('• Geopolitik: NEUES NATO-MITGLIED nahe Russland, Ostsee-Konfliktzone.')
doc.add_paragraph('• Klima: Kalt und dunkel - lange Winter, kurze Vegetationsperiode.')
doc.add_paragraph('• Selbstversorgung: Nur 3-4 Monate Anbausaison, schwierig.')
doc.add_paragraph('• Nur sinnvoll wenn ihr explizit nordisches Klima wollt.')

doc.add_paragraph()

# NICARAGUA
p = doc.add_paragraph()
p.add_run('NICARAGUA (14 Punkte / 35%): ').bold = True
p.add_run('Auf dem Papier günstig, aber zu riskant.')

doc.add_paragraph('• Geopolitik: ORTEGA-REGIME - Autoritäre Regierung, Verbindungen zu Russland.')
doc.add_paragraph('• Rechtssicherheit: SCHWACH - Willkürliche Enteignungen möglich, keine Rechtssicherheit.')
doc.add_paragraph('• ZJ-Gemeinschaft: Klein, unter Druck durch Regime.')
doc.add_paragraph('• EMPFEHLUNG: Wenn Zentralamerika, dann COSTA RICA statt Nicaragua.')

doc.add_paragraph()

# NIGERIA
p = doc.add_paragraph()
p.add_run('NIGERIA (10 Punkte / 25%): ').bold = True
p.add_run('Absolut NICHT empfohlen!')

doc.add_paragraph('• Sicherheit: MASSIVE PROBLEME - Entführungen, Terrorismus (Boko Haram im Norden), Bandenkriminalität.')
doc.add_paragraph('• Rechtssicherheit: KORRUPT - Schwaches Rechtssystem, Bestechung nötig.')
doc.add_paragraph('• ZJ-Gemeinschaft: Sehr groß (~407.000!) aber das wiegt Risiken NICHT auf.')
doc.add_paragraph('• Westlicher Lebensstil: BEGRENZT - Infrastruktur schlecht, Stromausfälle häufig.')
doc.add_paragraph('• FAZIT: Kein vernünftiger Grund, Nigeria anderen Optionen vorzuziehen.')

doc.add_page_break()

# =============================================================================
# AKTIONSPLAN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('AKTIONSPLAN - 24 MONATE BIS ZUM HOMESTEAD')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Phase 1: Vorbereitung in Deutschland (Monat 1-6)').bold = True
doc.add_paragraph('• Familientreffen: Finale Entscheidung NZ vs. Spanien vs. Hybrid')
doc.add_paragraph('• immigration.govt.nz studieren, Skills Assessment beantragen')
doc.add_paragraph('• Kontakt zu ZJ-Versammlungen im Zielland (jw.org)')
doc.add_paragraph('• Falls Spanien: Spanischkurs beginnen (Ziel A2)')
doc.add_paragraph('• Steuerberater: Wegzugsbesteuerung klären')
doc.add_paragraph('• Scouting-Reise buchen (4-6 Wochen)')

h3 = doc.add_paragraph()
h3.add_run('Phase 2: Scouting-Reise (Monat 6-9)').bold = True
doc.add_paragraph("• 4-6 Wochen im Zielland erkunden")
doc.add_paragraph('• Versammlungen besuchen (vorab Kontakt herstellen!)')
doc.add_paragraph('• Grundstücke besichtigen, Makler treffen')
doc.add_paragraph('• Einwanderungsanwälte vor Ort treffen')
doc.add_paragraph('• Familie 2: Mit Arbeitgebern sprechen')
doc.add_paragraph('• Internet testen, Schulen anschauen')
doc.add_paragraph('• Mit Expats sprechen: Was sind echte Probleme?')

h3 = doc.add_paragraph()
h3.add_run('Phase 3: Entscheidung und Visa (Monat 9-15)').bold = True
doc.add_paragraph('• Finale Länder-/Regionswahl treffen')
doc.add_paragraph('• Visa-Anträge einreichen')
doc.add_paragraph('• Immobilie identifizieren, Kaufverhandlung starten')
doc.add_paragraph('• Eigentumsstruktur rechtlich festlegen')
doc.add_paragraph('• Internationale Krankenversicherung')

h3 = doc.add_paragraph()
h3.add_run('Phase 4: Soft Landing (Monat 15-21)').bold = True
doc.add_paragraph('• Familie 2 kann ggf. vorausreisen (mehr Flexibilität)')
doc.add_paragraph('• Grundstückskauf abschließen')
doc.add_paragraph('• Familie 1: Umzug zum Schuljahresbeginn:')
doc.add_paragraph('  - NZ Schuljahr: Feb-Dez → Umzug Januar ideal')
doc.add_paragraph('  - Spanien Schuljahr: Sep-Jun → Umzug August ideal')
doc.add_paragraph('• Infrastruktur aufbauen, Remote-Work etablieren')
doc.add_paragraph('• Bei Versammlung vorstellen')

h3 = doc.add_paragraph()
h3.add_run('Phase 5: Aufbau und Etablierung (Monat 21-36)').bold = True
doc.add_paragraph('• Hausbau/Renovierung abschließen')
doc.add_paragraph('• Selbstversorgung starten: Garten, Tiere, Gewächshaus')
doc.add_paragraph('• Residency finalisieren')
doc.add_paragraph('• Kinder voll integriert')
doc.add_paragraph('• Familie 2: Stabiler Job oder eigenes Business')
doc.add_paragraph('• Aktiv in Versammlung und Gemeinschaft')

doc.add_paragraph()

# =============================================================================
# EMPFEHLUNG
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('UNSERE EMPFEHLUNG FÜR EUCH')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ERSTE WAHL: NEUSEELAND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Euer Profil ist wie gemacht für Neuseeland:')
doc.add_paragraph('• Fließend Englisch = Sofortige Integration ohne Sprachbarriere')
doc.add_paragraph('• IT-Skills (F1) = Skilled Migrant Visa realistisch')
doc.add_paragraph('• Arbeitskräftemangel = Familie 2 findet Vor-Ort-Jobs')
doc.add_paragraph('• Aktive ZJ-Gemeinschaft (~14.000)')
doc.add_paragraph('• MAXIMALE geopolitische Sicherheit (isoliert, neutral)')
doc.add_paragraph('• Exzellentes Klima und Selbstversorgungspotential')
doc.add_paragraph('• Voller westlicher Lebensstil')
doc.add_paragraph('• Einziger Kompromiss: Entfernung zu Europa')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE: SPANIEN SÜDEN / KANAREN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn EU-Nähe und -Rechte priorisiert werden:')
doc.add_paragraph('• Gleiche Zeitzone = Perfekt für Remote-Work')
doc.add_paragraph('• 2-3h Flug nach DE = Einfache Familienbesuche')
doc.add_paragraph('• ZJ-Versammlungen in bekannter Sprache (DE oder EN möglich)')
doc.add_paragraph('• Bestes Klima Europas (Kanaren: 2.800h Sonne)')
doc.add_paragraph('• EU-Rechte = Krankenversicherung, Rente')
doc.add_paragraph('• Spanisch lernen nötig (12-18 Monate)')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('NÄCHSTER SCHRITT:')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('4-6 Wochen Scouting-Reise nach Neuseeland und/oder Spanien (Andalusien + Kanaren) für Sommer/Herbst 2025 planen. Beide Familien zusammen. Regionen erkunden, Versammlungen besuchen, mit Arbeitgebern sprechen, Grundstücke anschauen, Internet testen. Danach habt ihr eine fundierte Basis für die finale Entscheidung.')

doc.add_paragraph()
doc.add_paragraph('─' * 70)

p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025 | Basierend auf Auswanderungsanalyse 2025').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Erweitert: ZJ | Englisch | Eigenkapital F2 | Vor-Ort-Jobs').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Familie 1: 40J, 2 Kinder (12+14) | Familie 2: 50J, keine Kinder').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_KOMPLETT_FINAL.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_KOMPLETT_FINAL.docx")

