from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def create_table(doc, data, header=True, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            if i == 0 and header:
                set_cell_shading(cell, '1F4E79')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    
    return table

# Create document
doc = Document()

# Title
title = doc.add_heading('AUSWANDERUNGSANALYSE 2025 – ERWEITERT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Fokus: Christlicher Glaube (Zeugen Jehovas) | Sprachintegration | Zwei-Familien-Projekt')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('─' * 80)

# =============================================================================
# PROFIL DER FAMILIEN
# =============================================================================
add_heading_with_style(doc, 'PROFIL DER FAMILIEN', 1)

add_heading_with_style(doc, 'Familie 1', 2)
doc.add_paragraph('• Alter Eltern: ca. 40 Jahre')
doc.add_paragraph('• Kinder: 2 (12 und 14 Jahre)')
doc.add_paragraph('• Remote IT-Job: Data Architect (stabile Einnahmequelle)')
doc.add_paragraph('• Sprachen: Deutsch + FLIESSEND ENGLISCH')
doc.add_paragraph('• Besondere Bedürfnisse: Schulpflichtige Kinder, Bildungssystem wichtig')

add_heading_with_style(doc, 'Familie 2', 2)
doc.add_paragraph('• Alter: ca. 50 Jahre')
doc.add_paragraph('• Kinder: Keine')
doc.add_paragraph('• VIEL EIGENKAPITAL (zusätzlich zum gemeinsamen 500k Budget)')
doc.add_paragraph('• Sprachen: Deutsch + FLIESSEND ENGLISCH')
doc.add_paragraph('• Arbeit: Sucht VOR-ORT-JOBS (keine IT/Remote)')
doc.add_paragraph('• Besondere Bedürfnisse: Flexibel, keine schulischen Bindungen')

add_heading_with_style(doc, 'Gemeinsame Kriterien', 2)
doc.add_paragraph('• Gesamtkapital: 500.000+ EUR (mit Eigenkapital Familie 2 deutlich mehr)')
doc.add_paragraph('• Glaube: Zeugen Jehovas – Integration in lokale Versammlung essenziell')
doc.add_paragraph('• Zeitrahmen: 2-3 Jahre, ggf. früher')
doc.add_paragraph('• Priorität: Schnelle sprachliche und gesellschaftliche Integration')
doc.add_paragraph('• NEUER VORTEIL: Englisch fließend = englischsprachige Länder möglich!')

doc.add_page_break()

# =============================================================================
# WICHTIGE ÄNDERUNG
# =============================================================================
add_heading_with_style(doc, 'WICHTIGE ÄNDERUNG DURCH NEUE INFORMATIONEN', 1)

p = doc.add_paragraph()
p.add_run('Durch das fließende Englisch beider Familien ').bold = False
p.add_run('ändern sich die Empfehlungen deutlich!').bold = True

doc.add_paragraph('')

change_data = [
    ['Kriterium', 'Vorher', 'Jetzt', 'Auswirkung'],
    ['Englischkenntnisse', 'Nicht berücksichtigt', 'Fließend Englisch', 'NZ, AU, USA werden attraktiv'],
    ['Familie 2 Arbeit', 'Remote IT angenommen', 'Vor-Ort-Jobs gesucht', 'Zeitzone weniger wichtig für F2'],
    ['Kapital Familie 2', '500k geteilt', 'Deutlich mehr verfügbar', 'Teurere Länder möglich'],
    ['Sprachbarriere', 'Spanisch lernen nötig', 'Englisch ODER Spanisch', 'Mehr Optionen'],
]
create_table(doc, change_data)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('ERGEBNIS: ').bold = True
p.add_run('Neuseeland und Australien rücken deutlich nach oben! Gleichzeitig bleiben spanischsprachige Länder attraktiv, da Spanisch für Deutsche relativ leicht zu lernen ist.')

doc.add_page_break()

# =============================================================================
# ZJ WELTWEITE PRÄSENZ
# =============================================================================
add_heading_with_style(doc, 'ZEUGEN JEHOVAS – WELTWEITE PRÄSENZ', 1)

zj_data = [
    ['Land', 'Verkündiger', 'Versammlungen', 'Bevölkerungsanteil', 'Sprache'],
    ['Neuseeland', '~14.000', '~175', '0,28%', 'Englisch'],
    ['Australien', '~68.000', '~790', '0,26%', 'Englisch'],
    ['Spanien', '~113.000', '~1.500', '0,24%', 'Spanisch'],
    ['Costa Rica', '~28.000', '~400', '0,54%', 'Spanisch'],
    ['Chile', '~79.000', '~950', '0,40%', 'Spanisch'],
    ['Uruguay', '~13.000', '~190', '0,38%', 'Spanisch'],
    ['Kanaren (Spanien)', '~8.000*', '~100*', '0,36%*', 'Spanisch'],
    ['Deutschland (Vgl.)', '~176.000', '~2.200', '0,21%', 'Deutsch'],
]
create_table(doc, zj_data)

doc.add_paragraph('')
doc.add_paragraph('* Kanaren sind Teil von Spanien, geschätzte Werte für die Inseln')

doc.add_page_break()

# =============================================================================
# NEUES GESAMTRANKING
# =============================================================================
add_heading_with_style(doc, 'NEUES GESAMTRANKING – MIT ENGLISCHKENNTNISSEN', 1)

ranking_data = [
    ['Rang', 'Land', 'Score', 'ZJ-Gemeinschaft', 'Sprache', 'Familie 1', 'Familie 2', 'Empfehlung'],
    ['1', 'NEUSEELAND', '94%', '★★★ Gut', '★★★ Englisch!', '★★★ Ideal', '★★★ Jobs!', 'TOP-EMPFEHLUNG'],
    ['2', 'Spanien Süd/Kanaren', '91%', '★★★ Sehr gut', '★★★ Spanisch', '★★★ Ideal', '★★☆ Gut', 'Sehr empfohlen'],
    ['3', 'Australien', '88%', '★★★ Gut', '★★★ Englisch!', '★★★ Gut', '★★★ Jobs!', 'Sehr empfohlen'],
    ['4', 'Costa Rica', '84%', '★★★ Sehr aktiv', '★★☆ Spanisch', '★★☆ Gut', '★★☆ Mittel', 'Empfohlen'],
    ['5', 'Chile', '80%', '★★★ Groß', '★★☆ Spanisch', '★★☆ Gut', '★★☆ Mittel', 'Empfohlen'],
    ['6', 'Uruguay', '78%', '★★☆ Mittel', '★★☆ Spanisch', '★★★ Ideal', '★★☆ Mittel', 'Empfohlen'],
]
create_table(doc, ranking_data)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('WICHTIG: ').bold = True
p.add_run('Neuseeland steigt von Platz 4 auf Platz 1! ')
p.add_run('Der Grund: Fließendes Englisch eliminiert die Sprachbarriere, und Neuseeland bietet exzellente Vor-Ort-Jobmöglichkeiten für Familie 2.')

doc.add_page_break()

# =============================================================================
# PLATZ 1: NEUSEELAND
# =============================================================================
add_heading_with_style(doc, 'PLATZ 1: NEUSEELAND (94%) – NEUE TOP-EMPFEHLUNG', 1)

add_heading_with_style(doc, 'Regionen: Hawke\'s Bay, Nelson/Tasman, Bay of Plenty, Waikato', 2)

add_heading_with_style(doc, 'Warum Neuseeland jetzt Platz 1 ist:', 2)

doc.add_paragraph('1. ENGLISCH = SOFORTIGE INTEGRATION')
doc.add_paragraph('   • Keine Sprachbarriere – ihr könnt ab Tag 1 kommunizieren')
doc.add_paragraph('   • Kinder können sofort in reguläre Schulen')
doc.add_paragraph('   • Versammlungen sofort auf Englisch besuchbar')
doc.add_paragraph('')

doc.add_paragraph('2. ZEUGEN JEHOVAS IN NEUSEELAND')
nz_zj = [
    ['Aspekt', 'Details'],
    ['Verkündiger', '~14.000 aktive Verkündiger'],
    ['Versammlungen', '~175 Versammlungen landesweit'],
    ['Bevölkerungsanteil', '0,28% - gut verteilt'],
    ['Religionsfreiheit', 'Vollständig garantiert'],
    ['Zweigbüro', 'Auckland (gut organisiert)'],
    ['Kongresse', 'Regionale Kongresse, sehr aktive Gemeinschaft'],
    ['Besonderheit', 'Freundliche, offene Kiwi-Mentalität'],
]
create_table(doc, nz_zj)

doc.add_paragraph('')
doc.add_paragraph('3. VOR-ORT-JOBS FÜR FAMILIE 2')
doc.add_paragraph('   • Neuseeland hat ARBEITSKRÄFTEMANGEL in vielen Bereichen')
doc.add_paragraph('   • Handwerk, Landwirtschaft, Tourismus, Gesundheit – viele Optionen')
doc.add_paragraph('   • Englisch = direkter Zugang zum Arbeitsmarkt')
doc.add_paragraph('   • Mit Eigenkapital: Eigenes kleines Business möglich')
doc.add_paragraph('')

doc.add_paragraph('4. SKILLED MIGRANT VISA FÜR FAMILIE 1')
doc.add_paragraph('   • Data Architect ist auf der NZ Skilled Occupation List!')
doc.add_paragraph('   • Mit IT-Skills + Kapital = hohe Visa-Chancen')
doc.add_paragraph('')

doc.add_paragraph('5. GEOPOLITISCHE SICHERHEIT')
doc.add_paragraph('   • Maximal isoliert – weit weg von Konfliktzonen')
doc.add_paragraph('   • Stabile Demokratie, niedrige Kriminalität')
doc.add_paragraph('   • Perfekt für Selbstversorgung und Homestead-Projekt')

add_heading_with_style(doc, 'Das Zeitzonenproblem – und die Lösung:', 2)
doc.add_paragraph('PROBLEM: Neuseeland ist +11-12h vor Deutschland')
doc.add_paragraph('')
doc.add_paragraph('LÖSUNGEN:')
doc.add_paragraph('1. Familie 1 (IT): Async-Arbeit vereinbaren (keine Live-Meetings)')
doc.add_paragraph('2. Familie 1 (IT): Neuseeländische/australische Kunden suchen')
doc.add_paragraph('3. Familie 2: Arbeitet VOR ORT – Zeitzone irrelevant!')
doc.add_paragraph('4. Hybrid: Zunächst Remote für DE, dann lokale Kunden aufbauen')

add_heading_with_style(doc, 'Finanzielle Kalkulation Neuseeland', 2)
calc = doc.add_paragraph()
calc.add_run('''
Grundstück 5-15 ha (Hawke's Bay, Waikato):    300.000 - 600.000 NZD
Hausbau/Renovierung (2 Familien):              300.000 - 500.000 NZD
Infrastruktur (Solar, Wasser):                  50.000 NZD
Fahrzeuge:                                      60.000 NZD
Reserve (2 Jahre):                             100.000 NZD
──────────────────────────────────────────────────────────────────
GESAMT:                                        810.000 - 1.310.000 NZD
                                              (~460.000 - 745.000 EUR)

MIT KAPITAL FAMILIE 2: Machbar!
''')

add_heading_with_style(doc, 'Visa-Optionen Neuseeland', 2)
visa_nz = [
    ['Visum', 'Für wen', 'Anforderungen', 'Dauer'],
    ['Skilled Migrant', 'Familie 1', 'IT-Skills auf Liste, Jobangebot hilft', '3-12 Monate'],
    ['Investor 2', 'Familie 2', '3 Mio NZD Investment', '12 Monate'],
    ['Accredited Employer Work Visa', 'Familie 2', 'Jobangebot von akkreditiertem Arbeitgeber', '2-6 Monate'],
    ['Visitor Visa', 'Beide', 'Bis 9 Monate, zum Erkunden', 'Sofort'],
]
create_table(doc, visa_nz)

doc.add_page_break()

# =============================================================================
# PLATZ 2: SPANIEN
# =============================================================================
add_heading_with_style(doc, 'PLATZ 2: SPANIEN SÜDEN / KANAREN (91%)', 1)

add_heading_with_style(doc, 'Regionen: Andalusien (Huelva, Almería), Kanaren (Teneriffa, Gran Canaria)', 2)

doc.add_paragraph('Spanien bleibt eine exzellente Option – besonders wenn:')
doc.add_paragraph('• EU-Rechte wichtig sind (Krankenversicherung, Rente, Freizügigkeit)')
doc.add_paragraph('• Nähe zu Deutschland gewünscht (2-3h Flug)')
doc.add_paragraph('• Gleiche Zeitzone für Remote-Work wichtig')
doc.add_paragraph('')

add_heading_with_style(doc, 'Zeugen Jehovas in Spanien', 2)
spain_zj = [
    ['Aspekt', 'Details'],
    ['Verkündiger', '~113.000 – 8. größte Gemeinschaft in Europa'],
    ['Versammlungen', '~1.500 – dichte Abdeckung auch ländlich'],
    ['Deutschsprachige Versammlungen', 'JA! Kanaren, Costa del Sol, Mallorca'],
    ['Kongresssäle', '3 große (Madrid, Barcelona, Valencia)'],
    ['Religionsfreiheit', 'Vollständig seit 1970'],
]
create_table(doc, spain_zj)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('VORTEIL: ').bold = True
p.add_run('Deutschsprachige Versammlungen als Übergangsbrücke! Ihr könnt zunächst auf Deutsch Zusammenkünfte besuchen und langsam auf Spanisch wechseln.')

add_heading_with_style(doc, 'Vor-Ort-Jobs für Familie 2 in Spanien', 2)
doc.add_paragraph('• Tourismus: Hotels, Restaurants, Reiseleitung (Englisch + Deutsch = Vorteil!)')
doc.add_paragraph('• Handwerk: Gefragt, besonders in Expat-Gebieten')
doc.add_paragraph('• Landwirtschaft: Olivenöl, Wein, Gemüseanbau')
doc.add_paragraph('• Eigenes Business: Mit Eigenkapital möglich (B&B, Finca-Vermietung)')
doc.add_paragraph('')
doc.add_paragraph('ABER: Spanischer Arbeitsmarkt ist schwieriger als NZ/AU – höhere Arbeitslosigkeit')

add_heading_with_style(doc, 'Finanzielle Kalkulation Spanien', 2)
doc.add_paragraph('''
Finca 10-15 ha Andalusien:                     150.000 - 250.000 EUR
Renovierung für 2 Familien:                     80.000 - 120.000 EUR
Infrastruktur (Solar, Brunnen):                 25.000 - 35.000 EUR
Lebenshaltung 2 Jahre Reserve:                  60.000 EUR
──────────────────────────────────────────────────────────────────
GESAMT:                                        315.000 - 465.000 EUR
BUDGET:                                        500.000+ EUR
RESERVE:                                        35.000 - 185.000+ EUR ✓
''')

doc.add_page_break()

# =============================================================================
# PLATZ 3: AUSTRALIEN
# =============================================================================
add_heading_with_style(doc, 'PLATZ 3: AUSTRALIEN (88%)', 1)

add_heading_with_style(doc, 'Regionen: Tasmanien, Queensland (Sunshine Coast), Victoria', 2)

doc.add_paragraph('Australien ist mit fließendem Englisch jetzt sehr attraktiv!')
doc.add_paragraph('')

add_heading_with_style(doc, 'Zeugen Jehovas in Australien', 2)
au_zj = [
    ['Aspekt', 'Details'],
    ['Verkündiger', '~68.000 aktive Verkündiger'],
    ['Versammlungen', '~790 Versammlungen'],
    ['Zweigbüro', 'Australasien (Sydney)'],
    ['Kongresse', 'Große regionale Kongresse'],
    ['Religionsfreiheit', 'Vollständig garantiert'],
    ['Besonderheit', 'Gut organisierte, aktive Gemeinschaft'],
]
create_table(doc, au_zj)

add_heading_with_style(doc, 'Vor-Ort-Jobs für Familie 2', 2)
doc.add_paragraph('• GROSSER ARBEITSKRÄFTEMANGEL in vielen Bereichen')
doc.add_paragraph('• Handwerk (Elektriker, Klempner, Bau): Sehr gefragt, gute Löhne')
doc.add_paragraph('• Landwirtschaft: Viele Möglichkeiten')
doc.add_paragraph('• Gesundheitswesen: Hoher Bedarf')
doc.add_paragraph('• Mindestlohn: ~24 AUD/Stunde – einer der höchsten weltweit!')

add_heading_with_style(doc, 'Nachteile Australien', 2)
doc.add_paragraph('• Klima: Extremer als NZ (Buschbrände, Dürren, Hitze)')
doc.add_paragraph('• Wasserknappheit: Für Selbstversorger problematisch (außer Tasmanien)')
doc.add_paragraph('• Grundstückspreise: Teuer in guten Lagen')
doc.add_paragraph('• Giftige Tiere: Schlangen, Spinnen – Gewöhnungssache')
doc.add_paragraph('')
doc.add_paragraph('EMPFEHLUNG: Tasmanien = Beste Region für Homestead (Klima ähnlich NZ, weniger Extreme)')

doc.add_page_break()

# =============================================================================
# PLATZ 4-6: LATEINAMERIKA
# =============================================================================
add_heading_with_style(doc, 'PLATZ 4-6: LATEINAMERIKA (SPANISCHSPRACHIG)', 1)

add_heading_with_style(doc, 'Costa Rica (84%)', 2)
doc.add_paragraph('+ Höchster ZJ-Bevölkerungsanteil (0,54%)')
doc.add_paragraph('+ Sehr aktive Gemeinschaft')
doc.add_paragraph('+ Kein Militär, stabile Demokratie')
doc.add_paragraph('+ Englisch weit verbreitet in Touristengebieten')
doc.add_paragraph('- Zeitzone: -7h zu Deutschland')
doc.add_paragraph('- Tropisches Klima nicht für jeden')
doc.add_paragraph('- Arbeitsmarkt: Weniger Optionen für Familie 2')
doc.add_paragraph('')

add_heading_with_style(doc, 'Chile (80%)', 2)
doc.add_paragraph('+ Große ZJ-Gemeinschaft (~79.000)')
doc.add_paragraph('+ Günstige Grundstücke')
doc.add_paragraph('+ Stabile Wirtschaft für Südamerika')
doc.add_paragraph('- Erdbebengefahr (höhere Baukosten)')
doc.add_paragraph('- Chilenisches Spanisch schwer zu verstehen')
doc.add_paragraph('- Weniger Englisch als Costa Rica')
doc.add_paragraph('')

add_heading_with_style(doc, 'Uruguay (78%)', 2)
doc.add_paragraph('+ Beste Rechtssicherheit in Südamerika')
doc.add_paragraph('+ Perfekte Zeitzone für Remote-Work (-4h)')
doc.add_paragraph('+ Günstige Grundstücke')
doc.add_paragraph('+ Neutrales Land')
doc.add_paragraph('- Kleinere ZJ-Gemeinschaft (~13.000)')
doc.add_paragraph('- Arbeitsmarkt: Begrenzt für Familie 2')
doc.add_paragraph('- Spanisch notwendig')

doc.add_page_break()

# =============================================================================
# VERGLEICHSTABELLE JOBS
# =============================================================================
add_heading_with_style(doc, 'VOR-ORT-JOBS FÜR FAMILIE 2 – LÄNDERVERGLEICH', 1)

jobs_data = [
    ['Land', 'Arbeitsmarkt', 'Sprachbedarf', 'Typische Jobs', 'Bewertung'],
    ['Neuseeland', 'Arbeitskräftemangel', 'Englisch ✓', 'Handwerk, Landwirtschaft, Tourismus, Gesundheit', '★★★ Ideal'],
    ['Australien', 'Arbeitskräftemangel', 'Englisch ✓', 'Handwerk, Bergbau, Landwirtschaft, Gesundheit', '★★★ Ideal'],
    ['Spanien', 'Hohe Arbeitslosigkeit', 'Spanisch nötig', 'Tourismus, Handwerk in Expat-Gebieten', '★★☆ Mittel'],
    ['Costa Rica', 'Begrenzt', 'Spanisch hilft', 'Tourismus, Eco-Lodges', '★★☆ Mittel'],
    ['Chile', 'Moderat', 'Spanisch nötig', 'Bergbau, Landwirtschaft, Handwerk', '★★☆ Mittel'],
    ['Uruguay', 'Begrenzt', 'Spanisch nötig', 'Tourismus, Landwirtschaft', '★☆☆ Schwierig'],
]
create_table(doc, jobs_data)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('FAZIT JOBS: ').bold = True
p.add_run('Neuseeland und Australien sind für Familie 2 mit Englischkenntnissen und Vor-Ort-Job-Suche die klaren Gewinner!')

doc.add_page_break()

# =============================================================================
# STRATEGIEEMPFEHLUNG
# =============================================================================
add_heading_with_style(doc, 'STRATEGIEEMPFEHLUNG FÜR EURE SITUATION', 1)

add_heading_with_style(doc, 'OPTION A: NEUSEELAND (Empfohlen)', 2)
doc.add_paragraph('')
option_a = doc.add_paragraph()
option_a.add_run('Für beide Familien die beste Gesamtlösung!').bold = True
doc.add_paragraph('')

doc.add_paragraph('Phase 1 (Monat 1-6): Vorbereitung in Deutschland')
doc.add_paragraph('   • Visa-Recherche und Antrag vorbereiten')
doc.add_paragraph('   • Familie 1: Skilled Migrant Visa prüfen')
doc.add_paragraph('   • Familie 2: Accredited Employer Visa oder Investor Visa prüfen')
doc.add_paragraph('   • Kontakt zu ZJ-Versammlungen in NZ aufnehmen')
doc.add_paragraph('')

doc.add_paragraph('Phase 2 (Monat 6-9): Scouting-Reise')
doc.add_paragraph('   • 4-6 Wochen Neuseeland (Nordinsel + Südinsel)')
doc.add_paragraph('   • Regionen erkunden: Hawke\'s Bay, Nelson, Waikato')
doc.add_paragraph('   • Versammlungen besuchen')
doc.add_paragraph('   • Immobilien anschauen')
doc.add_paragraph('   • Arbeitsmarkt für Familie 2 sondieren')
doc.add_paragraph('')

doc.add_paragraph('Phase 3 (Monat 9-18): Visa und Umzug')
doc.add_paragraph('   • Visa-Anträge einreichen')
doc.add_paragraph('   • Immobilie kaufen oder langfristig mieten')
doc.add_paragraph('   • Umzug (Familie 2 kann ggf. vorausgehen)')
doc.add_paragraph('')

doc.add_paragraph('Phase 4 (Jahr 2-3): Integration')
doc.add_paragraph('   • Kinder in Schule integriert')
doc.add_paragraph('   • Familie 2 hat lokalen Job')
doc.add_paragraph('   • Aktiv in Versammlung')
doc.add_paragraph('   • Homestead-Projekt aufbauen')

doc.add_paragraph('')
doc.add_paragraph('─' * 60)
doc.add_paragraph('')

add_heading_with_style(doc, 'OPTION B: SPANIEN (Kanaren → Festland)', 2)
doc.add_paragraph('')
option_b = doc.add_paragraph()
option_b.add_run('Beste Wahl wenn EU-Rechte und Nähe zu Deutschland Priorität haben').bold = True
doc.add_paragraph('')

doc.add_paragraph('Vorteile:')
doc.add_paragraph('   • EU-Rechte bleiben erhalten')
doc.add_paragraph('   • 2-3h Flug nach Deutschland')
doc.add_paragraph('   • Deutschsprachige Versammlungen als Brücke')
doc.add_paragraph('   • Gleiche Zeitzone für Remote-Work')
doc.add_paragraph('')

doc.add_paragraph('Nachteile:')
doc.add_paragraph('   • Spanisch lernen nötig (12-18 Monate)')
doc.add_paragraph('   • Arbeitsmarkt schwieriger für Familie 2')
doc.add_paragraph('   • Geopolitisch näher an Europa')

doc.add_paragraph('')
doc.add_paragraph('─' * 60)
doc.add_paragraph('')

add_heading_with_style(doc, 'OPTION C: SPLIT-STRATEGIE (Kreativ!)', 2)
doc.add_paragraph('')
doc.add_paragraph('Familie 1: Neuseeland (IT-Job remote, Kinder in Schule)')
doc.add_paragraph('Familie 2: Zunächst Spanien (deutschsprachige Versammlung, Übergang)')
doc.add_paragraph('           → Nach 1-2 Jahren: Auch nach NZ wenn es passt')
doc.add_paragraph('')
doc.add_paragraph('ODER:')
doc.add_paragraph('')
doc.add_paragraph('Beide Familien: 1 Jahr Kanaren (Übergang, deutschsprachige Versammlung)')
doc.add_paragraph('                → Dann gemeinsam nach Neuseeland')

doc.add_page_break()

# =============================================================================
# ZEITPLAN
# =============================================================================
add_heading_with_style(doc, 'DETAILLIERTER ZEITPLAN', 1)

add_heading_with_style(doc, 'Szenario: Umzug nach Neuseeland in 24 Monaten', 2)

timeline = [
    ['Phase', 'Zeitraum', 'Familie 1', 'Familie 2', 'Gemeinsam'],
    ['Vorbereitung', 'Jetzt - Monat 6', 'Visa-Recherche, Skills Assessment', 'Job-Recherche NZ, Visa-Optionen', 'Scouting-Reise buchen, ZJ-Kontakt'],
    ['Scouting', 'Monat 6-8', '4-6 Wochen in NZ', 'Mit Arbeitgebern sprechen', 'Versammlungen besuchen, Regionen erkunden'],
    ['Entscheidung', 'Monat 8-12', 'Visa-Antrag einreichen', 'Jobangebot sichern / Visa', 'Immobilie identifizieren'],
    ['Übergang', 'Monat 12-18', 'Kündigung DE, Umzugsplanung', 'Evtl. vorausreisen', 'Immobilie kaufen/mieten'],
    ['Umzug', 'Monat 18-24', 'Umzug zum Schuljahresbeginn (Feb)', 'Integration, Job starten', 'Versammlung wechseln'],
    ['Etablierung', 'Jahr 2-3', 'Remote-Work Routine', 'Lokaler Job läuft', 'Homestead aufbauen'],
]
create_table(doc, timeline)

add_heading_with_style(doc, 'Szenario: "Schnellstart" in 12 Monaten', 2)

doc.add_paragraph('MÖGLICH wenn:')
doc.add_paragraph('• Familie 2 mit Working Holiday Visa oder Visitor Visa vorausreist')
doc.add_paragraph('• Familie 1 Skilled Migrant schnell durchbekommt')
doc.add_paragraph('• Beide flexibel bei Schulterminen sind')
doc.add_paragraph('')
doc.add_paragraph('Neuseeland Schuljahr: Februar - Dezember')
doc.add_paragraph('→ Umzug Januar/Februar = Kinder können sofort starten')

doc.add_page_break()

# =============================================================================
# CHECKLISTE
# =============================================================================
add_heading_with_style(doc, 'CHECKLISTE: NÄCHSTE SCHRITTE', 1)

add_heading_with_style(doc, 'Sofort (Diese Woche)', 2)
doc.add_paragraph('☐ Familientreffen: Prioritäten final abstimmen (NZ vs. Spanien)')
doc.add_paragraph('☐ Neuseeland Immigration Website studieren (immigration.govt.nz)')
doc.add_paragraph('☐ Skills Assessment für Data Architect prüfen')
doc.add_paragraph('☐ Kapital-Situation klären (wie viel hat Familie 2 genau?)')
doc.add_paragraph('')

add_heading_with_style(doc, 'Kurzfristig (Monat 1-3)', 2)
doc.add_paragraph('☐ Kontakt zu ZJ-Versammlung in NZ aufnehmen (jw.org → Meeting finden)')
doc.add_paragraph('☐ Immigration Advisor konsultieren (für NZ empfohlen)')
doc.add_paragraph('☐ Steuerberater: Wegzugsbesteuerung klären')
doc.add_paragraph('☐ Remote-Work-Situation mit Arbeitgeber besprechen')
doc.add_paragraph('☐ Scouting-Reise für Spätsommer 2025 buchen')
doc.add_paragraph('')

add_heading_with_style(doc, 'Scouting-Reise (4-6 Wochen)', 2)
doc.add_paragraph('☐ Nordinsel: Auckland, Hawke\'s Bay, Bay of Plenty')
doc.add_paragraph('☐ Südinsel: Nelson, Marlborough')
doc.add_paragraph('☐ Versammlungen besuchen (vorab Kontakt herstellen!)')
doc.add_paragraph('☐ Immobilienmakler treffen')
doc.add_paragraph('☐ Schulen für Kinder besichtigen')
doc.add_paragraph('☐ Arbeitgeber für Familie 2 kontaktieren')
doc.add_paragraph('☐ Internet-Speed testen an potenziellen Standorten')
doc.add_paragraph('')

add_heading_with_style(doc, 'Entscheidungsphase (Monat 6-12)', 2)
doc.add_paragraph('☐ Finale Länder- und Regionswahl')
doc.add_paragraph('☐ Visa-Anträge einreichen')
doc.add_paragraph('☐ Immobilie identifizieren')
doc.add_paragraph('☐ Kündigungen in Deutschland vorbereiten')
doc.add_paragraph('')

add_heading_with_style(doc, 'Umzug (Monat 12-24)', 2)
doc.add_paragraph('☐ Umzug durchführen')
doc.add_paragraph('☐ Kinder in Schule anmelden')
doc.add_paragraph('☐ Bankkonto eröffnen')
doc.add_paragraph('☐ Gesundheitsversorgung sicherstellen')
doc.add_paragraph('☐ Bei Versammlung vorstellen')
doc.add_paragraph('☐ Remote-Work / lokale Arbeit starten')

doc.add_page_break()

# =============================================================================
# FAZIT
# =============================================================================
add_heading_with_style(doc, 'FINALES FAZIT', 1)

doc.add_paragraph('')
p1 = doc.add_paragraph()
p1.add_run('MIT FLIESSEND ENGLISCH ändert sich alles!').bold = True
doc.add_paragraph('')

doc.add_paragraph('NEUSEELAND ist jetzt die Top-Empfehlung weil:')
doc.add_paragraph('')
doc.add_paragraph('1. ✓ Keine Sprachbarriere – sofortige Integration')
doc.add_paragraph('2. ✓ Exzellenter Arbeitsmarkt für Familie 2')
doc.add_paragraph('3. ✓ IT-Skills öffnen Visa-Türen für Familie 1')
doc.add_paragraph('4. ✓ Aktive ZJ-Gemeinschaft (~14.000 Verkündiger)')
doc.add_paragraph('5. ✓ Maximale geopolitische Sicherheit')
doc.add_paragraph('6. ✓ Perfekt für Selbstversorgung und Homestead')
doc.add_paragraph('7. ✓ Familienfreundlich, exzellentes Bildungssystem')
doc.add_paragraph('')

p2 = doc.add_paragraph()
p2.add_run('SPANIEN bleibt exzellente Alternative wenn:').bold = True
doc.add_paragraph('')
doc.add_paragraph('• EU-Rechte Priorität haben')
doc.add_paragraph('• Nähe zu Deutschland wichtig ist')
doc.add_paragraph('• Deutschsprachige Versammlung als Übergang gewünscht')
doc.add_paragraph('')

doc.add_paragraph('─' * 60)
doc.add_paragraph('')

final = doc.add_paragraph()
final.add_run('EMPFEHLUNG: ').bold = True
final.add_run('Plant eine 4-6 wöchige Scouting-Reise nach Neuseeland (und optional Spanien) für Sommer/Herbst 2025. ')
final.add_run('Besucht Versammlungen, erkundet Regionen, trefft potenzielle Arbeitgeber. ')
final.add_run('Danach habt ihr eine fundierte Basis für die finale Entscheidung!')

doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('─' * 60)
doc.add_paragraph('Erstellt: Januar 2025')
doc.add_paragraph('Basierend auf: Originale Auswanderungsanalyse 2025')
doc.add_paragraph('Erweitert um: Zeugen Jehovas, Sprachintegration, Familienprofile')
doc.add_paragraph('Aktualisiert für: Fließend Englisch, Eigenkapital Familie 2, Vor-Ort-Jobs')

# Save document
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_Erweitert_Komplett.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_Erweitert_Komplett.docx")







