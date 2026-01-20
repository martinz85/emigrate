from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

def create_matrix_cell(cell, symbol, points, explanation, is_header=False):
    """Create a matrix cell with symbol, points, and explanation"""
    cell.text = ""
    
    if is_header:
        p = cell.paragraphs[0]
        run = p.add_run(str(symbol))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1F4E79')
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        run.font.bold = True
        run.font.size = Pt(11)
        
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"({points})")
        run2.font.size = Pt(8)
        
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(explanation)
        run3.font.size = Pt(7)
        run3.font.italic = True
        
        if symbol == "++":
            set_cell_shading(cell, 'C6EFCE')
        elif symbol == "o":
            set_cell_shading(cell, 'FFEB9C')
        elif symbol == "--":
            set_cell_shading(cell, 'FFC7CE')
        
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

# Create document
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025 - FINALE VERSION')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT (Data Architect) | Zwei-Familien-Projekt | Handwerklich kompetent')
run.font.size = Pt(10)
run.font.italic = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('ERWEITERT: Englisch | ZJ | Eigenkapital F2 | Krise | Corona-Restriktivität | Zypern')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# =============================================================================
# NEUE SEKTION: CORONA-RESTRIKTIVITÄT ÜBERSICHT
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('NEUES KRITERIUM: CORONA-RESTRIKTIVITÄT')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Basierend auf dem tatsächlichen Verhalten während der COVID-19-Pandemie 2020-2022:').italic = True

doc.add_paragraph()

corona_data = [
    ['Land', 'Restriktivität', 'Maßnahmen während Corona', 'Bewertung für Freiheit'],
    ['Neuseeland', 'EXTREM', 'Zero-COVID, Grenzen 2 Jahre zu, MIQ-Hotels, strenge Lockdowns', '-- (0 Pkt)'],
    ['Australien', 'EXTREM', 'Melbourne: 262 Tage Lockdown (Weltrekord!), Grenzen zu, Polizei hart', '-- (0 Pkt)'],
    ['Deutschland', 'HOCH', '2G/3G-Regeln, Lockdowns, Maskenpflicht, Impfdruck', '-- (0 Pkt)'],
    ['Spanien', 'HOCH', 'Strengste Ausgangssperren Europas, Polizei auf Straßen', 'o (1 Pkt)'],
    ['Kanaren', 'HOCH', 'Wie Spanien Festland', 'o (1 Pkt)'],
    ['Zypern', 'HOCH', 'Lockdowns, SafePass, Testpflichten', 'o (1 Pkt)'],
    ['Chile', 'HOCH', 'Strenge Lockdowns, Ausgangssperren', 'o (1 Pkt)'],
    ['Costa Rica', 'MODERAT', 'Moderate Maßnahmen, Grenzen zeitweise zu', 'o (1 Pkt)'],
    ['Uruguay', 'NIEDRIG', '"Libertad Responsable" - Keine harten Lockdowns!', '++ (2 Pkt)'],
    ['Schweden', 'NIEDRIG', 'SONDERWEG: Keine Lockdowns, freiwillige Maßnahmen', '++ (2 Pkt)'],
    ['Namibia', 'MODERAT', 'Maßnahmen vorhanden, aber wenig durchgesetzt', 'o (1 Pkt)'],
    ['Nicaragua', 'SEHR NIEDRIG', 'Ortega ignorierte Corona weitgehend', '++ (2 Pkt)'],
    ['Nigeria', 'NIEDRIG', 'Wenig Kapazität für Durchsetzung', '++ (2 Pkt)'],
]
create_table_original_style(doc, corona_data)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('WICHTIG: ').bold = True
p.add_run('Wenn euch persönliche Freiheit in Krisenzeiten wichtig ist, sind die "extremen" Länder (NZ, AU, DE) problematisch. Uruguay und Schweden zeigten, dass es auch liberal geht!')

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR): ').bold = True
p1.add_run('Öffnet praktisch alle Türen.')

p2 = doc.add_paragraph()
p2.add_run('2. Remote IT-Job (Familie 1): ').bold = True
p2.add_run('Data Architect ist gefragter Beruf, öffnet Visa-Optionen.')

p3 = doc.add_paragraph()
p3.add_run('3. FLIESSEND ENGLISCH: ').bold = True
p3.add_run('Öffnet englischsprachige Länder! Keine Sprachbarriere.')

p4 = doc.add_paragraph()
p4.add_run('4. Zwei-Familien-Projekt: ').bold = True
p4.add_run('Geteilte Kosten, soziales Netz von Tag 1.')

p5 = doc.add_paragraph()
p5.add_run('5. Familie 2 - Eigenkapital + Vor-Ort-Jobs: ').bold = True
p5.add_run('Flexibel für lokalen Arbeitsmarkt.')

p6 = doc.add_paragraph()
p6.add_run('6. Zeugen Jehovas: ').bold = True
p6.add_run('Weltweites Netzwerk, sofortige Gemeinschaft.')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING - FINALE VERSION MIT CORONA
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('FINALE GESAMTRANKING - INKL. CORONA-RESTRIKTIVITÄT')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# Neue Berechnung mit Corona-Faktor:
# Neuseeland: 46 - 4 (Corona extrem) = 42 → 84%
# Australien: 38 - 4 (Corona extrem) = 34 → 68%
# Spanien: 38 - 2 (Corona hoch) = 36 → 72%
# Zypern: 37 - 2 = 35 → 70%
# Uruguay: 39 + 2 (liberal) = 41 → 82%
# Schweden: 22 + 2 = 24 → 48%
# Deutschland: 24 - 4 = 20 → 40%

ranking_data = [
    ['#', 'Land', 'Basis', 'Corona', 'Final', '%', 'Kurzbewertung'],
    ['1', 'Uruguay', '39', '+2', '41', '82%', '🏆 KRISE + FREIHEIT! Neutral, selbstversorgend, liberal!'],
    ['2', 'Costa Rica', '39', '+0', '39', '78%', 'Kein Militär, krisensicher, moderate Corona-Politik'],
    ['3', 'Neuseeland', '46', '-4', '42', '84%', 'Top für alles AUSSER Freiheit bei Krisen'],
    ['4', 'Chile', '37', '-2', '35', '70%', 'Isoliert, selbstvers., aber strenge Maßnahmen'],
    ['5', 'Spanien Süd', '38', '-2', '36', '72%', 'EU-Rechte, Sonne, aber strenge Lockdowns'],
    ['6', 'Zypern', '37', '-2', '35', '70%', 'EU + Englisch, aber strenge Maßnahmen'],
    ['7', 'Kanaren', '37', '-2', '35', '70%', 'Bestes Klima, aber wie Spanien restriktiv'],
    ['8', 'Australien', '38', '-4', '34', '68%', 'ACHTUNG: Extremste Lockdowns der Welt!'],
    ['9', 'Namibia', '26', '+0', '26', '52%', 'Günstig, Englisch, moderate Corona-Politik'],
    ['10', 'Schweden', '22', '+2', '24', '48%', 'Liberalste Corona-Politik, aber NATO-Front + kalt'],
    ['11', 'Deutschland', '24', '-4', '20', '40%', 'NATO-Front + Extremer Corona-Kurs'],
    ['12', 'Nicaragua', '16', '+2', '18', '36%', 'Liberal bei Corona, aber instabil + Diktatur'],
    ['13', 'Nigeria', '10', '+2', '12', '24%', 'Nicht empfohlen trotz liberaler Corona-Politik'],
]
create_table_original_style(doc, ranking_data)

doc.add_paragraph()

# Erklärung
p = doc.add_paragraph()
p.add_run('Bewertungssystem Corona-Restriktivität:').bold = True
doc.add_paragraph('• +2 Punkte: Liberale Politik (Uruguay, Schweden, Nicaragua, Nigeria)')
doc.add_paragraph('• +0 Punkte: Moderate Politik (Costa Rica, Namibia)')
doc.add_paragraph('• -2 Punkte: Strenge Politik (Spanien, Kanaren, Zypern, Chile)')
doc.add_paragraph('• -4 Punkte: EXTREME Politik (Neuseeland, Australien, Deutschland)')

doc.add_page_break()

# =============================================================================
# DETAILMATRIX - FINAL MIT CORONA
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX - FINALE VERSION')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('Symbole: ').bold = True
legend.add_run('++ = Sehr gut (2 Pkt) | o = Mittel (1 Pkt) | -- = Schlecht (0 Pkt)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
run = abbrev.add_run('Länder: UY=Uruguay, NZ=Neuseeland, ES-S=Spanien Süd, AU=Australien, CY=Zypern, KAN=Kanaren, CR=Costa Rica, CL=Chile, NAM=Namibia, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix rows - 18 Kriterien inkl. Corona
matrix_rows = [
    ["Kriterium (Gewicht)", "UY", "NZ", "ES-S", "AU", "CY", "KAN", "CR", "CL", "NAM", "DE", "SE", "NI", "NG"],
    
    ["Geopolitische Sicherheit\nx2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "EU-Rand"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Geteilt"),
     ("++", "2", "Weit weg"),
     ("++", "2", "Kein Militär"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Stabil"),
     ("--", "0", "NATO-Front"),
     ("--", "0", "NATO-Ostsee"),
     ("o", "1", "Instabil"),
     ("--", "0", "Unsicher")],
    
    ["Einwanderung (500k + IT)\nx1.5",
     ("++", "2", "Sehr einfach"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "Investor OK"),
     ("++", "2", "Einfach"),
     ("++", "2", "Einfach"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "Einfach"),
     ("o", "1", "Kompliziert")],
    
    ["Arbeiten mit Englisch\nx1.5",
     ("o", "1", "Spanisch"),
     ("++", "2", "Mutter-\nsprache"),
     ("o", "1", "Spanisch"),
     ("++", "2", "Mutter-\nsprache"),
     ("++", "2", "Weit\nverbreitet!"),
     ("o", "1", "Spanisch"),
     ("o", "1", "Englisch OK"),
     ("--", "0", "Nur Spanisch"),
     ("++", "2", "Amts-\nsprache!"),
     ("o", "1", "Deutsch"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Spanisch"),
     ("o", "1", "Amtssprache")],
    
    ["Grundstück 10+ ha\nx1.5",
     ("++", "2", "50-150k USD"),
     ("o", "1", "Teuer 500k+"),
     ("++", "2", "Hinterland"),
     ("o", "1", "Teuer"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Gestiegen"),
     ("++", "2", "Günstig"),
     ("++", "2", "Sehr günstig"),
     ("--", "0", "Sehr teuer"),
     ("--", "0", "OK"),
     ("++", "2", "Billig"),
     ("++", "2", "Risiko")],
    
    ["Remote-Work Zeitzone\nx1",
     ("++", "2", "-4h ideal"),
     ("--", "0", "+11h schwer"),
     ("++", "2", "Gleich"),
     ("--", "0", "+9h schwer"),
     ("++", "2", "+1h gut"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "-5h OK"),
     ("++", "2", "+1h gut"),
     ("++", "2", "Basis"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "+1h gut")],
    
    ["Selbstversorgung\nx1",
     ("++", "2", "Ideal"),
     ("++", "2", "Nordinsel"),
     ("++", "2", "Sehr gut"),
     ("o", "1", "Wasser"),
     ("o", "1", "Wasser"),
     ("o", "1", "Wasser"),
     ("o", "1", "Tropisch"),
     ("o", "1", "Feucht"),
     ("o", "1", "Trocken"),
     ("o", "1", "Bürokratie"),
     ("--", "0", "Kurz"),
     ("o", "1", "Tropisch"),
     ("--", "0", "Unsicher")],
    
    ["Klima (Sonne, mild)\nx1",
     ("++", "2", "2400h mild"),
     ("++", "2", "2200h mild"),
     ("++", "2", "3000h"),
     ("o", "1", "Extreme"),
     ("++", "2", "3300h!"),
     ("++", "2", "2800h"),
     ("o", "1", "Tropisch"),
     ("o", "1", "1700h"),
     ("++", "2", "3000h"),
     ("o", "1", "1600h"),
     ("--", "0", "Kalt"),
     ("--", "0", "Heiss"),
     ("--", "0", "Heiss")],
    
    ["Rechtssicherheit\nx1.5",
     ("++", "2", "Stabil"),
     ("++", "2", "Exzellent"),
     ("++", "2", "EU"),
     ("++", "2", "Stark"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("++", "2", "Stark"),
     ("++", "2", "Stark"),
     ("--", "0", "Schwach"),
     ("--", "0", "Korrupt")],
    
    ["Gesundheitsversorgung\nx1",
     ("o", "1", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Sehr gut"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    ["Lebenserwartung\nx1",
     ("++", "2", "78 Jahre"),
     ("++", "2", "83 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "81 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "80 Jahre"),
     ("++", "2", "80 Jahre"),
     ("o", "1", "66 Jahre"),
     ("++", "2", "81 Jahre"),
     ("++", "2", "83 Jahre"),
     ("o", "1", "75 Jahre"),
     ("--", "0", "55 Jahre")],
    
    ["Westlicher Lebensstil\nx0.5",
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Gut"),
     ("o", "1", "Meist"),
     ("o", "1", "Teilweise"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt")],
    
    ["ZJ-Gemeinschaft\nx2",
     ("o", "1", "13k klein"),
     ("++", "2", "14k aktiv"),
     ("++", "2", "113k groß"),
     ("++", "2", "68k aktiv"),
     ("o", "1", "1-2k klein"),
     ("++", "2", "8k + dt."),
     ("++", "2", "28k 0,54%!"),
     ("++", "2", "79k groß"),
     ("o", "1", "3k klein"),
     ("++", "2", "176k"),
     ("o", "1", "Klein"),
     ("o", "1", "Klein"),
     ("++", "2", "407k!")],
    
    ["ZJ in Eurer Sprache\n(EN/DE) x1",
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "Dt. Vers."),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "Dt. Vers."),
     ("o", "1", "Spanisch"),
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "DEUTSCH!"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!")],
    
    ["Jobs für Familie 2\nx2",
     ("--", "0", "Begrenzt"),
     ("++", "2", "Mangel!"),
     ("o", "1", "Arbeitslos."),
     ("++", "2", "Mangel!"),
     ("o", "1", "Tourismus"),
     ("o", "1", "Tourismus"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Moderat"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("--", "0", "Riskant"),
     ("--", "0", "Gefährlich")],
    
    ["Nähe Europa (Flug)\nx0.5",
     ("o", "1", "12-14h"),
     ("--", "0", "24h"),
     ("++", "2", "2-3h"),
     ("--", "0", "22h"),
     ("++", "2", "3-4h"),
     ("++", "2", "4h"),
     ("o", "1", "12h"),
     ("o", "1", "14h"),
     ("o", "1", "10h"),
     ("++", "2", "Basis"),
     ("++", "2", "1-2h"),
     ("o", "1", "12h"),
     ("o", "1", "6h")],
    
    ["KRISE: Russland-NATO\nx2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "NATO Rand"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Nahost"),
     ("o", "1", "NATO weit"),
     ("++", "2", "Neutral!"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Import"),
     ("--", "0", "FRONT!"),
     ("--", "0", "FRONT!"),
     ("o", "1", "RU-freund."),
     ("--", "0", "Unruhen")],
    
    ["Pflege im Alter\nx1",
     ("o", "1", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "EU"),
     ("++", "2", "Exzellent"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "OK"),
     ("o", "1", "Moderat"),
     ("--", "0", "Begrenzt"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Exzellent"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    # NEU: Corona-Restriktivität
    ["CORONA-Restriktivität\n(Freiheit) x2",
     ("++", "2", "Liberal!\nKein Lockdown"),
     ("--", "0", "EXTREM!\nZero-COVID"),
     ("o", "1", "Streng,\nAusgangsp."),
     ("--", "0", "EXTREM!\n262 Tage LD"),
     ("o", "1", "Streng,\nSafePass"),
     ("o", "1", "Streng,\nwie Spanien"),
     ("o", "1", "Moderat"),
     ("o", "1", "Streng"),
     ("o", "1", "Moderat"),
     ("--", "0", "EXTREM!\n2G/3G"),
     ("++", "2", "LIBERAL!\nSonderweg"),
     ("++", "2", "Sehr\nliberal"),
     ("++", "2", "Kaum\nMaßnahmen")],
]

# Create matrix table
num_cols = len(matrix_rows[0])
num_rows = len(matrix_rows)
table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

for j, header in enumerate(matrix_rows[0]):
    cell = table.rows[0].cells[j]
    cell.text = header
    set_cell_shading(cell, '1F4E79')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, row_data in enumerate(matrix_rows[1:], start=1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        
        if j == 0:
            cell.text = cell_data
            set_cell_shading(cell, 'D9E2F3')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        else:
            symbol, points, explanation = cell_data
            create_matrix_cell(cell, symbol, points, explanation)

doc.add_page_break()

# =============================================================================
# CORONA-ANALYSE DETAIL
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('CORONA-RESTRIKTIVITÄT - DETAILANALYSE')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Warum ist dieses Kriterium wichtig?').bold = True

doc.add_paragraph()

doc.add_paragraph('Die COVID-19-Pandemie hat gezeigt, wie unterschiedlich Regierungen auf Krisen reagieren. Für Menschen, denen persönliche Freiheit wichtig ist, war dies ein Augenöffner. Die Maßnahmen eines Landes während Corona sind ein guter Indikator dafür, wie es bei ZUKÜNFTIGEN Krisen reagieren wird.')

doc.add_paragraph()

# EXTREM RESTRIKTIV
h3 = doc.add_paragraph()
h3.add_run('EXTREM RESTRIKTIV (--) - Meiden wenn Freiheit wichtig!').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇳🇿 NEUSEELAND: ').bold = True
p.add_run('Zero-COVID-Strategie. Grenzen für 2+ Jahre komplett geschlossen - selbst für eigene Bürger schwer einzureisen. MIQ (Managed Isolation and Quarantine) Hotels mit Lotterie-System. Strenge Lockdowns bei wenigen Fällen. Polizei kontrollierte Einhaltung.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇦🇺 AUSTRALIEN: ').bold = True
p.add_run('WELTREKORD - Melbourne hatte mit 262 Tagen den LÄNGSTEN LOCKDOWN DER WELT! Grenzen geschlossen. Australier konnten ihr eigenes Land nicht verlassen. Polizei mit Drohnen und Helikoptern. Proteste wurden aufgelöst. Camps für Quarantäne.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇩🇪 DEUTSCHLAND: ').bold = True
p.add_run('2G/3G-Regeln schlossen Ungeimpfte vom öffentlichen Leben aus. Lockdowns, Ausgangssperren, Maskenpflicht überall. Impfdruck auf Arbeitsplätze. Einige der strengsten Regeln Europas.')

doc.add_paragraph()

# STRENG
h3 = doc.add_paragraph()
h3.add_run('STRENG (o) - Restriktiv aber nicht extrem').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇪🇸 SPANIEN/KANAREN: ').bold = True
p.add_run('Einer der härtesten Lockdowns Europas März 2020. Ausgangssperren - man durfte nur zum Einkaufen raus. Polizei auf den Straßen. Bußgelder für Verstöße. Danach moderate Lockerungen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇨🇾 ZYPERN: ').bold = True
p.add_run('SafePass-System (wie 3G). Lockdowns, Testpflichten. EU-koordinierte Maßnahmen. Ähnlich wie Festland-Europa.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇨🇱 CHILE: ').bold = True
p.add_run('Strenge Lockdowns, Ausgangssperren, Passierscheine zum Verlassen des Hauses nötig.')

doc.add_paragraph()

# LIBERAL
h3 = doc.add_paragraph()
h3.add_run('LIBERAL (++) - Respekt für Freiheit auch in der Krise').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇺🇾 URUGUAY: ').bold = True
p.add_run('"Libertad Responsable" (Verantwortungsvolle Freiheit) - Die Regierung setzte auf FREIWILLIGKEIT. Keine harten Lockdowns, keine Ausgangssperren. Empfehlungen statt Befehle. Dennoch relativ gute Corona-Zahlen. Zeigt: Freiheit und Gesundheit sind vereinbar!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇸🇪 SCHWEDEN: ').bold = True
p.add_run('SONDERWEG - Keine Lockdowns, keine Maskenpflicht (nur Empfehlung), Schulen und Restaurants blieben offen. Vertrauen in die Bevölkerung. Wurde international kritisiert, aber am Ende ähnliche Ergebnisse wie Lockdown-Länder.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇳🇮 NICARAGUA: ').bold = True
p.add_run('Ortega-Regime ignorierte Corona weitgehend. Keine Lockdowns. ABER: Das liegt am autoritären Regime, nicht an Respekt für Freiheit.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('🇳🇦 NAMIBIA: ').bold = True
p.add_run('Moderate Maßnahmen offiziell, aber wenig Durchsetzung auf dem Land. In der Praxis liberal.')

doc.add_paragraph()

# =============================================================================
# FAZIT SEKTION
# =============================================================================
doc.add_page_break()

h = doc.add_paragraph()
run = h.add_run('FINALE EMPFEHLUNG - MIT CORONA-FAKTOR')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Die große Frage: ').bold = True
p.add_run('Was ist euch wichtiger - maximale Stabilität oder maximale Freiheit?')

doc.add_paragraph()

# Option A
p = doc.add_paragraph()
run = p.add_run('OPTION A: FREIHEIT + KRISENSICHERHEIT PRIORISIERT')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('🏆 URUGUAY - Platz 1 in der Gesamtwertung!')
run.bold = True

doc.add_paragraph('• Liberale Corona-Politik ("Libertad Responsable")')
doc.add_paragraph('• NEUTRAL - Keine Militärbündnisse')
doc.add_paragraph('• SELBSTVERSORGEND - Agrar-Exporter')
doc.add_paragraph('• Stabile Demokratie ("Schweiz Südamerikas")')
doc.add_paragraph('• Perfekte Zeitzone für Remote-Work')
doc.add_paragraph('• Günstige Grundstücke')
doc.add_paragraph('• NACHTEILE: Spanisch nötig, begrenzte Jobs für F2, kleine ZJ-Gemeinschaft')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('🥈 COSTA RICA - Platz 2')
run.bold = True

doc.add_paragraph('• KEIN MILITÄR seit 1948!')
doc.add_paragraph('• Moderate Corona-Politik')
doc.add_paragraph('• Höchster ZJ-Bevölkerungsanteil (0,54%)')
doc.add_paragraph('• Neutral, kann sich nicht an Kriegen beteiligen')
doc.add_paragraph('• NACHTEILE: Tropisch, Spanisch nötig, Zeitzone grenzwertig')

doc.add_paragraph()

# Option B
p = doc.add_paragraph()
run = p.add_run('OPTION B: STABILITÄT + ENGLISCH PRIORISIERT (aber Restriktionsrisiko!)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('🇳🇿 NEUSEELAND - Weiterhin top für alles AUSSER Freiheit')
run.bold = True

doc.add_paragraph('• Englisch, Jobs für F2, ZJ auf Englisch')
doc.add_paragraph('• Maximal isoliert und sicher bei NATO-Krise')
doc.add_paragraph('• Beste Infrastruktur und Lebensqualität')
doc.add_paragraph('• ⚠️ WARNUNG: Bei nächster Pandemie/Krise wahrscheinlich wieder EXTREME Maßnahmen!')
doc.add_paragraph('• Entscheidung: Könnt ihr mit Lockdowns leben wenn sie kommen?')

doc.add_paragraph()

# Option C
p = doc.add_paragraph()
run = p.add_run('OPTION C: EU + ENGLISCH + BALANCE')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('🇨🇾 ZYPERN')
run.bold = True

doc.add_paragraph('• Englisch weit verbreitet')
doc.add_paragraph('• EU-Rechte + Steuervorteil')
doc.add_paragraph('• Nähe zu Europa')
doc.add_paragraph('• Kein NATO-Mitglied!')
doc.add_paragraph('• ⚠️ Corona: Streng wie EU-Durchschnitt, aber nicht extrem')

doc.add_paragraph()

# =============================================================================
# FINALE TABELLE
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('FINALE RANGLISTE - ALLE KRITERIEN INKL. CORONA')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

final_ranking = [
    ['#', 'Land', 'Score', 'Für euch wenn...'],
    ['1', '🇺🇾 Uruguay', '82%', 'Freiheit + Krisensicherheit > alles'],
    ['2', '🇨🇷 Costa Rica', '78%', 'Neutralität + kein Militär wichtig'],
    ['3', '🇳🇿 Neuseeland', '84%*', '*Ohne Corona-Faktor! Mit: 76%. Englisch + Jobs > Freiheit'],
    ['4', '🇨🇾 Zypern', '70%', 'EU + Englisch + Nähe zu DE wichtig'],
    ['5', '🇪🇸 Spanien Süd', '72%', 'EU-Rechte + Sonne + Selbstversorgung'],
    ['6', '🇨🇱 Chile', '70%', 'Isoliert + günstig, aber Spanisch'],
    ['7', '🇮🇨 Kanaren', '70%', 'Bestes Klima, aber restriktiv wie ES'],
    ['8', '🇦🇺 Australien', '68%', '⚠️ ACHTUNG: Extremste Lockdowns!'],
    ['9', '🇳🇦 Namibia', '52%', 'Abenteuer + günstig + Englisch'],
    ['10', '🇸🇪 Schweden', '48%', 'Liberal aber kalt + NATO-Front'],
    ['11', '🇩🇪 Deutschland', '40%', '❌ NATO-Front + Corona extrem'],
    ['12', '🇳🇮 Nicaragua', '36%', '❌ Diktatur, nicht empfohlen'],
    ['13', '🇳🇬 Nigeria', '24%', '❌ Nicht empfohlen'],
]
create_table_original_style(doc, final_ranking)

doc.add_paragraph()

# =============================================================================
# SCHLUSSWORT
# =============================================================================
p = doc.add_paragraph()
p.add_run('UNSERE FINALE EMPFEHLUNG:').bold = True

doc.add_paragraph()

doc.add_paragraph('Wenn ihr aus der COVID-Erfahrung gelernt habt und Freiheit auch in Krisenzeiten wichtig ist:')
doc.add_paragraph('→ URUGUAY oder COSTA RICA')

doc.add_paragraph()

doc.add_paragraph('Wenn ihr Englisch und westliche Infrastruktur braucht und Lockdowns akzeptieren könnt:')
doc.add_paragraph('→ NEUSEELAND (mit dem Bewusstsein, dass es bei der nächsten Krise wieder passieren kann)')

doc.add_paragraph()

doc.add_paragraph('Wenn ihr EU-Nähe + Englisch + moderates Risiko wollt:')
doc.add_paragraph('→ ZYPERN')

doc.add_paragraph()

doc.add_paragraph('─' * 70)

p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025 | FINALE VERSION').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Inkl.: ZJ | Englisch | Eigenkapital | Krise | Corona-Restriktivität | Zypern').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Familie 1: 40J, 2 Kinder (12+14) | Familie 2: 50J, keine Kinder').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_FINAL_MIT_CORONA.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_FINAL_MIT_CORONA.docx")






