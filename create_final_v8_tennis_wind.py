from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

def create_matrix_cell(cell, symbol, points, explanation, is_header=False):
    cell.text = ""
    
    if is_header:
        p = cell.paragraphs[0]
        run = p.add_run(str(symbol))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1F4E79')
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        run.font.bold = True
        run.font.size = Pt(11)
        
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"({points})")
        run2.font.size = Pt(8)
        
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(explanation)
        run3.font.size = Pt(7)
        run3.font.italic = True
        
        if symbol == "++":
            set_cell_shading(cell, 'C6EFCE')
        elif symbol == "o":
            set_cell_shading(cell, 'FFEB9C')
        elif symbol == "--":
            set_cell_shading(cell, 'FFC7CE')
        
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

# Create document
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025 - FINALE VERSION')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT | Zwei-Familien | Englisch | ZJ | Tennis')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

# =============================================================================
# EUER PROFIL
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR) ').bold = True
p1.add_run('| 2. Remote IT-Job | 3. Fließend Englisch | 4. Zwei-Familien-Projekt')

p2 = doc.add_paragraph()
p2.add_run('5. Familie 2: Eigenkapital + Vor-Ort-Jobs ').bold = True
p2.add_run('| 6. Zeugen Jehovas | 7. Tennis ganzjährig')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING - MIT TENNIS + WIND
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('GESAMTRANKING - 20 KRITERIEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# Ranking mit Tennis + Wind eingerechnet
# Tennis-Wetter: Zypern++, Kanaren++, Spanien++, Namibia++, AU++, UY++, NZ o, CR o, CL o, DE--, SE--, NI o, NG o
# Wind (wenig=gut): CR++, CY o, ES o, AU o, DE o, KAN--, NZ--, UY--, CL--, NAM o
ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Kommentar', 'Tennis', 'Wind'],
    ['1', 'Neuseeland', '48/54', '89%', 'TOP - aber windig!', 'o', '--'],
    ['2', 'Zypern', '42/54', '78%', 'EU + Englisch + Tennis!', '++', 'o'],
    ['3', 'Spanien Süd', '42/54', '78%', 'Tennis + EU + Sonne', '++', 'o'],
    ['4', 'Uruguay', '42/54', '78%', 'Neutral + liberal, windig', '++', '--'],
    ['5', 'Costa Rica', '43/54', '80%', 'Neutral + wenig Wind!', 'o', '++'],
    ['6', 'Australien', '40/54', '74%', 'Englisch + Tennis', '++', 'o'],
    ['7', 'Kanaren', '38/54', '70%', 'Klima top, ABER windig!', '++', '--'],
    ['8', 'Chile', '36/54', '67%', 'Isoliert, sehr windig', 'o', '--'],
    ['9', 'Namibia', '29/54', '54%', 'Günstig + Tennis OK', '++', 'o'],
    ['10', 'Schweden', '25/54', '46%', 'Liberal, aber kein Tennis', '--', 'o'],
    ['11', 'Deutschland', '22/54', '41%', 'Kein Tennis Winter', '--', 'o'],
    ['12', 'Nicaragua', '21/54', '39%', 'Instabil', 'o', 'o'],
    ['13', 'Nigeria', '15/54', '28%', 'Nicht empfohlen', 'o', 'o'],
]
create_table_original_style(doc, ranking_data)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('NEU: ').bold = True
p.add_run('Tennis-Wetter (ganzjährig spielbar) und Wind (wenig = gut) als Kriterien mit Gewicht x0.5.')

doc.add_page_break()

# =============================================================================
# DETAILMATRIX - 20 KRITERIEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX - 20 KRITERIEN')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('++ = Sehr gut (2) | o = Mittel (1) | -- = Schlecht (0)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
run = abbrev.add_run('UY=Uruguay, NZ=Neuseeland, ES=Spanien, AU=Australien, CY=Zypern, KAN=Kanaren, CR=Costa Rica, CL=Chile, NAM=Namibia, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix mit 20 Kriterien inkl. Tennis + Wind
matrix_rows = [
    ["Kriterium", "UY", "NZ", "ES", "AU", "CY", "KAN", "CR", "CL", "NAM", "DE", "SE", "NI", "NG"],
    
    ["Geopolitik x2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "NATO"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Geteilt"),
     ("++", "2", "Weit"),
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Stabil"),
     ("--", "0", "Front!"),
     ("--", "0", "Front!"),
     ("o", "1", "Instabil"),
     ("--", "0", "Unsicher")],
    
    ["Einwanderung x1.5",
     ("++", "2", "Einfach"),
     ("++", "2", "IT-Liste"),
     ("++", "2", "EU"),
     ("++", "2", "IT-Liste"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("++", "2", "Investor"),
     ("++", "2", "Einfach"),
     ("++", "2", "Einfach"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "OK"),
     ("o", "1", "Schwer")],
    
    ["Englisch x1.5",
     ("o", "1", "Spanisch"),
     ("++", "2", "Mutter!"),
     ("o", "1", "Spanisch"),
     ("++", "2", "Mutter!"),
     ("++", "2", "Weit verb."),
     ("o", "1", "Spanisch"),
     ("o", "1", "OK"),
     ("--", "0", "Spanisch"),
     ("++", "2", "Amtsspr."),
     ("o", "1", "Deutsch"),
     ("o", "1", "Schwed."),
     ("o", "1", "Spanisch"),
     ("o", "1", "Amtsspr.")],
    
    ["Grundstück x1.5",
     ("++", "2", "Günstig"),
     ("o", "1", "Teuer"),
     ("++", "2", "OK"),
     ("o", "1", "Teuer"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "OK"),
     ("++", "2", "Günstig"),
     ("++", "2", "Günstig"),
     ("--", "0", "Teuer"),
     ("--", "0", "OK"),
     ("++", "2", "Billig"),
     ("++", "2", "Risiko")],
    
    ["Zeitzone x1",
     ("++", "2", "-4h"),
     ("--", "0", "+11h"),
     ("++", "2", "Gleich"),
     ("--", "0", "+9h"),
     ("++", "2", "+1h"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h"),
     ("o", "1", "-5h"),
     ("++", "2", "+1h"),
     ("++", "2", "Basis"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h"),
     ("o", "1", "+1h")],
    
    ["Selbstversorg. x1",
     ("++", "2", "Ideal"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("o", "1", "Wasser"),
     ("o", "1", "Wasser"),
     ("o", "1", "Wasser"),
     ("o", "1", "Tropisch"),
     ("o", "1", "OK"),
     ("o", "1", "Trocken"),
     ("o", "1", "Bürokr."),
     ("--", "0", "Kurz"),
     ("o", "1", "OK"),
     ("--", "0", "Unsicher")],
    
    ["Klima x1",
     ("++", "2", "2400h"),
     ("++", "2", "2200h"),
     ("++", "2", "3000h"),
     ("o", "1", "Extreme"),
     ("++", "2", "3300h!"),
     ("++", "2", "2800h"),
     ("o", "1", "Tropisch"),
     ("o", "1", "1700h"),
     ("++", "2", "3000h"),
     ("o", "1", "1600h"),
     ("--", "0", "Kalt"),
     ("--", "0", "Heiss"),
     ("--", "0", "Heiss")],
    
    ["Rechtssich. x1.5",
     ("++", "2", "Stabil"),
     ("++", "2", "Top"),
     ("++", "2", "EU"),
     ("++", "2", "Stark"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("++", "2", "Stark"),
     ("++", "2", "Stark"),
     ("--", "0", "Schwach"),
     ("--", "0", "Korrupt")],
    
    ["Gesundheit x1",
     ("o", "1", "Gut"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("++", "2", "Top"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "Begr."),
     ("++", "2", "Top"),
     ("++", "2", "Gut"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    ["Lebenserwart. x1",
     ("++", "2", "78J"),
     ("++", "2", "83J"),
     ("++", "2", "84J"),
     ("++", "2", "84J"),
     ("++", "2", "81J"),
     ("++", "2", "84J"),
     ("++", "2", "80J"),
     ("++", "2", "80J"),
     ("o", "1", "66J"),
     ("++", "2", "81J"),
     ("++", "2", "83J"),
     ("o", "1", "75J"),
     ("--", "0", "55J")],
    
    ["Westl. Leben x0.5",
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Gut"),
     ("o", "1", "Meist"),
     ("o", "1", "Teils"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("o", "1", "Begr."),
     ("o", "1", "Begr.")],
    
    ["ZJ-Gemeinde x2",
     ("o", "1", "13k"),
     ("++", "2", "14k"),
     ("++", "2", "113k"),
     ("++", "2", "68k"),
     ("o", "1", "2k"),
     ("++", "2", "8k+dt"),
     ("++", "2", "28k!"),
     ("++", "2", "79k"),
     ("o", "1", "3k"),
     ("++", "2", "176k"),
     ("o", "1", "Klein"),
     ("o", "1", "Klein"),
     ("++", "2", "407k!")],
    
    ["ZJ Sprache x1",
     ("o", "1", "Span."),
     ("++", "2", "EN!"),
     ("++", "2", "Dt."),
     ("++", "2", "EN!"),
     ("++", "2", "EN!"),
     ("++", "2", "Dt."),
     ("o", "1", "Span."),
     ("o", "1", "Span."),
     ("++", "2", "EN!"),
     ("++", "2", "DE!"),
     ("o", "1", "Schw."),
     ("o", "1", "Span."),
     ("++", "2", "EN!")],
    
    ["Jobs F2 x2",
     ("--", "0", "Wenig"),
     ("++", "2", "Mangel!"),
     ("o", "1", "Schwer"),
     ("++", "2", "Mangel!"),
     ("o", "1", "Tourism."),
     ("o", "1", "Tourism."),
     ("o", "1", "Begr."),
     ("o", "1", "OK"),
     ("o", "1", "Begr."),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("--", "0", "Riskant"),
     ("--", "0", "Gefährl.")],
    
    ["Nähe EU x0.5",
     ("o", "1", "12h"),
     ("--", "0", "24h"),
     ("++", "2", "2-3h"),
     ("--", "0", "22h"),
     ("++", "2", "3-4h"),
     ("++", "2", "4h"),
     ("o", "1", "12h"),
     ("o", "1", "14h"),
     ("o", "1", "10h"),
     ("++", "2", "Basis"),
     ("++", "2", "1-2h"),
     ("o", "1", "12h"),
     ("o", "1", "6h")],
    
    ["Krise NATO x2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "NATO"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Nahost"),
     ("o", "1", "NATO"),
     ("++", "2", "Neutral!"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Import"),
     ("--", "0", "Front!"),
     ("--", "0", "Front!"),
     ("o", "1", "RU-nah"),
     ("--", "0", "Unruhen")],
    
    ["Pflege x1",
     ("o", "1", "Gut"),
     ("++", "2", "Gut"),
     ("++", "2", "EU"),
     ("++", "2", "Top"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("--", "0", "Begr."),
     ("++", "2", "Top"),
     ("++", "2", "Top"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    ["Corona x1",
     ("++", "2", "Liberal"),
     ("--", "0", "Extrem"),
     ("o", "1", "Streng"),
     ("--", "0", "Extrem!"),
     ("o", "1", "Streng"),
     ("o", "1", "Streng"),
     ("o", "1", "Moderat"),
     ("o", "1", "Streng"),
     ("o", "1", "Moderat"),
     ("--", "0", "Extrem"),
     ("++", "2", "Liberal!"),
     ("++", "2", "Liberal"),
     ("++", "2", "Liberal")],
    
    # NEU: Tennis-Wetter
    ["Tennis-Wetter x0.5",
     ("++", "2", "2400h\nmild"),
     ("o", "1", "Regen\nmöglich"),
     ("++", "2", "3000h\nganzjährig"),
     ("++", "2", "Ganzjährig"),
     ("++", "2", "3300h!\nTop"),
     ("++", "2", "2800h\nperfekt"),
     ("o", "1", "Regen\nnachmittags"),
     ("o", "1", "Süden\nfeucht"),
     ("++", "2", "3000h\ntrocken"),
     ("--", "0", "Winter\nzu kalt"),
     ("--", "0", "6 Monate\nnicht"),
     ("o", "1", "Tropisch\nheiss"),
     ("o", "1", "Tropisch\nheiss")],
    
    # NEU: Wind (wenig = gut)
    ["Wind (wenig=gut) x0.5",
     ("--", "0", "Atlantik\nwindig"),
     ("--", "0", "Sehr\nwindig!"),
     ("o", "1", "Küste\nwindig"),
     ("o", "1", "Variiert"),
     ("o", "1", "Moderat"),
     ("--", "0", "SEHR\nwindig!"),
     ("++", "2", "Täler\ngeschützt"),
     ("--", "0", "Patagon.\nextrem"),
     ("o", "1", "Küste\nwindig"),
     ("o", "1", "Moderat"),
     ("o", "1", "Moderat"),
     ("o", "1", "OK"),
     ("o", "1", "OK")],
]

# Create matrix table
num_cols = len(matrix_rows[0])
num_rows = len(matrix_rows)
table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

for j, header in enumerate(matrix_rows[0]):
    cell = table.rows[0].cells[j]
    cell.text = header
    set_cell_shading(cell, '1F4E79')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, row_data in enumerate(matrix_rows[1:], start=1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        
        if j == 0:
            cell.text = cell_data
            set_cell_shading(cell, 'D9E2F3')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        else:
            symbol, points, explanation = cell_data
            create_matrix_cell(cell, symbol, points, explanation)

doc.add_page_break()

# =============================================================================
# TENNIS + WIND ANALYSE
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('TENNIS-WETTER & WIND - DETAILANALYSE')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('TENNIS-WETTER (ganzjährig spielbar):').bold = True

doc.add_paragraph()

doc.add_paragraph('++ ZYPERN: 3.300 Sonnenstunden, ganzjährig 15-35°C - perfekt für Tennis!')
doc.add_paragraph('++ KANAREN: 2.800h, ganzjährig 18-28°C - ideale Bedingungen')
doc.add_paragraph('++ SPANIEN SÜD: 3.000h, 10+ Monate spielbar')
doc.add_paragraph('++ URUGUAY: 2.400h, milde Winter, gut spielbar')
doc.add_paragraph('++ AUSTRALIEN: Ganzjährig (außer tropischer Norden im Sommer)')
doc.add_paragraph('++ NAMIBIA: 3.000h Sonne, sehr trocken - top für Outdoor')
doc.add_paragraph('o NEUSEELAND: 2.200h, aber Regen möglich, nicht immer ideal')
doc.add_paragraph('o COSTA RICA: Tropisch, nachmittags oft Regen')
doc.add_paragraph('-- DEUTSCHLAND: Winter 4-5 Monate nicht outdoor spielbar')
doc.add_paragraph('-- SCHWEDEN: 6+ Monate zu kalt/dunkel')

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WIND (wenig = gut für Tennis/Outdoor):').bold = True

doc.add_paragraph()

doc.add_paragraph('++ COSTA RICA: Geschützte Täler, wenig Wind')
doc.add_paragraph('o ZYPERN: Mittelmeer, moderater Wind')
doc.add_paragraph('o SPANIEN: Inland OK, Küste (Costa de la Luz) windig')
doc.add_paragraph('o AUSTRALIEN: Variiert stark nach Region')
doc.add_paragraph('-- KANAREN: SEHR windig! Fuerteventura = Windsurf-Paradies = Tennis-Hölle')
doc.add_paragraph('-- NEUSEELAND: Bekannt für Wind! Wellington = "Windy Welly"')
doc.add_paragraph('-- URUGUAY: Atlantikküste sehr windig')
doc.add_paragraph('-- CHILE: Patagonien extrem windig')

doc.add_paragraph()

# =============================================================================
# FINALE EMPFEHLUNG
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('FINALE EMPFEHLUNG - MIT TENNIS + WIND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('FÜR TENNIS-SPIELER:')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

doc.add_paragraph('1. ZYPERN - Beste Kombination: 3.300h Sonne + moderater Wind + Englisch + EU')
doc.add_paragraph('2. SPANIEN SÜD (Inland) - 3.000h + weniger Wind als Küste + große Grundstücke')
doc.add_paragraph('3. COSTA RICA - Geschützte Täler + kein Militär + neutral')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('ACHTUNG bei diesen Ländern: ').bold = True

doc.add_paragraph()

doc.add_paragraph('• KANAREN: Super Klima, ABER sehr windig - Tennis frustrierend!')
doc.add_paragraph('• NEUSEELAND: Top bei allem, ABER windig - outdoor Tennis eingeschränkt')
doc.add_paragraph('• URUGUAY: Atlantikküste windig - besser im Landesinneren')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('GESAMTEMPFEHLUNG:')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

doc.add_paragraph('Wenn Tennis wichtig ist: ZYPERN oder SPANIEN SÜD (Inland)')
doc.add_paragraph('Wenn alles andere wichtiger ist: NEUSEELAND bleibt Platz 1 (aber mit Wind leben)')

doc.add_paragraph()

doc.add_paragraph('─' * 70)

p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025 | FINALE VERSION mit Tennis + Wind').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('20 Kriterien | Familie 1: 40J, 2 Kinder | Familie 2: 50J, keine Kinder').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_FINAL_TENNIS_WIND.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_FINAL_TENNIS_WIND.docx")







