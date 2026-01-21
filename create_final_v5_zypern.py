from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table_original_style(doc, data, header_color='1F4E79'):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    return table

def create_matrix_cell(cell, symbol, points, explanation, is_header=False):
    """Create a matrix cell with symbol, points, and explanation"""
    cell.text = ""
    
    if is_header:
        p = cell.paragraphs[0]
        run = p.add_run(str(symbol))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1F4E79')
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(symbol)
        run.font.bold = True
        run.font.size = Pt(11)
        
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"({points})")
        run2.font.size = Pt(8)
        
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(explanation)
        run3.font.size = Pt(7)
        run3.font.italic = True
        
        if symbol == "++":
            set_cell_shading(cell, 'C6EFCE')
        elif symbol == "o":
            set_cell_shading(cell, 'FFEB9C')
        elif symbol == "--":
            set_cell_shading(cell, 'FFC7CE')
        
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

# Create document
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AUSWANDERUNGSANALYSE 2025')
run.bold = True
run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PERSONALISIERT FÜR EUER PROFIL')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('500k+ EUR Kapital | Remote IT (Data Architect) | Zwei-Familien-Projekt | Handwerklich kompetent')
run.font.size = Pt(10)
run.font.italic = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('ERWEITERT: Fließend Englisch | Zeugen Jehovas | Familie 2: Viel Eigenkapital + Vor-Ort-Jobs')
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# =============================================================================
# EUER PROFIL - STÄRKEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('EUER PROFIL - STÄRKEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.add_run('1. Kapital (500k+ EUR): ').bold = True
p1.add_run('Öffnet praktisch alle Türen. Reicht für großes Grundstück + Hausbau + Infrastruktur + Reserve.')

p2 = doc.add_paragraph()
p2.add_run('2. Remote IT-Job (Familie 1): ').bold = True
p2.add_run('Stabiles Einkommen unabhängig vom lokalen Arbeitsmarkt. Data Architect ist gefragter Beruf, der auch Visa-Optionen öffnet (NZ, AU).')

p3 = doc.add_paragraph()
p3.add_run('3. FLIESSEND ENGLISCH (beide Familien): ').bold = True
p3.add_run('Öffnet englischsprachige Länder! NZ und AU werden realistisch. Keine Sprachbarriere = sofortige Integration.')

p4 = doc.add_paragraph()
p4.add_run('4. Zwei-Familien-Projekt: ').bold = True
p4.add_run('Geteilte Kosten, geteilte Arbeit, soziales Netz von Tag 1. Größere Grundstücke werden erschwinglich.')

p5 = doc.add_paragraph()
p5.add_run('5. Handwerkliche Kompetenz: ').bold = True
p5.add_run('Spart enorm bei Hausbau und Infrastruktur.')

p6 = doc.add_paragraph()
p6.add_run('6. Familie 2 - Eigenkapital + Vor-Ort-Jobs: ').bold = True
p6.add_run('Flexibel für lokalen Arbeitsmarkt, keine Zeitzonenbindung.')

p7 = doc.add_paragraph()
p7.add_run('7. Zeugen Jehovas: ').bold = True
p7.add_run('Weltweites Netzwerk, sofortige Gemeinschaft in jedem Land.')

doc.add_paragraph()

# =============================================================================
# GESAMTRANKING - MIT ZYPERN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('GESAMTRANKING - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

ranking_data = [
    ['#', 'Land', 'Punkte', '%', 'Für euer Profil'],
    ['1', 'Neuseeland', '46 / 50', '92%', 'TOP - Englisch + Jobs + Krise + Pflege!'],
    ['2', 'Spanien (Süden)', '38 / 50', '76%', 'EU-Rechte + Sonne + gute Pflege'],
    ['3', 'Australien', '38 / 50', '76%', 'Englisch + Jobs + exzellente Pflege'],
    ['4', 'Zypern', '37 / 50', '74%', 'NEU: EU + Englisch + Sonne + Pflege!'],
    ['5', 'Kanarische Inseln', '37 / 50', '74%', 'EU + bestes Klima + Pflege'],
    ['6', 'Uruguay', '39 / 50', '78%', 'Krisensicher! Pflege OK mit Kapital'],
    ['7', 'Costa Rica', '39 / 50', '78%', 'Kein Militär + krisensicher!'],
    ['8', 'Chile (Mitte)', '37 / 50', '74%', 'Isoliert + selbstversorgend'],
    ['9', 'Deutschland', '24 / 50', '48%', 'Exzellente Pflege, aber NATO-FRONT!'],
    ['10', 'Schweden', '22 / 50', '44%', 'Beste Pflege, aber NATO-Ostsee!'],
    ['11', 'Namibia', '26 / 50', '52%', 'Englisch + günstig, Pflege begrenzt'],
    ['12', 'Nicaragua', '16 / 50', '32%', 'Instabil, keine Pflege-Infrastruktur'],
    ['13', 'Nigeria', '10 / 50', '20%', 'Nicht empfohlen - keine Pflege'],
]
create_table_original_style(doc, ranking_data)

doc.add_page_break()

# =============================================================================
# DETAILMATRIX - MIT ZYPERN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILMATRIX')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

legend = doc.add_paragraph()
legend.add_run('Symbole: ').bold = True
legend.add_run('++ = Sehr gut (2 Pkt) | o = Mittel (1 Pkt) | -- = Schlecht (0 Pkt)')

doc.add_paragraph()

abbrev = doc.add_paragraph()
run = abbrev.add_run('Länder: UY=Uruguay, NZ=Neuseeland, ES-S=Spanien Süd, AU=Australien, CY=Zypern, KAN=Kanaren, CR=Costa Rica, CL=Chile, NAM=Namibia, DE=Deutschland, SE=Schweden, NI=Nicaragua, NG=Nigeria')
run.font.size = Pt(8)

doc.add_paragraph()

# Matrix rows - now includes Zypern (CY)
matrix_rows = [
    ["Kriterium (Gewicht)", "UY", "NZ", "ES-S", "AU", "CY", "KAN", "CR", "CL", "NAM", "DE", "SE", "NI", "NG"],
    
    ["Geopolitische Sicherheit\nx2",
     ("++", "2", "Neutral"),
     ("++", "2", "Isoliert"),
     ("o", "1", "EU-Rand"),
     ("o", "1", "AUKUS"),
     ("o", "1", "Geteilt"),
     ("++", "2", "Weit weg"),
     ("++", "2", "Kein Militär"),
     ("++", "2", "Isoliert"),
     ("o", "1", "Stabil"),
     ("--", "0", "NATO-Front"),
     ("--", "0", "NATO-Ostsee"),
     ("o", "1", "Instabil"),
     ("--", "0", "Unsicher")],
    
    ["Einwanderung (500k + IT)\nx1.5",
     ("++", "2", "Sehr einfach"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "IT auf Liste"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "EU-Recht"),
     ("++", "2", "Investor OK"),
     ("++", "2", "Einfach"),
     ("++", "2", "Einfach"),
     ("++", "2", "EU"),
     ("++", "2", "EU"),
     ("o", "1", "Einfach"),
     ("o", "1", "Kompliziert")],
    
    ["Arbeiten mit Englisch\nx1.5",
     ("o", "1", "Spanisch\nnötig"),
     ("++", "2", "Mutter-\nsprache"),
     ("o", "1", "Spanisch\nbesser"),
     ("++", "2", "Mutter-\nsprache"),
     ("++", "2", "Weit\nverbreitet!"),
     ("o", "1", "Spanisch\nbesser"),
     ("o", "1", "Viel Englisch"),
     ("--", "0", "Nur Spanisch"),
     ("++", "2", "Amts-\nsprache!"),
     ("++", "2", "Deutsch\nnötig"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Etwas\nEnglisch"),
     ("o", "1", "Amtssprache")],
    
    ["Grundstück 10+ ha\nx1.5",
     ("++", "2", "50-150k USD"),
     ("o", "1", "Teuer 500k+"),
     ("++", "2", "Hinterland\nOK"),
     ("o", "1", "Teuer"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Sehr\nbegrenzt"),
     ("o", "1", "Gestiegen"),
     ("++", "2", "Sehr günstig"),
     ("++", "2", "Sehr\ngünstig!"),
     ("--", "0", "Sehr teuer"),
     ("--", "0", "Norden OK"),
     ("++", "2", "Billig"),
     ("++", "2", "Risiko")],
    
    ["Remote-Work Zeitzone\nx1",
     ("++", "2", "-4h ideal"),
     ("--", "0", "+11h schwer"),
     ("++", "2", "Gleich"),
     ("--", "0", "+9h schwer"),
     ("++", "2", "+1h gut"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "-5h OK"),
     ("++", "2", "+1h gut"),
     ("++", "2", "Basis"),
     ("++", "2", "Gleich"),
     ("o", "1", "-7h OK"),
     ("o", "1", "+1h gut")],
    
    ["Selbstversorgung\nx1",
     ("++", "2", "Ideal"),
     ("++", "2", "Nordinsel top"),
     ("++", "2", "Sehr gut"),
     ("o", "1", "Wasser knapp"),
     ("o", "1", "Wasser\nknapp"),
     ("o", "1", "Wasser\nknapp"),
     ("o", "1", "Tropisch"),
     ("o", "1", "Feucht/kühl"),
     ("o", "1", "Trocken"),
     ("o", "1", "Bürokratie"),
     ("--", "0", "Kurze Saison"),
     ("o", "1", "Tropisch"),
     ("--", "0", "Unsicher")],
    
    ["Klima (Sonne, mild)\nx1",
     ("++", "2", "2400h mild"),
     ("++", "2", "2200h mild"),
     ("++", "2", "3000h"),
     ("o", "1", "Extreme"),
     ("++", "2", "3300h!\nMittelmeer"),
     ("++", "2", "2800h\nperfekt"),
     ("o", "1", "Tropisch"),
     ("o", "1", "1700h feucht"),
     ("++", "2", "3000h\ntrocken"),
     ("o", "1", "1600h"),
     ("--", "0", "Kalt dunkel"),
     ("--", "0", "Heiss feucht"),
     ("--", "0", "Heiss feucht")],
    
    ["Rechtssicherheit\nx1.5",
     ("++", "2", "Stabil"),
     ("++", "2", "Exzellent"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "Stark"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "EU-Standard"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "OK"),
     ("++", "2", "Stark"),
     ("++", "2", "Stark"),
     ("--", "0", "Schwach"),
     ("--", "0", "Korrupt")],
    
    ["Gesundheitsversorgung\nx1",
     ("o", "1", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Gut"),
     ("++", "2", "Sehr gut"),
     ("++", "2", "Gut"),
     ("o", "1", "OK"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Exzellent"),
     ("++", "2", "Sehr gut"),
     ("--", "0", "Schwach"),
     ("--", "0", "Schwach")],
    
    ["Lebenserwartung\nx1",
     ("++", "2", "78 Jahre"),
     ("++", "2", "83 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "81 Jahre"),
     ("++", "2", "84 Jahre"),
     ("++", "2", "80 Jahre"),
     ("++", "2", "80 Jahre"),
     ("o", "1", "66 Jahre"),
     ("++", "2", "81 Jahre"),
     ("++", "2", "83 Jahre"),
     ("o", "1", "75 Jahre"),
     ("--", "0", "55 Jahre")],
    
    ["Westlicher Lebensstil\nx0.5",
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("++", "2", "Gut"),
     ("o", "1", "Meist"),
     ("o", "1", "Teilweise"),
     ("++", "2", "Voll"),
     ("++", "2", "Voll"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Begrenzt")],
    
    ["ZJ-Gemeinschaft\nx2",
     ("o", "1", "13k klein"),
     ("++", "2", "14k aktiv"),
     ("++", "2", "113k groß"),
     ("++", "2", "68k aktiv"),
     ("o", "1", "1-2k klein"),
     ("++", "2", "8k + dt.\nVersamml."),
     ("++", "2", "28k 0,54%!"),
     ("++", "2", "79k groß"),
     ("o", "1", "3k klein"),
     ("++", "2", "176k"),
     ("o", "1", "Klein"),
     ("o", "1", "Klein"),
     ("++", "2", "407k!")],
    
    ["ZJ in Eurer Sprache (EN/DE)\nx1",
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "Dt. Versam.\nvorhanden"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "Dt. Versam.\nvorhanden"),
     ("o", "1", "Spanisch,\nEN teils"),
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!"),
     ("++", "2", "DEUTSCH!"),
     ("o", "1", "Schwedisch"),
     ("o", "1", "Spanisch"),
     ("++", "2", "ENGLISCH!")],
    
    ["Jobs für Familie 2\nx2",
     ("--", "0", "Sehr\nbegrenzt"),
     ("++", "2", "Arbeits-\nkräftemangel"),
     ("o", "1", "Hohe\nArbeitslos."),
     ("++", "2", "Arbeits-\nkräftemangel"),
     ("o", "1", "Tourismus"),
     ("o", "1", "Tourismus"),
     ("o", "1", "Begrenzt"),
     ("o", "1", "Moderat"),
     ("o", "1", "Begrenzt"),
     ("++", "2", "Gut"),
     ("++", "2", "Gut"),
     ("--", "0", "Riskant"),
     ("--", "0", "Gefährlich")],
    
    ["Nähe Europa (Flug)\nx0.5",
     ("o", "1", "12-14h"),
     ("--", "0", "24h"),
     ("++", "2", "2-3h"),
     ("--", "0", "22h"),
     ("++", "2", "3-4h"),
     ("++", "2", "4h"),
     ("o", "1", "12h"),
     ("o", "1", "14h"),
     ("o", "1", "10h"),
     ("++", "2", "Basis"),
     ("++", "2", "1-2h"),
     ("o", "1", "12h"),
     ("o", "1", "6h")],
    
    ["KRISE: Russland-NATO\nx2",
     ("++", "2", "Neutral,\nselbstvers."),
     ("++", "2", "Isoliert,\nselbstvers."),
     ("o", "1", "NATO, aber\nRand"),
     ("o", "1", "AUKUS,\naber weit"),
     ("o", "1", "EU/NATO,\nNahost nah"),
     ("o", "1", "NATO, aber\nsehr weit"),
     ("++", "2", "Neutral,\nkein Militär"),
     ("++", "2", "Isoliert,\nselbstvers."),
     ("o", "1", "Neutral,\naber Import"),
     ("--", "0", "NATO-Front!\nSehr betroffen"),
     ("--", "0", "NATO-Front!\nOstsee"),
     ("o", "1", "Russland-\nfreundlich"),
     ("--", "0", "Import-abh.\nUnruhen")],
    
    ["Pflege im Alter\n(mit Kapital) x1",
     ("o", "1", "Gut für\nLateinam."),
     ("++", "2", "Sehr gut,\nwestl. Std."),
     ("++", "2", "EU-Standard,\nPflegeheime"),
     ("++", "2", "Exzellent,\ntop System"),
     ("++", "2", "EU-Standard,\nPrivatkliniken"),
     ("++", "2", "EU-Standard,\nExpat-Pflege"),
     ("o", "1", "OK mit\nKapital"),
     ("o", "1", "Moderat,\nPrivatkliniken"),
     ("--", "0", "Begrenzt,\nFamilie nötig"),
     ("++", "2", "Exzellent,\nPflegevers."),
     ("++", "2", "Exzellent,\nbekannt gut"),
     ("--", "0", "Schwach,\nkeine Infra"),
     ("--", "0", "Sehr schwach,\nnur Familie")],
]

# Create matrix table
num_cols = len(matrix_rows[0])
num_rows = len(matrix_rows)
table = doc.add_table(rows=num_rows, cols=num_cols)
table.style = 'Table Grid'

for j, header in enumerate(matrix_rows[0]):
    cell = table.rows[0].cells[j]
    cell.text = header
    set_cell_shading(cell, '1F4E79')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, row_data in enumerate(matrix_rows[1:], start=1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        
        if j == 0:
            cell.text = cell_data
            set_cell_shading(cell, 'D9E2F3')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        else:
            symbol, points, explanation = cell_data
            create_matrix_cell(cell, symbol, points, explanation)

doc.add_page_break()

# =============================================================================
# DETAILANALYSEN - VOLLSTÄNDIG MIT ALLEN KRITERIEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('DETAILANALYSEN - ANGEPASST AN EUER PROFIL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

# =============================================================================
# PLATZ 1: NEUSEELAND - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 1: Neuseeland (46 Punkte / 92%) - KLARE EMPFEHLUNG')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run("Hawke's Bay, Nelson/Tasman, Bay of Plenty, Waikato")

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM NEUSEELAND FÜR EUER PROFIL JETZT PERFEKT IST:').bold = True

doc.add_paragraph()

# Sprache/Englisch
p = doc.add_paragraph()
p.add_run('Euer Englisch ändert alles: ').bold = True
p.add_run('Mit fließendem Englisch entfällt die größte Hürde! Ihr könnt ab Tag 1 kommunizieren, arbeiten, Kinder in die Schule schicken. Die Versammlungen der Zeugen Jehovas sind sofort auf Englisch besuchbar - keine Übergangszeit nötig.')

doc.add_paragraph()

# Familie 1 IT
p = doc.add_paragraph()
p.add_run('Familie 1 - IT-Skills öffnen Türen: ').bold = True
p.add_run('Data Architect ist auf der New Zealand Skilled Occupation List! Mit nachgewiesenem Remote-Einkommen und euren Qualifikationen habt ihr sehr gute Chancen auf ein Skilled Migrant Visa.')

doc.add_paragraph()

# Familie 2 Jobs
p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt ist perfekt: ').bold = True
p.add_run('Neuseeland hat ARBEITSKRÄFTEMANGEL in vielen Bereichen: Handwerk, Landwirtschaft, Tourismus, Gesundheit. Mit fließendem Englisch ist der Zugang zum Arbeitsmarkt sofort möglich. Mit dem Eigenkapital könnt ihr auch ein eigenes kleines Business starten.')

doc.add_paragraph()

# ZJ
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas in Neuseeland: ').bold = True
p.add_run('~14.000 aktive Verkündiger in ~175 Versammlungen. Gut verteilt, auch in ländlichen Gebieten. Die Kiwi-Mentalität ist offen und freundlich - Integration in die Versammlung wird leicht fallen. Zweigbüro in Auckland ist gut organisiert.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('MAXIMAL ISOLIERT - Neuseeland liegt am Ende der Welt, weit weg von allen Konfliktzonen. Kein NATO-Mitglied, keine Militärbündnisse die zu Konflikten führen. Stabile Demokratie seit über 150 Jahren, friedliche Gesellschaft.')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('IDEAL - Neuseeland ist maximal weit weg vom Konfliktgebiet. Kein NATO-Mitglied, keine Verpflichtung zur Beteiligung. Das Land ist zu 80%+ selbstversorgend bei Lebensmitteln (Fleisch, Milch, Gemüse). Keine Abhängigkeit von russischem Gas/Öl. Innere Unruhen sehr unwahrscheinlich - stabile, homogene Gesellschaft. Lebensmittelpreise würden weniger steigen als in Europa.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('2.200 Sonnenstunden pro Jahr. Mildes, gemäßigtes Klima ohne Extreme. Nordinsel: Subtropisch im Norden, gemäßigt im Süden. Südinsel: Kühler, alpiner. Keine Hitzewellen wie Australien, keine Kälte wie Nordeuropa. Viel Grün, ausreichend Regen - ideal für Selbstversorgung.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('IDEAL - Besonders die Nordinsel. Ganzjährige Anbausaison möglich. Gute Böden, ausreichend Niederschlag. Keine Wasserknappheit wie Australien. Tierhaltung unproblematisch. Das Land ist dünn besiedelt - echte Selbstversorgung ist hier realistisch.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. Englischsprachig, britische Wurzeln. Supermärkte, Internet, Infrastruktur wie in Europa. Ihr vermisst nichts vom gewohnten Lebensstandard. Gleichzeitig entspannter und naturverbundener als Europa.')

doc.add_paragraph()

# RECHTSSICHERHEIT
p = doc.add_paragraph()
p.add_run('Rechtssicherheit/Eigentum: ').bold = True
p.add_run('EXZELLENT - Eines der besten Rechtssysteme weltweit. Transparente Bürokratie, wenig Korruption. Eigentumsrechte werden vollständig respektiert. Kaufvertrag = sicher. Keine Gefahr willkürlicher Enteignungen.')

doc.add_paragraph()

# GESUNDHEIT
p = doc.add_paragraph()
p.add_run('Gesundheitsversorgung: ').bold = True
p.add_run('SEHR GUT - Öffentliches Gesundheitssystem (ACC) für Unfälle kostenlos. Gute Krankenhäuser in allen größeren Städten. Private Zusatzversicherung empfohlen für schnellere Behandlung.')

doc.add_paragraph()

# PFLEGE IM ALTER
p = doc.add_paragraph()
p.add_run('Pflege im Alter: ').bold = True
p.add_run('SEHR GUT - Westlicher Standard. Pflegeheime (Rest Homes, Retirement Villages) in guter Qualität verfügbar. Mit Kapital: Private Pflegeeinrichtungen mit hohem Standard. Kosten: ca. 1.000-2.000 NZD/Woche für Vollzeitpflege. Häusliche Pflege auch möglich. Keine kulturelle Erwartung, dass Kinder pflegen müssen - institutionelle Pflege ist normal und akzeptiert.')

doc.add_paragraph()

# Zeitzone
p = doc.add_paragraph()
p.add_run('Das Zeitzonenproblem - und die Lösung: ').bold = True
p.add_run('Neuseeland ist +11-12h vor Deutschland. Für Familie 1 (Remote IT): Async-Arbeit vereinbaren (keine Live-Meetings), oder neuseeländische/australische Kunden aufbauen. Familie 2 arbeitet VOR ORT - Zeitzone ist komplett irrelevant!')

doc.add_paragraph()

# Zahlen
p = doc.add_paragraph()
p.add_run('Konkrete Zahlen für euer Projekt:').bold = True
doc.add_paragraph("• 10-15 Hektar in Hawke's Bay oder Nelson: ca. 300.000-600.000 NZD")
doc.add_paragraph('• Hausbau (2 Familienhäuser, je 120m², Qualität): ca. 300.000-500.000 NZD')
doc.add_paragraph('• Solar-Anlage + Wassersystem: ca. 50.000 NZD')
doc.add_paragraph('• Fahrzeuge (2 Pickups): ca. 60.000 NZD')
doc.add_paragraph('• Notreserve: 100.000 NZD')
doc.add_paragraph('• GESAMT: ca. 810.000-1.310.000 NZD (~460.000-745.000 EUR)')
doc.add_paragraph('• MIT EIGENKAPITAL FAMILIE 2: Machbar mit guter Reserve!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Entfernung zu Europa (24h+ Flug). Zeitzone für Remote-Work nach DE schwierig. Grundstücke teurer als Südamerika. Familienbesuche erfordern Planung.')

doc.add_paragraph()

# =============================================================================
# PLATZ 2: SPANIEN - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 2: Spanien Süden (38 Punkte / 76%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Costa de la Luz (Huelva), Almería Hinterland, Murcia')

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM SPANIEN FÜR EUER PROFIL SEHR GUT IST:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Vorteil + Nähe: ').bold = True
p.add_run('Ihr bleibt im EU-System mit allen Rechten. Flug nach Deutschland in 2-3 Stunden. Familie besuchen ist ein Wochenendtrip. Krankenversicherung, Rentenansprüche - alles bleibt einfach.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~113.000 Verkündiger in ~1.500 Versammlungen - eine der größten Gemeinschaften Europas. BONUS: An der Costa del Sol gibt es auch deutschsprachige Versammlungen. Aber ihr könnt mit eurem Englisch auch spanische Versammlungen besuchen und euch schnell integrieren.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('AM RAND EUROPAS - Spanien liegt weit weg von der NATO-Ostflanke. Bei einem europäischen Konflikt weniger exponiert als Deutschland. Aber: Immer noch EU und NATO-Mitglied.')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('MITTEL - Als NATO-Mitglied theoretisch betroffen, aber geographisch am Rand. Keine direkte Grenze zu Russland. Lebensmittel-Selbstversorgung gut (Obst, Gemüse, Olivenöl). ABER: Abhängig von Energieimporten, Preise würden steigen. Innere Unruhen möglich bei Wirtschaftskrise. Besser als Deutschland, schlechter als Südamerika.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('3.000+ Sonnenstunden im Süden - mehr als fast überall in Europa! Milde Winter (10-15°C), heiße Sommer (bis 40°C im Landesinneren). Perfekt für Solar. Anbausaison fast ganzjährig. Mittelmeer oder Atlantik erreichbar.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('SEHR GUT - Lange Anbausaison, viel Sonne. Oliven, Zitrusfrüchte, Gemüse wachsen hervorragend. ABER: Wasser kann in manchen Regionen knapp sein - eigener Brunnen und Zisterne wichtig. Mit Planung aber machbar.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. EU-Standard in allem. Supermärkte, Internet, Infrastruktur. Deutsche Produkte erhältlich. Große deutsche Expat-Community, besonders an der Küste.')

doc.add_paragraph()

# RECHTSSICHERHEIT
p = doc.add_paragraph()
p.add_run('Rechtssicherheit/Eigentum: ').bold = True
p.add_run('EU-STANDARD - Solides Rechtssystem. Eigentumsrechte geschützt. Notare für Kaufverträge. Bürokratie langsamer als in Deutschland, aber funktioniert.')

doc.add_paragraph()

# GESUNDHEIT
p = doc.add_paragraph()
p.add_run('Gesundheitsversorgung: ').bold = True
p.add_run('SEHR GUT - Öffentliches Gesundheitssystem mit EU-Karte nutzbar. Gute Krankenhäuser auch in ländlichen Gebieten. Private Zusatzversicherung für schnellere Termine möglich.')

doc.add_paragraph()

# PFLEGE IM ALTER
p = doc.add_paragraph()
p.add_run('Pflege im Alter: ').bold = True
p.add_run('EU-STANDARD - Gute Pflegeheime (Residencias) verfügbar. Mit Kapital: Private Seniorenresidenzen mit Pool, Garten, deutschsprachigem Personal in Expat-Gebieten! Kosten: 2.000-4.000 EUR/Monat für Vollzeitpflege. Viele deutsche Rentner nutzen Spanien für den Ruhestand. Häusliche Pflege auch günstig (Pflegekräfte aus Osteuropa/Lateinamerika).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Remote-Work: ').bold = True
p.add_run('PERFEKT - Gleiche Zeitzone wie Deutschland. Glasfaser gut ausgebaut, auch in ländlichen Gebieten. Starlink als Backup.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt: ').bold = True
p.add_run('Schwieriger als in NZ/AU wegen hoher Arbeitslosigkeit in Spanien. ABER: In Expat-Gebieten Vorteile durch Deutsch + Englisch (Tourismus, Immobilien, Handwerk). Mit Eigenkapital: Eigenes Business möglich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True
doc.add_paragraph('• Finca 10 ha mit Haus (renovierungsbedürftig) im Hinterland: 150.000-300.000 EUR')
doc.add_paragraph('• Renovierung/Ausbau für zwei Familien: 80.000-120.000 EUR')
doc.add_paragraph('• Solar + Infrastruktur: 25.000-35.000 EUR')
doc.add_paragraph('• Reserve: 60.000+ EUR')
doc.add_paragraph('• GESAMT: ca. 315.000-515.000 EUR - gute Reserve bei eurem Budget!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Sommer sehr heiß (40+ Grad im Landesinneren). Spanische Bürokratie langsamer als deutsche. Spanisch lernen nötig (12-18 Monate bis B2).')

doc.add_paragraph()

# =============================================================================
# ZYPERN - NEU!
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 4: ZYPERN (37 Punkte / 74%) - NEU HINZUGEFÜGT')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Paphos (beliebteste Expat-Region), Limassol, Larnaca, Troodos-Gebirge')

doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('WARUM ZYPERN FÜR EUER PROFIL INTERESSANT IST:').bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('ENGLISCH WEIT VERBREITET: ').bold = True
p.add_run('Zypern war bis 1960 britische Kolonie! Englisch ist zweite Amtssprache und wird überall verstanden. Besonders in Paphos und Limassol kann man problemlos nur mit Englisch leben. Das ist ein RIESENVORTEIL für euch!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EU-Mitglied mit Steuervorteil: ').bold = True
p.add_run('Volle EU-Freizügigkeit und Rechte. Non-Dom-Status: Keine Steuern auf ausländische Dividenden und Zinsen für 17 Jahre! Für Kapitalanleger sehr attraktiv. 12,5% Körperschaftsteuer für Unternehmen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas in Zypern: ').bold = True
p.add_run('Kleine, aber aktive Gemeinschaft (~1.500-2.000 Verkündiger). Versammlungen auf Englisch vorhanden! Auch griechische Versammlungen. Durch die Größe der Insel sind Versammlungen gut erreichbar.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('GETEILT - Die Insel ist seit 1974 geteilt. Der Norden (ca. 37%) ist türkisch besetzt. Die Republik Zypern (Süden) ist EU-Mitglied und stabil. Die Teilung ist "eingefroren" - kein aktiver Konflikt, aber auch keine Lösung in Sicht. Praktisch kein Risiko für Expats im Süden.')

doc.add_paragraph()

# KRISE
p = doc.add_paragraph()
p.add_run('BEI RUSSLAND-NATO-KONFLIKT: ').bold = True
p.add_run('MITTEL - Zypern ist EU-Mitglied aber KEIN NATO-Mitglied! Das ist ein Vorteil. Liegt nahe am Nahen Osten, aber nicht in unmittelbarer Konfliktzone. ABER: Kleine Insel, stark abhängig von Importen (Energie, viele Lebensmittel). Bei globaler Krise würden Versorgungsengpässe möglich sein. Selbstversorgung begrenzt durch Größe und Wasserknappheit.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('HERVORRAGEND - 3.300+ Sonnenstunden pro Jahr, eines der sonnigsten EU-Länder! Mediterranes Klima: Warme, trockene Sommer (25-35°C), milde Winter (10-17°C). Schnee nur im Troodos-Gebirge. Perfekt wenn ihr Sonne liebt.')

doc.add_paragraph()

# SELBSTVERSORGUNG
p = doc.add_paragraph()
p.add_run('Selbstversorgung: ').bold = True
p.add_run('BEGRENZT - Wasserknappheit ist ein Problem. Die Insel hat kaum natürliche Wasserreserven. Entsalzungsanlagen versorgen die Bevölkerung. Eigener Anbau möglich, aber mit Bewässerung. Große Grundstücke begrenzt verfügbar.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. Britischer Einfluss überall sichtbar: Linksverkehr, englische Supermärkte (M&S, Debenhams), britische Produkte. Große britische Expat-Community (ca. 60.000). Infrastruktur modern.')

doc.add_paragraph()

# RECHTSSICHERHEIT
p = doc.add_paragraph()
p.add_run('Rechtssicherheit/Eigentum: ').bold = True
p.add_run('EU-STANDARD - Britisches Rechtssystem als Basis. Eigentumsrechte gut geschützt. ABER: Vorsicht beim Kauf - türkische Ansprüche auf einige Grundstücke im Süden möglich. Immer mit Anwalt prüfen!')

doc.add_paragraph()

# GESUNDHEIT
p = doc.add_paragraph()
p.add_run('Gesundheitsversorgung: ').bold = True
p.add_run('GUT - GESY (allgemeines Gesundheitssystem seit 2019). Private Krankenhäuser in guter Qualität. Viele Ärzte sprechen Englisch (oft in UK ausgebildet). Mit Kapital: Exzellente private Versorgung.')

doc.add_paragraph()

# LEBENSERWARTUNG
p = doc.add_paragraph()
p.add_run('Lebenserwartung: ').bold = True
p.add_run('81 Jahre - auf EU-Niveau. Gute Grundversorgung, mediterraner Lebensstil mit gesunder Ernährung.')

doc.add_paragraph()

# PFLEGE IM ALTER
p = doc.add_paragraph()
p.add_run('Pflege im Alter: ').bold = True
p.add_run('EU-STANDARD - Pflegeheime vorhanden, oft mit englischsprachigem Personal. Private Seniorenresidenzen verfügbar. Kosten günstiger als in UK oder Nordeuropa (ca. 2.000-3.500 EUR/Monat). Viele britische Rentner verbringen ihren Lebensabend auf Zypern. Häusliche Pflege auch eine Option.')

doc.add_paragraph()

# REMOTE-WORK
p = doc.add_paragraph()
p.add_run('Remote-Work: ').bold = True
p.add_run('GUT - Nur +1h zu Deutschland (noch besser als Spanien um 1h). Glasfaser in Städten gut ausgebaut. Starlink verfügbar. Digital Nomad Visa verfügbar.')

doc.add_paragraph()

# FAMILIE 2
p = doc.add_paragraph()
p.add_run('Familie 2 - Arbeitsmarkt: ').bold = True
p.add_run('Tourismus-basiert. Viele Jobs im Hospitality-Bereich, besonders in Paphos/Limassol. Mit Englisch + Deutsch Vorteile bei internationalen Unternehmen. Viele Offshore-Firmen mit Englisch als Arbeitssprache.')

doc.add_paragraph()

# KOSTEN
p = doc.add_paragraph()
p.add_run('Konkrete Zahlen:').bold = True
doc.add_paragraph('• Grundstück 1-3 ha (größer schwer zu finden): 100.000-250.000 EUR')
doc.add_paragraph('• Haus/Villa bauen oder kaufen: 200.000-400.000 EUR')
doc.add_paragraph('• Lebenshaltungskosten: Niedriger als Nordeuropa, ähnlich wie Spanien')
doc.add_paragraph('• Keine Erbschaftssteuer!')
doc.add_paragraph('• GESAMT für Projekt: ca. 350.000-650.000 EUR')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Nachteile: ').bold = True
p.add_run('Geteilte Insel (psychologischer Faktor). Kleine Insel - kann nach Jahren "eng" wirken. Wasserknappheit limitiert Selbstversorgung. Große Grundstücke (10+ ha) schwer zu finden. Nahost-Nähe bei Krise ein Risiko. Heiße Sommer (35°C+).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('FAZIT ZYPERN: ').bold = True
p.add_run('SEHR GUTE WAHL wenn Englisch + EU + Sonne Priorität haben und Selbstversorgung weniger wichtig ist. Besser als Kanaren für Englisch-Sprecher. Steuervorteil mit Non-Dom-Status. Kleinere ZJ-Gemeinschaft als Spanien, aber Versammlungen auf Englisch!')

doc.add_paragraph()

# =============================================================================
# PLATZ 3: KANAREN - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 5: Kanarische Inseln (37 Punkte / 74%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Inseln: ').bold = True
p.add_run('Teneriffa Süd, Gran Canaria Süd, La Palma, Fuerteventura')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Bestes Klima Europas: ').bold = True
p.add_run('Ganzjährig 18-28 Grad, keine Extreme. 2.800 Sonnenstunden - perfekt. Kein heißer Sommer wie auf dem Festland, kein kalter Winter. Ideal für Gesundheit und Wohlbefinden.')

doc.add_paragraph()

# ZJ
p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~8.000 Verkündiger auf den Inseln. Auf Teneriffa und Gran Canaria gibt es auch deutschsprachige Versammlungen als Option. Mit eurem Englisch könnt ihr aber auch direkt spanische Versammlungen besuchen.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('WEIT WEG - Die Kanaren liegen vor Afrika im Atlantik, weit weg vom europäischen Festland. Bei einem Konflikt in Europa relativ geschützt.')

doc.add_paragraph()

# WESTLICHER LEBENSSTIL
p = doc.add_paragraph()
p.add_run('Westlicher Lebensstil: ').bold = True
p.add_run('VOLL westlich. EU-Rechte, deutsche Supermärkte (Lidl, Aldi), deutsches Fernsehen. Große deutsche Community.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Das Problem - Grundstücke: ').bold = True
p.add_run('Große Grundstücke sind begrenzt und teuer. Die Inseln sind klein. 10+ Hektar in Küstennähe sind kaum zu finden oder unbezahlbar.')

doc.add_paragraph()

# =============================================================================
# PLATZ 4: AUSTRALIEN - VOLLSTÄNDIG
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 3: Australien (38 Punkte / 76%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Tasmanien (beste Wahl!), Sunshine Coast (Queensland), Victoria')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Euer Englisch öffnet Türen: ').bold = True
p.add_run('Data Architect ist auf der Australian Skilled Occupation List. Mit fließendem Englisch sofortiger Zugang zu allem.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~68.000 Verkündiger in ~790 Versammlungen. Gut organisierte, aktive Gemeinschaft.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('AUKUS-MITGLIED - Bei einem Konflikt mit China potentiell exponiert. Weniger neutral als Neuseeland.')

doc.add_paragraph()

# KLIMA
p = doc.add_paragraph()
p.add_run('Klima: ').bold = True
p.add_run('EXTREME möglich - Buschbrände, Dürren, Hitzewellen auf dem Festland. Tasmanien ist die Ausnahme: Gemäßigtes Klima ähnlich Neuseeland.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Familie 2 - SUPER Arbeitsmarkt: ').bold = True
p.add_run('Großer Arbeitskräftemangel! Handwerk sehr gefragt mit guten Löhnen.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EMPFEHLUNG: ').bold = True
p.add_run('Wenn Australien, dann TASMANIEN! Beste Region für Homestead - Klima ähnlich NZ, weniger Extreme, günstiger als Festland.')

doc.add_paragraph()

# =============================================================================
# PLATZ 5: COSTA RICA
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 6-7: Costa Rica & Uruguay (39 Punkte / 78%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Costa Rica: ').bold = True
p.add_run('KEIN MILITÄR seit 1948! Kann sich nicht an Kriegen beteiligen. Höchster ZJ-Bevölkerungsanteil (0,54%). Bei NATO-Krise: Sehr gut geschützt.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Uruguay: ').bold = True
p.add_run('NEUTRAL, keine Bündnisse! Agrar-Exporteur - selbstversorgend. "Schweiz Südamerikas". Bei NATO-Krise: Ideal geschützt. ABER: Arbeitsmarkt für Familie 2 sehr begrenzt.')

doc.add_paragraph()

# =============================================================================
# NAMIBIA
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Platz 11: Namibia (26 Punkte / 52%)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 112, 192)

p = doc.add_paragraph()
p.add_run('Beste Regionen: ').bold = True
p.add_run('Windhoek Umgebung, Swakopmund, Omaruru')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Englisch als Amtssprache: ').bold = True
p.add_run('Namibia ist eines der wenigen afrikanischen Länder mit Englisch als Amtssprache! Außerdem gibt es eine deutschsprachige Minderheit (~20.000).')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Zeugen Jehovas: ').bold = True
p.add_run('~3.100 Verkündiger. Kleine, aber aktive Gemeinschaft. Versammlungen auf Englisch.')

doc.add_paragraph()

# GEOPOLITIK
p = doc.add_paragraph()
p.add_run('Geopolitische Sicherheit: ').bold = True
p.add_run('STABIL - Friedliche Demokratie seit 1990. Keine Konflikte, keine Militärbündnisse.')

doc.add_paragraph()

# LEBENSERWARTUNG
p = doc.add_paragraph()
p.add_run('Lebenserwartung: ').bold = True
p.add_run('NUR 66 JAHRE - Deutlich niedriger als Europa oder NZ/AU. Dies ist ein wichtiger Nachteil!')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Grundstücke: ').bold = True
p.add_run('SEHR GÜNSTIG - Große Farmen (1.000+ ha!) für 100.000-300.000 EUR möglich.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('FAZIT: ').bold = True
p.add_run('Interessant wegen Englisch + Deutsch + günstigen Preisen. ABER: Die niedrige Lebenserwartung und Wasserprobleme sind ernste Nachteile. Nur empfohlen wenn ihr Abenteuer sucht.')

doc.add_paragraph()

# =============================================================================
# PLATZ 9-12: NICHT EMPFOHLEN
# =============================================================================
h2 = doc.add_paragraph()
run = h2.add_run('Nicht empfohlen: Deutschland, Schweden, Nicaragua, Nigeria')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# DEUTSCHLAND
p = doc.add_paragraph()
p.add_run('DEUTSCHLAND: ').bold = True
p.add_run('NATO-FRONTSTAAT - Bei einem Konflikt mit Russland direkt exponiert. US-Militärbasen = primäre Ziele. 100% abhängig von Energieimporten.')

doc.add_paragraph()

# SCHWEDEN
p = doc.add_paragraph()
p.add_run('SCHWEDEN: ').bold = True
p.add_run('NEUES NATO-MITGLIED nahe Russland. Kalt und dunkel - lange Winter.')

doc.add_paragraph()

# NICARAGUA
p = doc.add_paragraph()
p.add_run('NICARAGUA: ').bold = True
p.add_run('ORTEGA-REGIME - Autoritäre Regierung. Schwache Rechtssicherheit. Wenn Zentralamerika, dann COSTA RICA.')

doc.add_paragraph()

# NIGERIA
p = doc.add_paragraph()
p.add_run('NIGERIA: ').bold = True
p.add_run('NICHT EMPFOHLEN - Massive Sicherheitsprobleme, Korruption. ZJ-Gemeinschaft sehr groß (~407.000!) aber das wiegt Risiken NICHT auf.')

doc.add_page_break()

# =============================================================================
# VERGLEICH ZYPERN vs KANAREN vs SPANIEN
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('VERGLEICH: ZYPERN vs. KANAREN vs. SPANIEN SÜDEN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

comparison_data = [
    ['Kriterium', 'Zypern', 'Kanaren', 'Spanien Süd'],
    ['Englisch', '++ Überall!', '- Spanisch nötig', '- Spanisch nötig'],
    ['Sonne', '++ 3.300h', '++ 2.800h', '++ 3.000h'],
    ['Klima', '++ Mittelmeer heiß', '++ Perfekt mild', '+ Sommerhitze'],
    ['EU-Rechte', '++ Ja', '++ Ja', '++ Ja'],
    ['Steuern', '++ Non-Dom!', '+ IGIC 7%', 'o Normal'],
    ['ZJ Englisch', '++ Ja', '- Nein', '- Nein (dt. ja)'],
    ['Grundstücke 10ha', '- Schwer', '-- Kaum', '+ Möglich'],
    ['Selbstversorgung', '- Begrenzt', '- Begrenzt', '+ Gut'],
    ['Nähe DE (Flug)', '+ 3-4h', '+ 4h', '++ 2-3h'],
    ['NATO-Mitglied', '- Nein!', '+ Ja (aber weit)', '+ Ja (aber weit)'],
    ['Arbeitsmarkt F2', 'o Tourismus', 'o Tourismus', 'o Begrenzt'],
]
create_table_original_style(doc, comparison_data)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('EMPFEHLUNG:').bold = True

doc.add_paragraph()

doc.add_paragraph('• Wenn ENGLISCH Priorität hat: ZYPERN > Kanaren/Spanien')
doc.add_paragraph('• Wenn KLIMA (mild) Priorität hat: KANAREN > Zypern > Spanien')
doc.add_paragraph('• Wenn SELBSTVERSORGUNG Priorität hat: SPANIEN > Kanaren > Zypern')
doc.add_paragraph('• Wenn STEUERN Priorität haben: ZYPERN (Non-Dom) > Kanaren (IGIC)')

doc.add_paragraph()

# =============================================================================
# EMPFEHLUNG
# =============================================================================
h = doc.add_paragraph()
run = h.add_run('UNSERE EMPFEHLUNG FÜR EUCH')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ERSTE WAHL: NEUSEELAND')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Euer Profil ist wie gemacht für Neuseeland:')
doc.add_paragraph('• Fließend Englisch = Sofortige Integration ohne Sprachbarriere')
doc.add_paragraph('• IT-Skills (F1) = Skilled Migrant Visa realistisch')
doc.add_paragraph('• Arbeitskräftemangel = Familie 2 findet Vor-Ort-Jobs')
doc.add_paragraph('• Aktive ZJ-Gemeinschaft (~14.000)')
doc.add_paragraph('• MAXIMALE geopolitische Sicherheit (isoliert, neutral)')
doc.add_paragraph('• Einziger Kompromiss: Entfernung zu Europa')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE A: ZYPERN (NEU!)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn EU + Englisch + Nähe zu DE Priorität haben:')
doc.add_paragraph('• Englisch überall verbreitet - keine Sprachbarriere!')
doc.add_paragraph('• EU-Rechte + Non-Dom Steuervorteil')
doc.add_paragraph('• 3-4h Flug nach Deutschland')
doc.add_paragraph('• ZJ-Versammlungen auf Englisch')
doc.add_paragraph('• Bestes Sonnenwetter in der EU')
doc.add_paragraph('• Kein NATO-Mitglied!')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('ALTERNATIVE B: SPANIEN SÜDEN / KANAREN')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 112, 192)

doc.add_paragraph('Wenn Selbstversorgung + große Grundstücke Priorität haben:')
doc.add_paragraph('• Mehr Platz als auf Zypern')
doc.add_paragraph('• Spanisch lernen nötig (12-18 Monate)')
doc.add_paragraph('• ZJ deutschsprachige Versammlungen vorhanden')

doc.add_paragraph()

doc.add_paragraph('─' * 70)

p = doc.add_paragraph()
p.add_run('Erstellt: Januar 2025 | Basierend auf Auswanderungsanalyse 2025').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Erweitert: ZJ | Englisch | Eigenkapital F2 | Vor-Ort-Jobs | ZYPERN NEU').font.size = Pt(9)
p = doc.add_paragraph()
p.add_run('Familie 1: 40J, 2 Kinder (12+14) | Familie 2: 50J, keine Kinder').font.size = Pt(9)

# Save
doc.save(r'C:\Project A\Auswanderungsanalyse_2025_MIT_ZYPERN.docx')
print("Word-Dokument erfolgreich erstellt!")
print(r"Gespeichert unter: C:\Project A\Auswanderungsanalyse_2025_MIT_ZYPERN.docx")











